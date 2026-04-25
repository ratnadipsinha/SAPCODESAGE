import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')

def find_idx(fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def insert_para_before(ref_idx, text, style='Heading 1'):
    ref = doc.paragraphs[ref_idx]
    new_p = OxmlElement('w:p')
    ref._element.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            try:
                p.style = doc.styles[style]
            except:
                pass
            p.add_run(text)
            return p
    return None

def append_section(anchor_fragment, heading, lines):
    """Append heading + lines after the paragraph containing anchor_fragment."""
    idx = find_idx(anchor_fragment)
    if idx < 0:
        # fallback: append at end before last paragraph
        idx = len(doc.paragraphs) - 2

    ref = doc.paragraphs[idx]
    items = [(heading, 'Heading 1')] + [(l, 'Heading 2' if l.startswith('###') else 'List Paragraph' if l.startswith('  ') else '') for l in lines]

    prev = ref
    for text, style in items:
        text = text.lstrip('#').strip()
        if not text:
            continue
        new_p = OxmlElement('w:p')
        prev._element.addnext(new_p)
        for p in doc.paragraphs:
            if p._element is new_p:
                try:
                    if style:
                        p.style = doc.styles[style]
                except:
                    pass
                p.add_run(text)
                prev = p
                break

# ── Anchor: insert after last existing section (Compliance Checklist) ─────────
anchor = 'Who accessed what'

heading = '7. SAP Clean Core & ABAP Cloud Compatibility'

content = [
    '### What is Clean Core?',
    'SAP Clean Core is a mandatory architecture principle for S/4HANA Cloud. It restricts how ABAP extensions are written, moving away from system modifications toward released APIs, RAP, and BAdI implementations.',
    '',
    '### Does the Base Model Know Clean Core?',
    'CodeLlama has partial knowledge only. Critical gaps exist in RAP, released API classification, ABAP Cloud profile restrictions, and BTP extensions. Without addressing these gaps, the assistant risks suggesting non-compliant code that would fail a Clean Core audit.',
    '',
    '### Clean Core Topics the Assistant Actively Teaches',
    '  CDS Views: Define, expose, and consume Core Data Services. Basic, Composite, Consumption, and Interface views. UI, OData, and Search annotations.',
    '  EML (Entity Manipulation Language): READ ENTITIES, MODIFY ENTITIES, DELETE ENTITIES, COMMIT ENTITIES — replacing direct DB statements.',
    '  Released APIs: How to find C1/C2 classified APIs on api.sap.com. Alternatives when a classic FM has no released equivalent.',
    '  BAdI / Enhancement Framework: Implementing BAdIs correctly in S/4HANA Cloud. Switching from classic user exits to Enhancement Spots.',
    '  RAP (ABAP RESTful Application Model): Full guidance — CDS data model, Behaviour Definition, Behaviour Implementation, OData service binding, Fiori app.',
    '  Key User Extensibility: When to use no-code KUE instead of ABAP — custom fields, custom logic, custom apps via Fiori.',
    '  BTP Side-by-Side Extensions: When to move logic to BTP. Calling S/4HANA released APIs from Node.js or Python on BTP.',
    '',
    '### Deprecated & Non-Released API Handling',
    'The assistant warns automatically when a suggested API is deprecated, non-released, or restricted:',
    '',
    '  CONVERSION_EXIT_ALPHA_INPUT — NOT released for cloud (no C1/C2). Risk: syntax error in ABAP Cloud profile. Alternative: use string template |{ lv_value ALPHA = IN }|',
    '  RFC_READ_TABLE — DEPRECATED. Removed from released API list. Performance issues, no cloud support. Alternative: CDS view with @OData.publish or RAP-based OData V4.',
    '  MODIFY ekko FROM TABLE lt_ekko — FORBIDDEN in Clean Core. Bypasses SAP business logic. Alternative: BAPI_PO_CHANGE (C1 released) or EML MODIFY ENTITIES.',
    '',
    'The assistant maintains a local deprecation index built from: SAP Released API catalogue (api.sap.com), SAP Notes on deprecated FMs, and your project own deprecated Z-objects list.',
    '',
    '### Clean Core Compliance Checker',
    'Developer pastes code — assistant scans and reports:',
    '  Line 12: CLASS ZCL_HANDLER — Clean Core compliant',
    '  Line 28: CALL FUNCTION ENQUE_READ — verify C1 release status',
    '  Line 45: MODIFY SAPMV45A — FORBIDDEN in Clean Core',
    '  Line 67: SELECT * FROM T001W — use I_Plant CDS view instead',
    '  Overall: NOT CLEAN CORE COMPLIANT — 2 critical issues. Suggested fixes shown.',
    '',
    '### Model Strategy by Landscape',
    '  ECC / On-Premise S/4: CodeLlama-7B + SAP Help + BAPI docs. Focus: Classic ABAP patterns.',
    '  S/4HANA Cloud (Clean Core): CodeLlama-7B + Clean Core + RAP + api.sap.com. Focus: RAP, BAdI, released APIs.',
    '  Hybrid (moving to cloud): Both RAG layers combined. Focus: Both patterns + migration guidance.',
    '  BTP Side-by-Side: CodeLlama-7B + BTP docs + CAP model. Focus: Node.js/Python on BTP.',
    '',
    '### Additional SAP Docs to Index for Clean Core',
    '  ABAP Cloud Development Guide — help.sap.com: ABAP Cloud profile, restrictions',
    '  RAP Programming Model — help.sap.com: CDS, Behaviour Definitions, OData',
    '  Released API Catalogue — api.sap.com: C1/C2 classification lookup',
    '  Clean Core Extensibility Guide — help.sap.com: Tier 1/2/3 model',
    '  Key User Extensibility — help.sap.com: Fiori-based no-code extensions',
    '  BTP Extension Patterns — help.sap.com/btp: Side-by-side architecture',
]

append_section(anchor, heading, content)

doc.save(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')
print('Clean Core section added and saved.')
