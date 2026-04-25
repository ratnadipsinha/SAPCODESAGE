import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = Document(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')

def find_idx(fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def append_paras_after(anchor_fragment, paragraphs):
    """Append list of (text, style) tuples after the anchor paragraph."""
    idx = find_idx(anchor_fragment)
    if idx < 0:
        idx = len(doc.paragraphs) - 2
    prev = doc.paragraphs[idx]
    for text, style in paragraphs:
        if not text.strip():
            text = ' '
        new_p = OxmlElement('w:p')
        prev._element.addnext(new_p)
        for p in doc.paragraphs:
            if p._element is new_p:
                try:
                    p.style = doc.styles[style]
                except:
                    pass
                p.add_run(text)
                prev = p
                break

# ── Find last paragraph of doc to append after ───────────────────────────────
anchor = 'BTP Extension Patterns'

clean_core_paras = [
    # Section heading
    ('7. SAP Clean Core & ABAP Cloud Compatibility', 'Heading 1'),

    ('7.1  What is Clean Core?', 'Heading 2'),
    ('SAP Clean Core is a mandatory architecture principle for S/4HANA Cloud and increasingly expected in on-premise S/4HANA. It restricts how ABAP extensions are written — moving away from system modifications toward released APIs, RAP, and BAdI implementations.', ''),
    ('Traditional ABAP (ECC/On-Prem): Modify SAP tables, user exits, unrestricted Z-programs, direct DB SELECT on SAP tables, non-released FMs, classic BAPIs.', 'List Paragraph'),
    ('Clean Core (S/4HANA Cloud): All of the above are FORBIDDEN or restricted. Allowed: BAdI implementations, ABAP RAP, Key User Extensibility, BTP side-by-side, released APIs only.', 'List Paragraph'),

    ('7.2  Base Model Knowledge Gaps', 'Heading 2'),
    ('CodeLlama has partial Clean Core knowledge. Critical gaps exist in: ABAP RAP, released API classification (C1/C2), ABAP Cloud profile restrictions, BTP extensions, and Key User Extensibility. Without addressing these gaps, the assistant risks suggesting non-compliant code that would fail a Clean Core audit.', ''),

    ('7.3  Clean Core Topics the Assistant Actively Teaches', 'Heading 2'),
    ('CDS Views: Define, expose, consume Core Data Services. Basic, Composite, Consumption, Interface views. Annotations: @UI, @OData, @Search.', 'List Paragraph'),
    ('EML (Entity Manipulation Language): READ ENTITIES, MODIFY ENTITIES, DELETE ENTITIES, COMMIT ENTITIES — replacing all direct DB statements in RAP context.', 'List Paragraph'),
    ('Released APIs: Find C1/C2 classified APIs on api.sap.com. Check release status before use. Alternatives when a classic FM has no released equivalent.', 'List Paragraph'),
    ('BAdI / Enhancement Framework: Implement BAdIs correctly in S/4HANA Cloud. Switch from classic user exits to Enhancement Spots. Find the right Enhancement Spot for each business process.', 'List Paragraph'),
    ('RAP (ABAP RESTful Application Model): Full guidance — CDS data model, Behaviour Definition, Behaviour Implementation, OData service binding, Fiori app.', 'List Paragraph'),
    ('Key User Extensibility (KUE): When to use no-code KUE instead of ABAP — custom fields, custom logic, custom apps via Fiori.', 'List Paragraph'),
    ('BTP Side-by-Side Extensions: When to move logic to BTP. Calling S/4HANA released APIs from Node.js or Python on BTP.', 'List Paragraph'),

    ('7.4  Cloudification Repository Viewer — Primary Reference', 'Heading 2'),
    ('The SAP Cloudification Repository Viewer is SAP\'s official tool that classifies every SAP API, FM, BAPI, table, and object by its cloud readiness status. It is the single most important source for Clean Core compliance checking.', ''),
    ('For every SAP object it provides: Release status (Released / Not Released / Deprecated), Contract level (C1 stable / C2 use with caution), Successor API recommended by SAP, Availability (S/4HANA Cloud / BTP ABAP Env / On-Prem), and Migration guidance.', 'List Paragraph'),
    ('Access points: ADT (Eclipse) — right-click object, Where Used, Cloud Classification. SAP web viewer via BTP ABAP Environment. Programmatic via /UI2/CL_CR_VIEWER or SAP ATC checks.', 'List Paragraph'),
    ('Example: Developer asks to use ADDRESS_INTO_PRINTFORM. CodeSage checks Cloudification Repository — Status: NOT RELEASED for cloud, no C1/C2 classification. Successor: CL_S4_ADDRESS_FACADE=>GET_PRINT_FORMAT (C1 released). Response: "Use CL_S4_ADDRESS_FACADE=>GET_PRINT_FORMAT — C1 released and available in S/4HANA Cloud and BTP."', 'List Paragraph'),

    ('7.5  Deprecated & Non-Released API Handling', 'Heading 2'),
    ('CONVERSION_EXIT_ALPHA_INPUT: NOT released for cloud (no C1/C2). Risk: syntax error in ABAP Cloud profile. Alternative: string template |{ lv_value ALPHA = IN }|', 'List Paragraph'),
    ('RFC_READ_TABLE: DEPRECATED — removed from released API list. Performance issues, no cloud support. Alternative: CDS view with @OData.publish annotation or RAP-based OData V4.', 'List Paragraph'),
    ('MODIFY ekko FROM TABLE lt_ekko: FORBIDDEN in Clean Core. Bypasses SAP business logic. Alternative: BAPI_PO_CHANGE (C1 released) or EML MODIFY ENTITIES OF i_purchaseorder.', 'List Paragraph'),

    ('7.6  Everything Stays Local — Including Clean Core Knowledge', 'Heading 2'),
    ('A critical design principle: no Clean Core check ever calls an external service at query time. All knowledge is downloaded once during setup and stored in local ChromaDB collections.', ''),
    ('Cloudification Repository (exported CSV from ADT) → parsed and embedded into ChromaDB collection: cloudification_db', 'List Paragraph'),
    ('api.sap.com Released API Catalogue → scraped and chunked during setup → stored in ChromaDB collection: released_apis', 'List Paragraph'),
    ('SAP Help pages (RAP, Clean Core, BTP docs) → downloaded and indexed → stored in ChromaDB collection: sap_help_docs', 'List Paragraph'),
    ('Your project deprecated Z-objects → maintained in scan_config.yaml → checked locally at every code suggestion', 'List Paragraph'),
    ('At query time: ZERO external calls. Everything answered from local ChromaDB + local fine-tuned model. No SAP connection. No internet. No external API call.', ''),

    ('7.7  Model Strategy by Client Landscape', 'Heading 2'),
    ('ECC / On-Premise S/4: CodeLlama-7B + SAP Help + BAPI docs. Focus: Classic ABAP patterns.', 'List Paragraph'),
    ('S/4HANA Cloud (Clean Core): CodeLlama-7B + Clean Core + RAP + Cloudification Repository + api.sap.com. Focus: RAP, BAdI, released APIs.', 'List Paragraph'),
    ('Hybrid (moving to cloud): Both RAG layers combined. Focus: Both patterns + migration guidance from Cloudification Repository.', 'List Paragraph'),
    ('BTP Side-by-Side: CodeLlama-7B + BTP docs + CAP model. Focus: Node.js/Python on BTP + S/4 released API consumption.', 'List Paragraph'),
]

append_paras_after(anchor, clean_core_paras)

doc.save(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')
print('Clean Core + Cloudification Repository section added and saved.')
