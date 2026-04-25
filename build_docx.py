"""
Builds CodeSage_Proposal.docx from proposal.md
- Fully editable Word document (native charts, styled tables, text)
- Phase diagrams as editable Word tables
- Problem chart as native Word bar chart (right-click → Edit Data)
"""

import re, io
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml

# ═══════════════════════════════════════════════════════════════
#  DATA CHART TABLE  (editable Word table styled as bar chart)
# ═══════════════════════════════════════════════════════════════

def add_native_bar_chart(doc, categories, series1_vals, series2_vals,
                          s1_name, s2_name, chart_title,
                          width_in=6.3, height_in=3.6):
    """
    Inserts an editable Word table styled as a horizontal comparison chart.
    Each row shows a coloured bar proportional to the value.
    """
    MAX_VAL = 4.0  # max hours for scaling
    BAR_COLS = 20  # number of block columns representing the bar
    S1_COLOR = 'FFCDD2'   # light red — Without
    S2_COLOR = 'C8E6C9'   # light green — With
    S1_FILL  = 'EF9A9A'   # bar fill — Without
    S2_FILL  = '81C784'   # bar fill — With

    # Title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(4)
    tr = title_p.add_run(f'  {chart_title}  ')
    tr.bold = True; tr.font.size = Pt(10.5); tr.font.name = 'Calibri'
    tr.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

    # Legend
    leg_p = doc.add_paragraph()
    leg_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    leg_p.paragraph_format.space_after = Pt(4)
    for txt, bg in [(f'  {s1_name}  ', S1_FILL), ('   ', 'FFFFFF'), (f'  {s2_name}  ', S2_FILL)]:
        run = leg_p.add_run(txt)
        run.font.size = Pt(9); run.font.name = 'Calibri'; run.bold = True

    # Table: header + one row per category (2 sub-rows: s1 and s2)
    # Columns: [Activity label | bar blocks... | value]
    LABEL_W = Cm(4.0)
    BAR_W   = Cm(0.28)
    VAL_W   = Cm(1.0)

    # Header row
    hdr_p = doc.add_paragraph()
    hdr_p.paragraph_format.space_after = Pt(2)
    hr = hdr_p.add_run(f'  Activity{" " * 30}{s1_name} / {s2_name} (hours per day, max {MAX_VAL:.0f}h)  ')
    hr.bold = True; hr.font.size = Pt(8.5); hr.font.name = 'Calibri'
    hr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    def shade_para_local(p, fill):
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), fill)
        pPr.append(shd)

    shade_para_local(hdr_p, '1A3A5C')
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Data rows table
    tbl = doc.add_table(rows=len(categories) * 2 + 1, cols=BAR_COLS + 2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    def set_shd(cell, fill):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
        tcPr.append(shd)

    def no_border(cell):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for side in ('top', 'bottom', 'left', 'right'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        tcPr.append(tcBorders)

    def write_cell(cell, text, bold=False, size=8, color='222222', fill=None, align=WD_ALIGN_PARAGRAPH.LEFT):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(*bytes.fromhex(color))
        if fill:
            set_shd(cell, fill)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        no_border(cell)

    # Column widths
    for row in tbl.rows:
        row.cells[0].width = LABEL_W
        for c in range(1, BAR_COLS + 1):
            row.cells[c].width = BAR_W
        row.cells[BAR_COLS + 1].width = VAL_W

    # Scale: how many bar columns to fill
    def bar_len(val): return max(1, round((val / MAX_VAL) * BAR_COLS))

    # Sub-header row (row 0)
    write_cell(tbl.cell(0, 0), 'Activity', bold=True, size=8, color='FFFFFF', fill='1A3A5C')
    for c in range(1, BAR_COLS + 1):
        write_cell(tbl.cell(0, c), '', fill='1A3A5C')
    write_cell(tbl.cell(0, BAR_COLS + 1), 'hrs', bold=True, size=8, color='FFFFFF', fill='1A3A5C', align=WD_ALIGN_PARAGRAPH.CENTER)

    for ri, (cat, v1, v2) in enumerate(zip(categories, series1_vals, series2_vals)):
        row1 = ri * 2 + 1   # Without row
        row2 = ri * 2 + 2   # With row
        b1 = bar_len(v1)
        b2 = bar_len(v2)
        row_bg = 'F8FAFB' if ri % 2 == 0 else 'FFFFFF'

        # Label merged across both sub-rows
        lbl = tbl.cell(row1, 0).merge(tbl.cell(row2, 0))
        write_cell(lbl, cat, bold=True, size=8, color='1A3A5C', fill=row_bg)

        # Without row bars
        for c in range(1, BAR_COLS + 1):
            fill = S1_FILL if c <= b1 else row_bg
            write_cell(tbl.cell(row1, c), '', fill=fill)
        write_cell(tbl.cell(row1, BAR_COLS + 1), f'{v1:.1f}', size=8, color='C62828', fill=row_bg, align=WD_ALIGN_PARAGRAPH.CENTER)

        # With row bars
        for c in range(1, BAR_COLS + 1):
            fill = S2_FILL if c <= b2 else row_bg
            write_cell(tbl.cell(row2, c), '', fill=fill)
        write_cell(tbl.cell(row2, BAR_COLS + 1), f'{v2:.1f}', size=8, color='2E7D32', fill=row_bg, align=WD_ALIGN_PARAGRAPH.CENTER)

    cap = doc.add_paragraph('Red bars = Without CodeSage   |   Green bars = With CodeSage   |   Edit any cell to update values')
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after  = Pt(10)
    if cap.runs:
        cap.runs[0].font.size = Pt(8); cap.runs[0].italic = True
        cap.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)


# ═══════════════════════════════════════════════════════════════
#  EDITABLE PHASE DIAGRAMS  (Word tables — fully editable)
# ═══════════════════════════════════════════════════════════════

def set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def set_cell_border(cell, sides=('top','bottom','left','right'), color='CCCCCC', sz='4'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in sides:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    sz)
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def cell_write(cell, text, bold=False, size=9, color_hex=None,
               align=WD_ALIGN_PARAGRAPH.CENTER, bg_hex=None, italic=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    r.font.name = 'Calibri'
    if color_hex:
        r.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    if bg_hex:
        set_cell_shading(cell, bg_hex)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_arrow_row(tbl, col_count, label='▼', bg='FFFFFF'):
    """Single merged arrow row between flow steps."""
    row = tbl.add_row()
    for i, cell in enumerate(row.cells):
        if i == 0:
            cell_write(cell, label, bold=True, size=10, color_hex='888888', bg_hex=bg)
        else:
            cell_write(cell, '', bg_hex=bg)
    # Merge all cells
    merged = row.cells[0].merge(row.cells[-1])
    cell_write(merged, label, bold=True, size=10, color_hex='888888', bg_hex=bg)
    for cell in row.cells:
        set_cell_border(cell, color='FFFFFF')


def make_phase1_table(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after  = Pt(6)
    r = title.add_run('  Phase 1 — Scan & Extract: How it flows  ')
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_para(title, '1565C0')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    steps = [
        ('1  CONFIG', 'scan_config.yaml\nDefine packages, object types,\nand filters to apply', 'D0E8F7', '1565C0'),
        ('2  CONNECT', 'SAP System\nECC / S/4HANA\nread-only RFC user\nZABAP_SCANNER', 'D0E8F7', '1565C0'),
        ('3  DISCOVER', 'Object Inventory\nRS_PROGRAM_INDEX\nList all ABAP objects\nin each package', 'E8F5E9', '2E7D32'),
        ('4  EXTRACT', 'Source Code\nREAD_REPORT\nRPY_FUNCTIONMODULE_READ\nSEO_CLASS_GET_SOURCE', 'FFF3E0', 'E65100'),
        ('5  SAVE', 'Raw Files\n.abap source code\n+ .json metadata\nsidecar per object', 'F3E5F5', '6A1B9A'),
        ('6  LOG', 'Audit Trail\ngit commit per run\nTimestamp + object count\nCompliance record', 'FCE4EC', 'C62828'),
    ]

    tbl = doc.add_table(rows=1, cols=len(steps))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Header row
    for ci, (step_num, step_text, bg, fc) in enumerate(steps):
        cell = tbl.cell(0, ci)
        lines = step_text.split('\n')
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Step badge
        r = p.add_run(step_num + '\n')
        r.bold = True; r.font.size = Pt(8.5); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(*bytes.fromhex(fc))
        # Content
        r2 = p.add_run('\n'.join(lines))
        r2.font.size = Pt(8); r2.font.name = 'Calibri'
        r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        set_cell_shading(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Arrow row below
    row2 = tbl.add_row()
    for ci in range(len(steps)):
        cell = row2.cells[ci]
        arrow = '→' if ci < len(steps) - 1 else '✓'
        cell_write(cell, arrow, bold=True, size=12, color_hex='888888', bg_hex='FAFAFA')
        set_cell_border(cell, color='EEEEEE')

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph('Components are fully editable above — change text directly in the table cells.').paragraph_format.space_after = Pt(10)
    cap = doc.paragraphs[-1]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs: cap.runs[0].font.size = Pt(8.5); cap.runs[0].italic = True; cap.runs[0].font.color.rgb = RGBColor(0x88,0x88,0x88)


def make_phase2_table(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    title = doc.add_paragraph()
    r = title.add_run('  Phase 2 — Index & Embed: How it flows  ')
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_para(title, '006064')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after  = Pt(6)

    # Two-section table: inputs (left) → processing (centre) → output (right)
    tbl = doc.add_table(rows=3, cols=5)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    col_widths = [Cm(4.0), Cm(1.0), Cm(3.5), Cm(1.0), Cm(5.0)]
    for i, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[i].width = w

    # Input col items
    inputs = [
        ('Raw ABAP Files\nfrom Phase 1\n(programs, classes,\nFMs, BAdIs)', 'E8F5E9', '2E7D32'),
        ('SAP Help Pages\nRAP · BAPIs · Clean Core\nABAP Cloud guide', 'FFF3E0', 'E65100'),
        ('Cloudification Repo\nC1 / C2 / Not Released\nsuccessor API per object', 'FFEBEE', 'C62828'),
    ]
    for ri, (txt, bg, fc) in enumerate(inputs):
        cell = tbl.cell(ri, 0)
        lines = txt.split('\n')
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(lines[0])
        r.bold = True; r.font.size = Pt(8.5); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(*bytes.fromhex(fc))
        if len(lines) > 1:
            r2 = p.add_run('\n' + '\n'.join(lines[1:]))
            r2.font.size = Pt(7.5); r2.font.name = 'Calibri'
        set_cell_shading(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Arrow col
        cell_write(tbl.cell(ri, 1), '→', bold=True, size=12, color_hex='888888', bg_hex='FAFAFA')
        set_cell_border(tbl.cell(ri, 1), color='EEEEEE')

    # Merge processing col (rows 0-2)
    proc = tbl.cell(0, 2).merge(tbl.cell(2, 2))
    cell_write(proc, 'Chunker +\nEmbedder\n\nnomic-embed-text\nvia Ollama\n(CPU only)\n\n~2 hrs\n/ 10k objects',
               bold=False, size=9, color_hex='1A1A1A', bg_hex='F3E5F5')
    proc.paragraphs[0].runs[0].bold = True

    # Arrow col (merged)
    arr = tbl.cell(0, 3).merge(tbl.cell(2, 3))
    cell_write(arr, '→', bold=True, size=12, color_hex='888888', bg_hex='FAFAFA')
    set_cell_border(arr, color='EEEEEE')

    # ChromaDB output (merged)
    out = tbl.cell(0, 4).merge(tbl.cell(2, 4))
    out.text = ''
    p = out.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('ChromaDB\n')
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x00, 0x60, 0x64)
    r2 = p.add_run('4 local collections:\n\n· abap_codebase\n· sap_help_docs\n· released_apis\n· cloudification_db')
    r2.font.size = Pt(8.5); r2.font.name = 'Calibri'
    set_cell_shading(out, 'E0F7FA')
    out.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def make_phase3_table(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    title = doc.add_paragraph()
    r = title.add_run('  Phase 3 — Fine-Tune: How it flows  ')
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_para(title, '1B5E20')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after  = Pt(6)

    steps = [
        ('TRAINING DATA\nGenerate Q&A pairs\nvia Claude API\n+ Clean Core\nexamples', 'FFF3E0', 'E65100'),
        ('BASE MODEL\nCodeLlama-7B\nopen-source\npre-trained on code\nincl. ABAP', 'E3F2FD', '1565C0'),
        ('QLoRA TRAINING\nKaggle T4 GPU\n(free — 16 GB VRAM)\n10M of 7B params\n6–12 hrs', 'E8F5E9', '2E7D32'),
        ('LoRA ADAPTER\n~50 MB file\nmerged to GGUF\nformat for\nOllama', 'FFFDE7', 'F57F17'),
        ('OLLAMA SERVER\nlocalhost:11434\nfully offline\nno cloud\ndependency', 'ECEFF1', '546E7A'),
    ]

    tbl = doc.add_table(rows=1, cols=len(steps) * 2 - 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    for ci in range(len(steps) * 2 - 1):
        cell = tbl.cell(0, ci)
        if ci % 2 == 0:
            si = ci // 2
            txt, bg, fc = steps[si]
            lines = txt.split('\n')
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lines[0] + '\n')
            r.bold = True; r.font.size = Pt(8); r.font.name = 'Calibri'
            r.font.color.rgb = RGBColor(*bytes.fromhex(fc))
            r2 = p.add_run('\n'.join(lines[1:]))
            r2.font.size = Pt(7.5); r2.font.name = 'Calibri'
            r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = Cm(2.8)
        else:
            cell_write(cell, '→', bold=True, size=12, color_hex='888888', bg_hex='FAFAFA')
            set_cell_border(cell, color='EEEEEE')
            cell.width = Cm(0.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def make_phase4_table(doc):
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    title = doc.add_paragraph()
    r = title.add_run('  Phase 4 — Runtime Query: How it flows  ')
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade_para(title, 'BF360C')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after  = Pt(6)

    steps = [
        ('DEVELOPER\nAsk in plain\nEnglish via\nchat interface', 'FFEBEE', 'C62828'),
        ('STREAMLIT UI\nlocalhost:8501\nbrowser — no\ninstall needed', 'FFF3E0', 'E65100'),
        ('ChromaDB\nRETRIEVER\nsemantic search\n<200ms · top 5\nrelevant chunks', 'E0F7FA', '00838F'),
        ('RAG PROMPT\nBUILDER\ncode + SAP docs\n+ compliance\ncheck', 'F3E5F5', '7B1FA2'),
        ('Fine-tuned\nCODESAGE LLM\nanswers in your\nteam\'s style\n2–8 seconds', 'E8F5E9', '2E7D32'),
        ('RESPONSE\nAnswer + Code\n+ Citations\n+ Compliance\n✅ / ❌', 'FFFDE7', 'F9A825'),
    ]

    tbl = doc.add_table(rows=2, cols=len(steps) * 2 - 1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    # Forward flow (row 0)
    for ci in range(len(steps) * 2 - 1):
        cell = tbl.cell(0, ci)
        if ci % 2 == 0:
            si = ci // 2
            txt, bg, fc = steps[si]
            lines = txt.split('\n')
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(lines[0] + '\n')
            r.bold = True; r.font.size = Pt(8); r.font.name = 'Calibri'
            r.font.color.rgb = RGBColor(*bytes.fromhex(fc))
            r2 = p.add_run('\n'.join(lines[1:]))
            r2.font.size = Pt(7.5); r2.font.name = 'Calibri'
            set_cell_shading(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        else:
            cell_write(cell, '→', bold=True, size=11, color_hex='888888', bg_hex='FAFAFA')
            set_cell_border(cell, color='EEEEEE')

    # Return row (row 1) — response back arrow
    ret_row = tbl.rows[1]
    for ci, cell in enumerate(ret_row.cells):
        cell_write(cell, '', bg_hex='FAFAFA')
        set_cell_border(cell, color='EEEEEE')
    # Label in merged bottom row
    merged = ret_row.cells[0].merge(ret_row.cells[-1])
    cell_write(merged, '◀─────────────── response returned to developer ───────────────',
               size=8.5, color_hex='888888', bg_hex='F9F9F9', italic=True)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


def make_overview_table(doc):
    """Four-phase overview as an editable summary table for the cover."""
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'

    phases = [
        ('Phase 1\nScan & Extract', 'Half a day\nAny laptop\nCPU only', 'D0E8F7', '1565C0'),
        ('Phase 2\nIndex & Embed', '~2 hours\nAny laptop\nCPU only', 'E0F7FA', '006064'),
        ('Phase 3\nFine-Tune', '6–12 hours\nKaggle T4 GPU\n(free)', 'E8F5E9', '1B5E20'),
        ('Phase 4\nRuntime Query', 'Live forever\n2–8 sec response\nfully offline', 'FFF3E0', 'BF360C'),
    ]

    concurrency_note = ['', '← concurrent →', '← concurrent →', '']

    for ci, (title, detail, bg, fc) in enumerate(phases):
        cell = tbl.cell(0, ci)
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title + '\n')
        r.bold = True; r.font.size = Pt(10); r.font.name = 'Calibri'
        r.font.color.rgb = RGBColor(*bytes.fromhex(fc))
        r2 = p.add_run(detail)
        r2.font.size = Pt(8.5); r2.font.name = 'Calibri'
        set_cell_shading(cell, bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Concurrency note row
    row2 = tbl.add_row()
    for ci in range(4):
        cell = row2.cells[ci]
        note = '← Phase 2 & 3 run at the same time →' if ci == 1 else ''
        color = 'BF360C' if note else '888888'
        cell_write(cell, note, size=8, color_hex=color, bg_hex='FFFDE7', italic=bool(note))
        set_cell_border(cell, color='EEEEEE')
    row2.cells[1].merge(row2.cells[2])

    doc.add_paragraph().paragraph_format.space_after = Pt(10)


# ═══════════════════════════════════════════════════════════════
#  DOCUMENT HELPERS
# ═══════════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    set_cell_shading(cell, hex_color)

def set_table_borders(table, color='DDDDDD'):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, color=color)

def shade_para(p, fill='F3F4F6'):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)

def add_inline_para(doc, text, before=3, after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'; r.font.size = Pt(9.5)
        else:
            p.add_run(re.sub(r'&amp;','&', re.sub(r'&lt;','<', re.sub(r'&gt;','>',part))))
    return p


# ═══════════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════════

doc = Document()

for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

for (name, size, color, sb, sa) in [
    ('Heading 1', 18, '1A3A5C', 14, 4),
    ('Heading 2', 13, '0D47A1', 10, 3),
    ('Heading 3', 11, '1B5E20',  8, 2),
]:
    s = doc.styles[name]
    s.font.name = 'Calibri'; s.font.size = Pt(size); s.font.bold = True
    s.font.color.rgb = RGBColor(*bytes.fromhex(color))
    s.paragraph_format.space_before = Pt(sb)
    s.paragraph_format.space_after  = Pt(sa)

CHAPTER_COLORS = {
    'CHAPTER 1': ('1565C0', 'Phase 1 — Scan & Extract'),
    'CHAPTER 2': ('006064', 'Phase 2 — Index & Embed'),
    'CHAPTER 3': ('1B5E20', 'Phase 3 — Fine-Tune'),
    'CHAPTER 4': ('BF360C', 'Phase 4 — Runtime Query'),
    'APPENDIX':  ('4A148C', 'Appendix — Code Listings'),
}

PHASE_TABLE_FN = {
    'CHAPTER 1': make_phase1_table,
    'CHAPTER 2': make_phase2_table,
    'CHAPTER 3': make_phase3_table,
    'CHAPTER 4': make_phase4_table,
}

# ── COVER PAGE ────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('CodeSage for SAP')
r.bold = True; r.font.size = Pt(34)
r.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Know Your Codebase. Build Faster. Stay Compliant.')
r.italic = True; r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x00, 0x97, 0xA7)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Custom AI Coding Assistant for ABAP Developers')
r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

doc.add_paragraph().paragraph_format.space_after = Pt(16)
make_overview_table(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Document Version: 2.0   |   March 2026   |   Classification: Confidential')
r.font.size = Pt(9); r.italic = True
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

doc.add_page_break()

# ── PARSE MARKDOWN ────────────────────────────────────────────

with open('proposal.md', encoding='utf-8') as f:
    lines = f.readlines()

in_code    = False
in_table   = False
code_lines = []
table_rows = []
current_chapter = None

def flush_table():
    global table_rows, in_table
    if not table_rows:
        table_rows = []; in_table = False; return
    data = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
    if not data:
        table_rows = []; in_table = False; return
    parsed = []
    for row in data:
        cols = [c.strip() for c in row.strip().strip('|').split('|')]
        parsed.append(cols)
    ncols = max(len(r) for r in parsed)
    tbl = doc.add_table(rows=len(parsed), cols=ncols)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(parsed):
        for ci in range(ncols):
            ct = row[ci] if ci < len(row) else ''
            cell = tbl.cell(ri, ci)
            cell.text = ''
            p = cell.paragraphs[0]
            cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1',
                     re.sub(r'`(.+?)`', r'\1',
                     re.sub(r'&amp;','&', ct)))
            r = p.add_run(cleaned)
            r.font.size = Pt(9.5); r.font.name = 'Calibri'
            if ri == 0:
                r.bold = True; set_cell_shading(cell, 'D6E4F0')
            elif ri % 2 == 0:
                set_cell_shading(cell, 'F8FAFB')
    set_table_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    table_rows = []; in_table = False

def flush_code():
    global code_lines, in_code
    if not code_lines:
        code_lines = []; in_code = False; return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Cm(0.3)
    p.paragraph_format.right_indent = Cm(0.3)
    shade_para(p, 'F0F2F4')
    r = p.add_run('\n'.join(code_lines))
    r.font.name = 'Courier New'; r.font.size = Pt(8)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    code_lines = []; in_code = False

i = 0
while i < len(lines):
    line = lines[i].rstrip('\n')
    stripped = line.strip()

    if i < 6: i += 1; continue  # skip cover lines

    # Code fence
    if stripped.startswith('```'):
        if in_code: flush_code()
        else: in_code = True
        i += 1; continue
    if in_code:
        code_lines.append(line); i += 1; continue

    # Table
    if stripped.startswith('|'):
        if not in_table: in_table = True
        table_rows.append(stripped); i += 1; continue
    elif in_table:
        flush_table()

    if not stripped or re.match(r'^---+$', stripped):
        i += 1; continue

    # Chapter banner
    m = re.match(r'^# (CHAPTER \d+|APPENDIX)$', stripped)
    if m:
        current_chapter = m.group(1)
        color, label = CHAPTER_COLORS.get(current_chapter, ('1A3A5C', current_chapter))
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after  = Pt(2)
        shade_para(p, color)
        r = p.add_run(f'  {label}  ')
        r.bold = True; r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1; continue

    # Phase title  (# Phase N — ...)
    if re.match(r'^# Phase \d', stripped):
        text = re.sub(r'&amp;','&', re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[2:]))
        doc.add_heading(text, level=1)
        # Insert editable phase diagram table
        if current_chapter in PHASE_TABLE_FN:
            PHASE_TABLE_FN[current_chapter](doc)
            del PHASE_TABLE_FN[current_chapter]
        i += 1; continue

    # H1
    if stripped.startswith('# ') and not stripped.startswith('## '):
        text = re.sub(r'&amp;','&', re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[2:]))
        doc.add_heading(text, level=1); i += 1; continue

    # H2
    if stripped.startswith('## '):
        text = re.sub(r'&amp;','&', re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[3:]))
        doc.add_heading(text, level=2)
        if 'Problem' in text and 'Numbers' in text:
            cats = ['Searching (BAPI)', 'Understanding code',
                    'Writing new code', 'Testing & debug',
                    'Documentation', 'New features freed']
            add_native_bar_chart(doc, cats,
                                 [2.5, 1.5, 2.0, 1.5, 0.5, 0.0],
                                 [0.5, 0.5, 3.0, 1.5, 0.5, 2.0],
                                 'Without CodeSage', 'With CodeSage',
                                 'Developer Day: Before vs After CodeSage')
        i += 1; continue

    # H3
    if stripped.startswith('### '):
        text = re.sub(r'&amp;','&', re.sub(r'\*\*(.+?)\*\*', r'\1', stripped[4:]))
        doc.add_heading(text, level=3); i += 1; continue

    # Step headings
    sm = re.match(r'^\*\*(Step \d+\.\d+[^*]*)\*\*$', stripped)
    if sm:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(sm.group(1))
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        i += 1; continue

    # Blockquote > ...
    if stripped.startswith('> '):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after  = Pt(6)
        p.paragraph_format.left_indent  = Cm(0.8)
        shade_para(p, 'E3F2FD')
        text = stripped[2:]
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        r = p.add_run(text)
        r.bold = True; r.font.size = Pt(11)
        r.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        i += 1; continue

    # Bullet
    if stripped.startswith('- '):
        text = stripped[2:]
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(2)
        p.paragraph_format.left_indent  = Cm(0.5)
        for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
            if part.startswith('**') and part.endswith('**'):
                r = p.add_run(part[2:-2]); r.bold = True
            elif part.startswith('`') and part.endswith('`'):
                r = p.add_run(part[1:-1])
                r.font.name = 'Courier New'; r.font.size = Pt(9.5)
            else:
                p.add_run(re.sub(r'&amp;','&', re.sub(r'&lt;','<', re.sub(r'&gt;','>',part))))
        i += 1; continue

    add_inline_para(doc, stripped)
    i += 1

if in_table: flush_table()
if in_code:  flush_code()

# ── SAVE ─────────────────────────────────────────────────────

out = r'c:\Users\ratna\apps\ABAP 2\CodeSage_Proposal.docx'
doc.save(out)
print(f'Saved: {out}')
