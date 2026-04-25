import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement

doc = Document('C:/Users/ratna/OneDrive/Desktop/RAG_SAP_Developer_Guide_Template.docx')

def find_idx(fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def set_text(idx, new_text):
    if idx < 0:
        return
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.add_run(new_text)

# ── Subtitle ─────────────────────────────────────────────────────────────────
set_text(find_idx('A conceptual and implementation guide for SAP Basis'),
    'A practical implementation guide for ABAP developers — fine-tuned local LLM + RAG over your SAP codebase.')

# ── 1.1 What is RAG ──────────────────────────────────────────────────────────
set_text(find_idx('Retrieval-Augmented Generation (RAG) is an AI architecture'),
    'Retrieval-Augmented Generation (RAG) is an AI architecture pattern that combines a base language model '
    'with a dynamic retrieval step over your own data. Instead of training a model from scratch, RAG retrieves '
    'relevant context at query time from a local vector store and injects it into the prompt — giving grounded, '
    'accurate answers without exposing your data to external services.')

# ── 1.2 Why RAG for SAP ───────────────────────────────────────────────────────
set_text(find_idx('SAP environments hold vast amounts of mission-critical'),
    'SAP environments contain millions of lines of custom ABAP code, BAPIs, function modules, and business '
    'logic accumulated over years. Developers waste time searching for the right BAPI, understanding legacy code, '
    'or rewriting logic that already exists elsewhere in the system. An ABAP Coding Assistant powered by '
    'RAG + fine-tuning solves this by making the entire codebase instantly searchable and queryable in natural language.')

bullets_12 = [
    'Search and explain existing ABAP programs, function modules, and classes.',
    'Suggest the right BAPI for any business scenario (sales order, goods receipt, invoice posting).',
    "Generate new ABAP code following your team's existing coding patterns.",
    'Answer questions grounded in YOUR codebase — not generic internet examples.',
]
start = find_idx('Query ERP data in natural language')
for j, b in enumerate(bullets_12):
    set_text(start + j, b)

# ── 1.3 Scope ─────────────────────────────────────────────────────────────────
set_text(find_idx('This document covers:'),
    'This guide covers the full implementation of a local ABAP Coding Assistant using fine-tuned SLM + RAG:')

scope = [
    'Extracting ABAP source code from SAP via RFC or SE80 export — filtered by selected packages',
    'Fine-tuning CodeLlama-7B on your custom ABAP codebase using QLoRA',
    'Building a local RAG pipeline with ChromaDB and nomic-embed-text embeddings',
    'Data protection, local operation, and audit control using local Git',
]
start = find_idx('RAG concepts and architecture')
for j, b in enumerate(scope):
    set_text(start + j, b)

# ── 2.1 Pipeline ─────────────────────────────────────────────────────────────
set_text(find_idx('A RAG pipeline has five logical stages:'),
    'The ABAP assistant combines three knowledge layers at query time:')

pipeline = [
    'Layer 1 — Base Model: CodeLlama-7B pre-trained on public SAP/ABAP code. Knows standard BAPIs, syntax, and patterns.',
    'Layer 2 — Indexed SAP Docs (RAG): Your SAP Help pages embedded in ChromaDB. Retrieves official BAPI parameters and usage.',
    "Layer 3 — Your Z Code (RAG): All custom programs, FMs, and classes from selected packages. Answers 'how does OUR system do this?'",
    "Fine-tuned Adapter: A QLoRA adapter on top of CodeLlama trained on your team's Q&A pairs. Writes code in your style.",
    'Final Answer: Combines all three — standard knowledge + SAP docs + your codebase example.',
]
start = find_idx('Ingest — Load source documents')
for j, b in enumerate(pipeline):
    set_text(start + j, b)

# ── 3.1 Component Mapping ────────────────────────────────────────────────────
set_text(find_idx('Map each RAG stage to SAP-native or SAP-compatible tooling:'),
    'Each component of the ABAP assistant runs locally with no cloud dependency:')

# ── 3.2 Architecture Variants ────────────────────────────────────────────────
set_text(find_idx('Choose the right pattern based on your data and latency requirements:'),
    'Two extraction options depending on SAP access level:')

arch = [
    'Option A — RFC/API Extraction (Recommended)',
    'Fully automated. pyrfc connects to SAP and pulls source by package. Supports scheduled refresh.',
    'Pattern: pyrfc connection to SAP → RS_PROGRAM_INDEX (package list) → READ_REPORT / RPY_CLASS_READ → local .abap files',
    'Option B — Manual SE80 Export (Quick Start, No Setup)',
    'Open SE80, select package, right-click Export. All .abap files saved to a local folder instantly.',
    'Pattern: SE80 export → local folder → indexer runs automatically → ChromaDB updated → assistant ready',
]
start = find_idx('Simple RAG')
for j, b in enumerate(arch):
    set_text(start + j, b)

# ── 4.1 Prerequisites ────────────────────────────────────────────────────────
prereqs = [
    'Python 3.10+ with pyrfc, transformers, peft, trl, chromadb, llama-index installed',
    'Ollama installed locally with CodeLlama-7B pulled (ollama pull codellama:7b)',
    'SAP RFC-enabled read-only user (ZABAP_SCANNER) OR SE80 access for manual export',
    'GPU recommended: RTX 3060 (8GB VRAM) minimum; cloud GPU (RunPod A100) for faster fine-tuning',
    'Local Git installed — for audit trail, version control, and rollback',
]
start = find_idx('SAP BTP subaccount with Generative AI Hub')
for j, b in enumerate(prereqs):
    set_text(start + j, b)

# ── 4.2 Install ───────────────────────────────────────────────────────────────
set_text(find_idx('Install the required Python packages:'), 'Install required Python packages for the full pipeline:')
set_text(find_idx('# Python dependencies'), '# Install all dependencies')
set_text(find_idx('pip install generative-ai-hub-sdk'),
    'pip install pyrfc chromadb llama-index transformers peft trl torch accelerate bitsandbytes')
set_text(find_idx('# Verify SAP AI Core'), '# Pull the base model locally via Ollama')
set_text(find_idx('--auth-url <AUTH_URL> --resource-group default'), 'ollama pull codellama:7b')

# ── 4.3 heading + content ────────────────────────────────────────────────────
set_text(find_idx('4.3  Step 2 — Connect to SAP HANA Vector Engine'), '4.3  Step 2 — Extract ABAP Source from SAP')
set_text(find_idx('Python — hana_connect.py'), 'Python — sap_connector.py')
set_text(find_idx('# hana_connect.py'), '# sap_connector.py')
set_text(find_idx('from hdbcli import dbapi'), 'import pyrfc, yaml, os')
set_text(find_idx('conn = dbapi.connect('), 'conn = pyrfc.Connection(')
set_text(find_idx("address='<HANA_HOST>.hanacloud.ondemand.com',"),
    "    ashost=os.getenv('SAP_HOST'), sysnr='00',")
set_text(find_idx('port=443,'),
    "    client='100', user=os.getenv('SAP_USER'), passwd=os.getenv('SAP_PASS'))")
set_text(find_idx("user='<HANA_USER>',"), '')
set_text(find_idx("password='<HANA_PASSWORD>',"),
    'config = yaml.safe_load(open("config/scan_config.yaml"))')
set_text(find_idx('encrypt=True'), 'for pkg in config["packages"]["include"]:')
set_text(find_idx('# Create vector table (run once)'),
    '    objects = conn.call("RS_PROGRAM_INDEX", DEVCLASS=pkg)')
for frag in ['cursor = conn.cursor()', "cursor.execute('''", 'ID          BIGINT',
             'SOURCE      NVARCHAR', 'CHUNK_TEXT  NCLOB', 'EMBEDDING   REAL_VECTOR',
             'METADATA    NCLOB', "''')", "conn.commit()","print('Vector table created"]:
    set_text(find_idx(frag), '')

# ── 4.4 heading ──────────────────────────────────────────────────────────────
set_text(find_idx('4.4  Step 3 — Ingest and Embed Documents'), '4.4  Step 3 — Index Code into ChromaDB')
set_text(find_idx('The ingest pipeline chunks your source documents and writes embeddings to HANA:'),
    'The indexer reads all .abap files, chunks by logical unit (class/method/FM), embeds with nomic-embed-text, '
    'and stores vectors in ChromaDB. Run once after extraction; re-run after any package refresh.')
set_text(find_idx('Python — ingest.py'), 'Python — index_rag.py')

# ── 4.5 heading ──────────────────────────────────────────────────────────────
set_text(find_idx('4.5  Step 4 — Retrieval Function'), '4.5  Step 4 — Query Engine (Retrieval + Answer)')
set_text(find_idx('Python — retrieve.py'), 'Python — query_engine.py')
set_text(find_idx('# retrieve.py'), '# query_engine.py — retrieve relevant code chunks from ChromaDB')

# ── 4.6 heading ──────────────────────────────────────────────────────────────
set_text(find_idx('4.6  Step 5 — Generate Grounded Response'), '4.6  Step 5 — Fine-tune CodeLlama with QLoRA')
set_text(find_idx('Python — generate.py'), 'Python — finetune.py')
set_text(find_idx('# generate.py'), '# finetune.py — QLoRA fine-tuning on your ABAP Q&A pairs')

# ── 4.7 heading ──────────────────────────────────────────────────────────────
set_text(find_idx('4.7  ABAP Integration via RFC / OData'), '4.7  Developer UI (Streamlit Chat Interface)')
set_text(find_idx('For ABAP developers, expose the RAG service via an OData endpoint and call from ABAP:'),
    'The Streamlit UI runs on localhost. Developers interact via a browser-based chat interface. '
    'Includes a Package Picker screen for admins to configure which SAP packages to scan, '
    'a Preview screen showing estimated object count before any extraction, and the main chat for queries.')

# ── 5.1 Security ─────────────────────────────────────────────────────────────
set_text(find_idx('Authentication & Authorization'), 'Local Operation & Network Isolation')
set_text(find_idx('Secure all RAG API endpoints with SAP XSUAA OAuth 2.0 tokens'),
    'All components (model, vector store, UI) run on an internal server — zero internet calls at query time.')
set_text(find_idx('Apply Role-Based Access Control (RBAC)'),
    'Firewall rule: block all outbound traffic from the assistant server.')
set_text(find_idx('Never embed SAP credentials in code; use SAP Credential Store'),
    'SAP credentials stored in .env file, excluded from Git via .gitignore, rotated quarterly.')

set_text(find_idx('Data Residency & Privacy'), 'SAP Access Control — Minimum Permissions')
set_text(find_idx("Confirm your LLM provider's data processing agreement"),
    'Create dedicated read-only RFC user ZABAP_SCANNER with S_DEVELOP ACTVT=03 (Display only, no write/transport).')
set_text(find_idx('Strip PII (employee IDs, vendor data) from chunks'),
    'Scanner never touches SAP standard code or objects outside selected packages.')
set_text(find_idx('Use SAP AI Core with on-premise model deployment'),
    'Local objects ($TMP) excluded by default. Full scan log committed to Git after every extraction run.')

set_text(find_idx('Network Security'), 'Data Protection')
set_text(find_idx('Route all LLM calls through SAP API Management'),
    'All .abap source files, ChromaDB vector store, and fine-tuned model adapter remain on-premise at all times.')
set_text(find_idx('Use SAP Private Link for HANA Cloud access'),
    'No Git remote configured — local Git only. No code or data ever pushed to GitHub or any external service.')

# ── 5.3 heading + content ─────────────────────────────────────────────────────
set_text(find_idx('5.3  Evaluation & Quality Metrics'), '5.3  Audit Trail & Version Control (Local Git)')
set_text(find_idx('Track these metrics to ensure your RAG pipeline meets production standards:'),
    'Local Git tracks every change to the system — extractions, retrains, config changes — for compliance and rollback:')

audit = [
    'git commit after every SAP scan — records exactly which packages and objects were extracted, with timestamp.',
    'git commit after every retrain — records which training pairs and model version were used.',
    'git tag for stable releases (v1.0, v1.1) — any version restorable in minutes with git checkout.',
    'git blame on scan_config.yaml — shows who changed package selection and when.',
    'git diff between scans — shows exactly what ABAP code changed between extraction runs.',
]
start = find_idx('Retrieval Precision@k')
for j, b in enumerate(audit):
    set_text(start + j, b)

# ── 6.1 Glossary ──────────────────────────────────────────────────────────────
set_text(find_idx('[Add project-specific or organisation-specific terms here.'),
    'RAG (Retrieval-Augmented Generation): Combines a base model with dynamic retrieval from a local vector store. '
    'QLoRA (Quantized Low-Rank Adaptation): Efficient fine-tuning that runs on a single consumer GPU using 4-bit quantization. '
    'ChromaDB: Local open-source vector database — stores embeddings as a file on disk, no server needed. '
    'RFC (Remote Function Call): SAP protocol used to extract source code programmatically via pyrfc. '
    'Z-objects: Custom SAP developments prefixed with Z or Y — the primary training and indexing data source. '
    'nomic-embed-text: Free, local embedding model run via Ollama — converts ABAP code chunks to vectors.')

# ── 6.2 References ────────────────────────────────────────────────────────────
refs = [
    'CodeLlama — Meta AI: huggingface.co/codellama',
    'QLoRA Fine-tuning: github.com/artidoro/qlora',
    'ChromaDB Documentation: docs.trychroma.com',
    'LlamaIndex RAG Framework: docs.llamaindex.ai',
    'pyrfc — SAP RFC for Python: github.com/SAP/PyRFC',
]
start = find_idx('SAP Help: SAP HANA Cloud Vector Engine')
for j, b in enumerate(refs):
    set_text(start + j, b)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('C:/Users/ratna/OneDrive/Desktop/RAG_SAP_Developer_Guide_Template.docx')
print('Saved successfully.')
