import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import nsmap
import copy

# ── Colour Palette (Canva-inspired dark navy + teal) ─────────────────────────
NAV  = RGBColor(0x1B, 0x2B, 0x4B)   # Deep navy
TEAL = RGBColor(0x00, 0xB4, 0xD8)   # Bright teal accent
MINT = RGBColor(0x06, 0xD6, 0xA0)   # Mint green
ORG  = RGBColor(0xFF, 0x6B, 0x35)   # Orange accent
WHT  = RGBColor(0xFF, 0xFF, 0xFF)   # White
LGY  = RGBColor(0xF4, 0xF7, 0xFB)   # Light grey bg
MGY  = RGBColor(0xD0, 0xDC, 0xEE)   # Mid grey
DGY  = RGBColor(0x44, 0x55, 0x6B)   # Dark grey text
YLW  = RGBColor(0xFF, 0xD1, 0x66)   # Yellow highlight

doc = Document(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')

# ── XML helpers ───────────────────────────────────────────────────────────────
def hex_color(rgb: RGBColor) -> str:
    return f'{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}'

def set_run_color(run, rgb):
    run.font.color.rgb = rgb

def set_para_shading(para, fill_rgb: RGBColor):
    """Set paragraph background fill colour."""
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color(fill_rgb))
    pPr.append(shd)

def set_para_border_bottom(para, rgb: RGBColor, sz=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(sz))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color(rgb))
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_table_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color(rgb))
    tcPr.append(shd)

def set_cell_text_color(cell, rgb: RGBColor):
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.color.rgb = rgb

def set_cell_bold(cell, bold=True):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold

def set_para_space(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    pSp = OxmlElement('w:spacing')
    pSp.set(qn('w:before'), str(before))
    pSp.set(qn('w:after'),  str(after))
    pPr.append(pSp)

def add_left_bar(para, rgb: RGBColor, sz=24):
    """Add a thick left border bar to a paragraph (blockquote style)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(sz))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), hex_color(rgb))
    pBdr.append(left)
    pPr.append(pBdr)

def indent_para(para, left_twips=360):
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement('w:ind')
    ind.set(qn('w:left'), str(left_twips))
    pPr.append(ind)

# ── Style all paragraphs ──────────────────────────────────────────────────────
for i, para in enumerate(doc.paragraphs):
    style = para.style.name if para.style else ''
    text  = para.text.strip()
    if not text:
        continue

    # ── TITLE PAGE ────────────────────────────────────────────────────────────
    if text in ('ABAP AI CODING ASSISTANT', 'CodeSage for SAP',
                'Know Your Codebase. Build Faster. Stay Secure.'):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_shading(para, NAV)
        for run in para.runs:
            run.font.color.rgb = WHT
            run.font.bold       = True
            if text == 'ABAP AI CODING ASSISTANT':
                run.font.size = Pt(32)
                run.font.name = 'Calibri'
            elif text == 'CodeSage for SAP':
                run.font.size = Pt(22)
                run.font.color.rgb = TEAL
                run.font.name = 'Calibri'
            else:
                run.font.size = Pt(13)
                run.font.color.rgb = YLW
                run.font.name = 'Calibri'
        set_para_space(para, before=120, after=80)
        continue

    if 'A practical implementation guide for ABAP developers' in text:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_para_shading(para, NAV)
        for run in para.runs:
            run.font.color.rgb = MGY
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            run.font.italic = True
        set_para_space(para, before=60, after=200)
        continue

    # ── HEADING 1 ─────────────────────────────────────────────────────────────
    if 'Heading 1' in style:
        set_para_shading(para, NAV)
        set_para_space(para, before=240, after=120)
        for run in para.runs:
            run.font.color.rgb = WHT
            run.font.bold       = True
            run.font.size       = Pt(16)
            run.font.name       = 'Calibri'
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        continue

    # ── HEADING 2 ─────────────────────────────────────────────────────────────
    if 'Heading 2' in style:
        set_para_border_bottom(para, TEAL, sz=16)
        set_para_space(para, before=180, after=80)
        for run in para.runs:
            run.font.color.rgb = NAV
            run.font.bold       = True
            run.font.size       = Pt(13)
            run.font.name       = 'Calibri'
        continue

    # ── HEADING 3 ─────────────────────────────────────────────────────────────
    if 'Heading 3' in style or 'Heading 4' in style:
        set_para_space(para, before=120, after=60)
        for run in para.runs:
            run.font.color.rgb = TEAL
            run.font.bold       = True
            run.font.size       = Pt(11)
            run.font.name       = 'Calibri'
        continue

    # ── LIST ITEMS ────────────────────────────────────────────────────────────
    if 'List' in style:
        add_left_bar(para, TEAL, sz=12)
        indent_para(para, left_twips=480)
        for run in para.runs:
            run.font.color.rgb = DGY
            run.font.size       = Pt(10)
            run.font.name       = 'Calibri'
        set_para_space(para, before=40, after=40)
        continue

    # ── CODE BLOCKS (monospace / terminal style) ──────────────────────────────
    if style in ('', 'None') and any(c in text for c in ['█', '═', '│', '┌', '└', '►', '→', 'pip ', 'import ', 'def ', 'conn =', 'SELECT', 'CALL FUNCTION']):
        set_para_shading(para, RGBColor(0x1E, 0x1E, 0x2E))
        indent_para(para, left_twips=360)
        for run in para.runs:
            run.font.color.rgb = RGBColor(0xA6, 0xE2, 0x2E)
            run.font.size       = Pt(9)
            run.font.name       = 'Courier New'
        set_para_space(para, before=60, after=60)
        continue

    # ── BODY TEXT ─────────────────────────────────────────────────────────────
    for run in para.runs:
        if run.font.color.rgb is None or run.font.color.rgb == RGBColor(0,0,0):
            run.font.color.rgb = DGY
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
    set_para_space(para, before=60, after=60)

# ── Style all Tables ──────────────────────────────────────────────────────────
for tbl_idx, table in enumerate(doc.tables):
    # Alternate accent per table
    accent = [TEAL, MINT, ORG][tbl_idx % 3]

    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Header row
            if row_idx == 0:
                set_table_cell_bg(cell, NAV)
                set_cell_text_color(cell, WHT)
                set_cell_bold(cell, True)
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
            # First column
            elif col_idx == 0:
                set_table_cell_bg(cell, LGY)
                set_cell_bold(cell, True)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = NAV
                        run.font.size       = Pt(9)
                        run.font.name       = 'Calibri'
            # Alternating rows
            elif row_idx % 2 == 0:
                set_table_cell_bg(cell, LGY)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = DGY
            else:
                set_table_cell_bg(cell, WHT)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
                        run.font.color.rgb = DGY

    # Table borders
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), hex_color(MGY))
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Table width = full page
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '9360')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

# ── Page margins ──────────────────────────────────────────────────────────────
from docx.oxml import OxmlElement as OE
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(r'C:\Users\ratna\OneDrive\Desktop\RAG_SAP_Developer_Guide_Template.docx')
print('Beautified and saved.')
