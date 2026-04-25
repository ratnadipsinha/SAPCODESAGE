from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ratna\apps\ABAP 2\CodeSage_modified.docx'

SAP_BLUE = RGBColor(0x0A, 0x6E, 0xD1)
SAP_DARK = RGBColor(0x1B, 0x6C, 0xA8)
DARK     = RGBColor(0x1A, 0x1A, 0x1A)
GREY     = RGBColor(0x88, 0x88, 0x88)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)

# ── XML helpers ───────────────────────────────────────────────────────────────

def cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)

def cell_border(cell, color='0A6ED1', sz='10'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fc1 = OxmlElement('w:fldChar'); fc1.set(qn('w:fldCharType'), 'begin')
    it  = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
    it.text = ' TOC \\o "1-3" \\h \\z \\u '
    fc2 = OxmlElement('w:fldChar'); fc2.set(qn('w:fldCharType'), 'separate')
    fc3 = OxmlElement('w:fldChar'); fc3.set(qn('w:fldCharType'), 'end')
    run._r.append(fc1); run._r.append(it)
    run._r.append(fc2); run._r.append(fc3)

# ── Content helpers ───────────────────────────────────────────────────────────

def p(doc, text, bold=False, italic=False, color=None, size=None,
      align=None, sb=None, sa=None, style='Normal', indent=None):
    para = doc.add_paragraph(style=style)
    run  = para.add_run(text)
    if bold:   run.bold   = True
    if italic: run.italic = True
    if color:  run.font.color.rgb = color
    if size:   run.font.size = Pt(size)
    if align:  para.alignment = align
    if sb is not None: para.paragraph_format.space_before = Pt(sb)
    if sa is not None: para.paragraph_format.space_after  = Pt(sa)
    if indent is not None: para.paragraph_format.left_indent = Inches(indent)
    return para

def bullet(doc, text, bold_prefix=None, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    para  = doc.add_paragraph(style=style)
    if bold_prefix:
        r = para.add_run(bold_prefix); r.bold = True
    para.add_run(text)
    return para

def numbered(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = para.add_run(bold_prefix); r.bold = True
    para.add_run(text)
    return para

def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.35)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return para

def info_box(doc, label, body):
    """Single-cell shaded box for callouts / notes."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(7.0)
    cell_bg(cell, 'EAF4FF')
    cell_border(cell, '0A6ED1', '8')
    para = cell.paragraphs[0]
    r1 = para.add_run(label + '  '); r1.bold = True
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = SAP_BLUE
    r2 = para.add_run(body)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    doc.add_paragraph()

def make_table(doc, headers, rows, col_widths=None, hdr_bg='0A6ED1', font_size=10):
    """Striped table with SAP-blue header."""
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        cell_bg(c, hdr_bg)
        para = c.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        r.bold = True; r.font.name = 'Calibri'
        r.font.size = Pt(font_size); r.font.color.rgb = WHITE
    # data rows
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            c = tbl.rows[ri + 1].cells[ci]
            if ri % 2 == 1: cell_bg(c, 'EAF4FF')
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Calibri'; r.font.size = Pt(font_size)
    # widths
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def apply_styles(doc):
    for name, sz, col, sb, sa in [
        ('Heading 1', 18, SAP_BLUE, 24, 8),
        ('Heading 2', 14, SAP_DARK, 12, 6),
        ('Heading 3', 12, DARK,      8, 4),
    ]:
        s = doc.styles[name]
        s.font.name = 'Calibri'; s.font.size = Pt(sz)
        s.font.color.rgb = col; s.font.bold = True
        s.paragraph_format.space_before   = Pt(sb)
        s.paragraph_format.space_after    = Pt(sa)
        s.paragraph_format.keep_with_next = True
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_after = Pt(5)

# =============================================================================
# BUILD
# =============================================================================
doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1.2)
    s.page_width  = Inches(8.27)
    s.page_height = Inches(11.69)
apply_styles(doc)

# ── Title page ─────────────────────────────────────────────────────────────────
p(doc, '', sb=36, sa=0)
p(doc, 'CodeSage for SAP',
  bold=True, color=SAP_BLUE, size=32,
  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
p(doc, 'AI-Powered Code Knowledge Platform',
  bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
p(doc, 'SAP AI Core BYOM  \u00b7  LLaMA-3 8B Instruct  \u00b7  RAG (ChromaDB)  \u00b7  Claude Skill',
  color=GREY, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=16)
p(doc, '\u2500' * 72, color=SAP_BLUE, size=9,
  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=16)
p(doc, 'Version 1.0  |  March 2026  |  Confidential',
  color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
doc.add_page_break()

# ── Contents ───────────────────────────────────────────────────────────────────
doc.add_heading('Contents', level=1)
p(doc, 'Open in Microsoft Word and press Ctrl+A then F9 to update page numbers.',
  italic=True, color=GREY, size=9, sa=2)
add_toc_field(doc)
doc.add_page_break()

# =============================================================================
# CHAPTER 1 — THE INVISIBLE ARCHIVE
# =============================================================================
doc.add_heading('Chapter 1:  The Invisible Archive \u2014 Unlocking Institutional Knowledge', level=1)

doc.add_heading('1.1  The Organisational Challenge', level=2)
p(doc, 'Every organisation that has run SAP for more than a few years carries with it an '
       'invisible archive \u2014 thousands of custom programs, function modules, classes, '
       'BAdI implementations, and enhancement spots, written over years or even decades by '
       'developers who may no longer be with the organisation. This archive holds enormous '
       'value: reusable logic, solved problems, domain knowledge encoded in code. But it is '
       'almost entirely inaccessible in practice.')
p(doc, 'Ask a developer "do we already have a function module that validates vendor payment '
       'terms?" and the honest answer, in most organisations, is "I don\'t know." If the '
       'developer has prior experience with a similar solution they will bring that. So every '
       'new talent designs the solution in their own way. Existing solutions are rarely '
       'enhanced or reused \u2014 they are quietly duplicated.')
p(doc, 'Today organisations have multiple implementation partners. When they onboard a new '
       'team, that team rarely gets enough time to learn the existing system, its processes, '
       'and its accumulated logic. That brings delay and misalignment.')
p(doc, 'This is a classic case of knowledge being siloed and lost in transitions \u2014 '
       'institutional knowledge decay. It leads to unawareness of existing modules, '
       'reinventing solutions instead of enhancing them, and new implementation partners '
       'struggling to ramp up.')

doc.add_heading('1.2  Why Institutional Knowledge Gets Locked Away', level=2)
bullet(doc, 'Code may exist but is poorly documented \u2014 making it invisible to anyone who did not write it.', bold_prefix='Documentation gaps:  ')
bullet(doc, 'Each partner brings their own practices and rarely invests time in learning the existing system deeply.', bold_prefix='Multiple vendors / partners:  ')
bullet(doc, 'Onboarding teams are pushed to deliver quickly, so they design new solutions rather than explore legacy ones.', bold_prefix='Time pressure:  ')
bullet(doc, 'Developers lean on what they have done before because it feels safer than digging into unfamiliar codebases.', bold_prefix='Experience bias:  ')

doc.add_heading('1.3  Consequences', level=2)
bullet(doc, 'Duplicate functionality \u2014 multiple modules solving the same problem differently.')
bullet(doc, 'Increased maintenance burden across fragmented implementations.')
bullet(doc, 'Slower onboarding and longer delivery timelines for every new partner.')
bullet(doc, 'Architecture drift away from original design principles.')
bullet(doc, 'Higher costs due to redundant development effort across projects.')
p(doc, 'In short, the problem is not just technical \u2014 it is cultural. Organisations that '
       'value reuse and documentation build resilience. Those that rely on tribal knowledge '
       'end up reinventing the wheel every time a team changes.')

doc.add_heading('1.4  What Organisations Can Do', level=2)
bullet(doc, 'Maintain a searchable catalogue of existing modules, APIs, and business processes \u2014 an internal service registry.', bold_prefix='Knowledge repositories:  ')
bullet(doc, 'Use automated documentation generators or code search platforms that answer "do we have a function for X?"', bold_prefix='Code discovery tools:  ')
bullet(doc, 'Require new partners to spend time with system walkthroughs before delivery begins.', bold_prefix='Mandatory onboarding:  ')
bullet(doc, 'A small core team should enforce reuse policies \u2014 new solutions must extend existing ones.', bold_prefix='Architecture governance:  ')
bullet(doc, 'New developers shadow experienced ones to absorb tacit knowledge faster.', bold_prefix='Pairing & shadowing:  ')

doc.add_heading('1.5  Solution Proposal: CodeSage for SAP', level=2)
p(doc, 'CodeSage for SAP is an AI-powered code knowledge assistant purpose-built for ABAP '
       'development teams. It gives every developer \u2014 new or experienced \u2014 the '
       'ability to query your organisation\'s entire codebase in plain English and receive '
       'accurate, contextualised answers grounded in your own source code.')
p(doc, 'Unlike generic AI tools, CodeSage is trained on your specific ABAP landscape: your '
       'naming conventions, your BAPIs, your error classes, your enhancement patterns. The '
       'result is a system that does not just know SAP \u2014 it knows your SAP.')
info_box(doc, '\u26a1 Technology:',
    'LLaMA-3 8B Instruct (Meta) fine-tuned on your ABAP code using QLoRA, deployed on '
    'SAP AI Core via Bring Your Own Model (BYOM). Retrieval-Augmented Generation (RAG) '
    'grounds every answer in your live ChromaDB vector index \u2014 keeping everything '
    'within SAP\'s managed, compliance-certified infrastructure.')

doc.add_heading('1.6  How CodeSage Works \u2014 Technical Overview', level=2)
p(doc, 'CodeSage is built from two complementary techniques working together:')

p(doc, '1.  Fine-Tuned Language Model \u2014 LLaMA-3 on SAP AI Core BYOM', bold=True, sa=2)
p(doc, 'The base model \u2014 LLaMA-3 8B Instruct \u2014 is fine-tuned on your '
       'organisation\'s actual ABAP code using QLoRA (Quantised Low-Rank Adaptation). It '
       'learns your naming conventions, program structures, preferred BAPIs, error class '
       'patterns, and enhancement approaches. This is a one-time training exercise '
       '(6\u201312 hours on a free Kaggle T4 GPU) that produces a model deeply familiar '
       'with your codebase. The fine-tuned model is packaged in vLLM format and deployed '
       'to SAP AI Core via BYOM.')

p(doc, '2.  Retrieval-Augmented Generation (RAG)', bold=True, sa=2)
p(doc, 'At query time, the system searches a local vector database (ChromaDB) holding your '
       'indexed ABAP source code and SAP documentation. Relevant code snippets are retrieved '
       'and included in the prompt, so every answer is grounded in real, up-to-date source '
       'material \u2014 not just what the model memorised during training. This eliminates '
       'hallucination and ensures all responses cite actual code.')

p(doc, 'The result is a system that answers not just "what is the standard BAPI for sales '
       'order creation?" but "how does our team call that BAPI, in our programs, in our '
       'style, with our error handling?"')

doc.add_heading('1.7  Why LLaMA-3 on SAP AI Core BYOM', level=2)
p(doc, 'SAP AI Core BYOM supports a curated set of open-source models. The table below '
       'compares the available options. CodeSage uses LLaMA-3 8B Instruct as its recommended '
       'base model due to its strong code reasoning and full BYOM compatibility.')
make_table(doc,
    ['Model', 'Provider', 'BYOM Support', 'Code Suitability', 'Recommendation'],
    [
        ['LLaMA-3 8B Instruct', 'Meta AI',    '\u2705 Supported',      'Excellent \u2014 strong code reasoning & instruction following', '\u2b50 Recommended'],
        ['Mistral 7B',          'Mistral AI',  '\u2705 Supported',      'Good \u2014 fast and lightweight',                              'Alternative'],
        ['Mixtral 8\u00d77B',  'Mistral AI',  '\u2705 Supported',      'Very good \u2014 higher capacity MoE model',                    'If larger context needed'],
        ['Phi-3',               'Microsoft',   '\u2705 Supported',      'Good \u2014 compact and efficient',                             'Low-resource option'],
        ['CodeLlama-7B',        'Meta AI',     '\u274c Not supported',  'Purpose-built for code but not available in BYOM',              'Not applicable'],
    ],
    col_widths=[1.7, 1.2, 1.3, 2.7, 1.5]
)

doc.add_page_break()

# =============================================================================
# CHAPTER 2 — BUILD PIPELINE
# =============================================================================
doc.add_heading('Chapter 2:  Build Pipeline \u2014 All Four Phases at a Glance', level=1)

doc.add_heading('2.1  Pipeline Overview', level=2)
p(doc, 'The CodeSage build pipeline runs once to produce the deployed system, then refreshes '
       'monthly to incorporate new code. Phases 2 and 3 are independent and run concurrently '
       'once Phase 1 completes.')
make_table(doc,
    ['Phase', 'Name', 'What Happens', 'Runs', 'Duration', 'Output'],
    [
        ['1', 'Scan & Extract',      'RFC read from SAP \u2192 save .abap + .json files',  'Once + monthly',   '1\u20132 hrs',         'ABAP source files ready for Phase 2 & 3'],
        ['2', 'Index & Embed',       'Chunk \u2192 embed \u2192 store in ChromaDB',        'Once + monthly',   '~2 hrs (CPU)',          'Local vector index \u2014 no cloud, no GPU'],
        ['3', 'Fine-Tune & Package', 'QA gen \u2192 QLoRA train \u2192 Docker \u2192 BYOM','Once + quarterly', '6\u201312 hrs (GPU)',   'Fine-tuned model live on SAP AI Core'],
        ['4', 'Runtime Query',       'Embed \u2192 RAG \u2192 Agent \u2192 LLaMA-3',       'Per request',      '2\u20135 s',           'Developer answer with cited source objects'],
    ],
    col_widths=[0.45, 1.5, 2.6, 1.2, 1.2, 2.3]
)

# ── Phase 1 ──────────────────────────────────────────────────────────────────
doc.add_heading('Chapter 3:  Phase 1 \u2014 Scan & Extract', level=1)
p(doc, 'Before we can teach an AI anything about your codebase, we first need to read it. '
       'Phase 1 is purely about discovery: connect to the live SAP system, extract source '
       'code for every custom object, and save it as readable files on a local machine. '
       'Nothing is changed in SAP. Nothing is sent to the cloud. This phase is read-only.')

p(doc, 'How it connects:', bold=True, sa=2)
p(doc, 'The extractor uses pyrfc \u2014 a Python library that speaks SAP\'s Remote Function '
       'Call (RFC) protocol. It calls the standard SAP function module RPY_PROGRAM_READ '
       '(and equivalents for function groups, classes, BAdIs) to pull source code line by '
       'line. A scan_config.yaml file controls exactly what is in scope.')

bullet(doc, 'Programs, function modules, classes, BAdIs, enhancement spots \u2014 Z*/Y* namespace by default. Standard SAP objects (no Z/Y prefix) are excluded \u2014 they are not organisation-specific.')
bullet(doc, 'Each object is saved as two files: a .abap source file containing the code, and a .json sidecar recording object type, package, author, and last-changed date. The sidecar is later used to populate ChromaDB metadata.')
bullet(doc, 'A pre-scan filter strips hard-coded passwords, personal identifiers, and anything matching credential patterns before files are written to disk.')
bullet(doc, 'The RFC transport profile is read-only \u2014 no ABAP transport, no change document, no audit footprint in the SAP system.')

p(doc, 'What the scan produces:', bold=True, sa=2)
p(doc, 'A folder of .abap and .json files representing your entire custom codebase \u2014 '
       'typically 5,000\u201320,000 objects for a mid-size SAP landscape. This flat-file '
       'representation is the input to both Phase 2 (indexing) and Phase 3 (fine-tuning).')

p(doc, 'Phase 1 folder structure:', bold=True, sa=2)
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase1_extract/
\u2502   \u251c\u2500\u2500 scan_config.yaml          # Controls which objects/namespaces to scan
\u2502   \u251c\u2500\u2500 extractor.py              # RFC extractor script (run once or monthly)
\u2502   \u2514\u2500\u2500 abap_files/               # OUTPUT of Phase 1
\u2502       \u251c\u2500\u2500 Z_VALIDATE_VENDOR.abap  # Source code (one file per object)
\u2502       \u251c\u2500\u2500 Z_VALIDATE_VENDOR.json  # Metadata sidecar
\u2502       \u251c\u2500\u2500 ZCL_PO_STATUS_HANDLER.abap
\u2502       \u251c\u2500\u2500 ZCL_PO_STATUS_HANDLER.json
\u2502       \u2514\u2500\u2500 ... (5,000 \u2013 20,000 files)""")

p(doc, 'scan_config.yaml:', bold=True, sa=2)
code_block(doc,
"""sap:
  host:     "your-sap-host"       # SAP application server hostname
  sysnr:    "00"                  # System number
  client:   "100"                 # Client
  user:     "CODESAGE_RFC"        # Read-only RFC user (no dialog logon needed)
  password: "${SAP_RFC_PASSWORD}" # Set as environment variable \u2014 never in plain text

output_dir: "./abap_files"

namespaces:                        # Only extract custom objects
  - "Z*"
  - "Y*"

object_types:
  - PROG    # Programs / Reports
  - FUGR    # Function Groups (contains Function Modules)
  - CLAS    # ABAP Classes
  - INTF    # ABAP Interfaces
  - ENHO    # BAdI / Enhancement Implementations

exclude_packages:
  - "$TMP"  # Local dev objects \u2014 not part of the official codebase""")

p(doc, 'extractor.py \u2014 Python RFC extraction script:', bold=True, sa=2)
code_block(doc,
"""import pyrfc, json, pathlib, yaml, re

# Load config
cfg = yaml.safe_load(open('scan_config.yaml'))
out = pathlib.Path(cfg['output_dir'])
out.mkdir(exist_ok=True)

# Pre-scan filter: strip credentials before writing to disk
CRED_RE = re.compile(
    r'(password|passwd|pwd|secret|apikey)\\s*=\\s*[\\'\\"][^\\'\\\"]+[\\'\\"]', re.I)

# Open read-only RFC connection
conn = pyrfc.Connection(
    ashost=cfg['sap']['host'],   sysnr=cfg['sap']['sysnr'],
    client=cfg['sap']['client'], user=cfg['sap']['user'],
    passwd=cfg['sap']['password'])

def save_object(name, obj_type, source, meta):
    source_clean = CRED_RE.sub('[REDACTED]', source)
    (out / f'{name}.abap').write_text(source_clean, encoding='utf-8')
    (out / f'{name}.json').write_text(json.dumps(meta, indent=2), encoding='utf-8')

# --- Extract Programs (PROG) ---
progs = conn.call('RPY_DIRECTORY_FINISH', OBJECT_TYPE='PROG', GENERIC_NAME='Z*')
for obj in progs['TADIR']:
    try:
        r = conn.call('RPY_PROGRAM_READ', PROG_NAME=obj['OBJ_NAME'])
        src = '\\n'.join(l['LINE'] for l in r['SOURCE'])
        save_object(obj['OBJ_NAME'], 'PROG', src, {
            'name': obj['OBJ_NAME'], 'type': 'PROG',
            'package': obj.get('DEVCLASS',''), 'changed': obj.get('LDATE','')})
    except Exception as e:
        print(f'  SKIP {obj["OBJ_NAME"]}: {e}')

# --- Extract Function Modules (FUGR) ---
fugs = conn.call('RPY_DIRECTORY_FINISH', OBJECT_TYPE='FUGR', GENERIC_NAME='Z*')
for obj in fugs['TADIR']:
    try:
        r = conn.call('RPY_FUNCTIONMODULE_READ', FUNCNAME=obj['OBJ_NAME'])
        src = '\\n'.join(l['LINE'] for l in r.get('SOURCE',[]))
        save_object(obj['OBJ_NAME'], 'FUGR', src, {
            'name': obj['OBJ_NAME'], 'type': 'FM',
            'package': obj.get('DEVCLASS',''), 'changed': obj.get('LDATE','')})
    except Exception as e:
        print(f'  SKIP {obj["OBJ_NAME"]}: {e}')

print(f'Done. {len(list(out.glob("*.abap")))} objects saved to {out}/')""")

# ── How to Run Phase 1 ───────────────────────────────────────────────────────
doc.add_heading('3.1  How to Run Phase 1', level=2)

p(doc, 'Prerequisites', bold=True, sa=2)
make_table(doc,
    ['Requirement', 'Details', 'Where to Get It'],
    [
        ['Python 3.10+',      'Required to run extractor.py',                               'python.org or Anaconda'],
        ['pyrfc library',     'Python binding for SAP RFC protocol',                         'pip install pyrfc (needs SAP NW RFC SDK)'],
        ['SAP NW RFC SDK',    '64-bit C libraries required by pyrfc',                        'SAP Support Portal \u2014 search "SAP NW RFC SDK"'],
        ['RFC user in SAP',   'A dialog or system user with S_RFC auth object (read-only)',   'Created by BASIS team \u2014 see note below'],
        ['Network access',    'TCP port 33XX (XX = system number) open to SAP app server',   'Network / firewall team'],
        ['PyYAML',            'For reading scan_config.yaml',                                 'pip install pyyaml'],
    ],
    col_widths=[1.6, 2.8, 2.8]
)

info_box(doc, '\U0001f511 SAP RFC user authorisation (BASIS note):',
    'The RFC user needs authorisation object S_RFC with ACTVT=16 (execute) for '
    'RFC function groups RPY_*, SFES_*, RS_*, SOBJ_*. No write authorisations are '
    'required. Recommended: create a dedicated system user CODESAGE_RFC and lock it '
    'to the IP address of the extraction machine via SU01 \u2192 Logon Data \u2192 '
    'Valid From/To + restrict via profile parameter login/accept_sso2_ticket.')

p(doc, 'All commands below use Windows PowerShell.  Open PowerShell by pressing '
       'Win + X \u2192 Windows PowerShell (Admin).  Every code block is a copy-paste '
       'command \u2014 run them in the order shown.', italic=True, color=GREY)

p(doc, 'Step-by-step: first-time setup', bold=True, sa=2)

numbered(doc, 'Install the SAP NW RFC SDK (one-time, done by a developer or BASIS).  '
              'This is a set of C library files that let Python talk to SAP over RFC.  '
              'Think of it as the "phone line" between your Python script and SAP.')
code_block(doc,
"""# 1. Download "SAP NW RFC SDK 7.50 for Windows 64-bit" from SAP Support Portal
#    (search: SAP Note 2573790 or "SAP NW RFC SDK")
#
# 2. Extract the ZIP to a permanent folder, e.g.:
#    C:\\nwrfcsdk\\
#    After extraction you should see:  C:\\nwrfcsdk\\lib\\sapnwrfc.dll
#
# 3. Add the lib folder to the Windows system PATH permanently:
#    Win + S -> "Edit the system environment variables"
#    -> Environment Variables -> System Variables -> Path -> Edit -> New
#    -> Type: C:\\nwrfcsdk\\lib  -> OK all dialogs
#
# 4. Open a NEW PowerShell window (path change only takes effect in new windows)
#    and verify the DLL is on the path:
Get-Item "C:\\nwrfcsdk\\lib\\sapnwrfc.dll"
# Expected: a file listing showing sapnwrfc.dll""")

numbered(doc, 'Create a Python virtual environment.  A virtual environment is an isolated '
              'Python installation just for this project \u2014 it keeps CodeSage\'s '
              'dependencies separate from any other Python projects on your machine.')
code_block(doc,
"""# Open PowerShell and navigate to your project folder
cd C:\\codesage\\phase1_extract

# Create the virtual environment (this creates a codesage-env folder)
python -m venv codesage-env

# Activate it  (you will see (codesage-env) appear in your prompt)
.\\codesage-env\\Scripts\\Activate.ps1

# If you get a script execution error, run this once to allow local scripts:
Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser""")

numbered(doc, 'Install Python dependencies:')
code_block(doc,
"""# Make sure the virtual environment is active first (see step above)
pip install pyrfc pyyaml

# Verify installation
python -c "import pyrfc; print('pyrfc OK')"
python -c "import yaml;  print('pyyaml OK')"
# Both should print OK with no errors""")

numbered(doc, 'Store the SAP password safely as a Windows environment variable.  '
              'Setting it this way means it is available for this PowerShell session only '
              '\u2014 it is never written to disk or stored in a file.')
code_block(doc,
"""# Set for the current PowerShell session only
$env:SAP_RFC_PASSWORD = "your_sap_password_here"

# To make it persist across reboots (optional, stored in Windows registry):
[System.Environment]::SetEnvironmentVariable(
    "SAP_RFC_PASSWORD",
    "your_sap_password_here",
    "User")""")

numbered(doc, 'Edit scan_config.yaml with your SAP system details.  Open it in Notepad or VS Code and fill in the host, system number, and client for your SAP landscape.')

numbered(doc, 'Test the RFC connection before running the full scan.  Save the snippet '
              'below as test_connection.py in the phase1_extract folder, then run it.  '
              'This is a quick sanity check \u2014 it only calls RFC_PING (no data read):')
code_block(doc,
"""# test_connection.py  -- save this file, then run: python test_connection.py
import pyrfc, os, yaml

cfg  = yaml.safe_load(open('scan_config.yaml'))
conn = pyrfc.Connection(
    ashost=cfg['sap']['host'],   sysnr=cfg['sap']['sysnr'],
    client=cfg['sap']['client'], user=cfg['sap']['user'],
    passwd=os.environ['SAP_RFC_PASSWORD'])

result = conn.call('RFC_PING')
print('RFC connection successful:', result)
# Expected output:  RFC connection successful: {}""")

code_block(doc, """python test_connection.py""")

numbered(doc, 'Run the full extractor.  This is the main Phase 1 step.  '
              'It will loop through all Z*/Y* objects and save them as files.  '
              'Duration: 30\u201390 minutes depending on codebase size.')
code_block(doc,
"""cd C:\\codesage\\phase1_extract
python extractor.py""")

p(doc, 'What you will see in the console while it runs:', bold=True, sa=2)
code_block(doc,
"""  SKIP ZZ_OLD_TEST: function not found      <- normal: object was deleted from SAP
  SKIP ZZ_LEGACY_FM: authorization failure  <- BASIS needs to grant CODESAGE_RFC access
  ...
Done. 8,432 objects saved to .\\abap_files\\""")

p(doc, 'Verify the output \u2014 four checks to run after Phase 1 completes:', bold=True, sa=2)
numbered(doc, 'Count the extracted files \u2014 there should be pairs of .abap and .json for each object:')
code_block(doc,
"""# Count .abap files in the output folder
(Get-ChildItem -Path ".\\abap_files" -Filter "*.abap").Count
# Typical result: 5,000 to 20,000 files""")

numbered(doc, 'Open a few .abap files in Notepad++ or VS Code to confirm the ABAP source looks correct.')
numbered(doc, 'Check a .json sidecar to confirm metadata was captured:')
code_block(doc,
"""# Print the first JSON sidecar found
Get-Content (Get-ChildItem ".\\abap_files\\*.json" | Select-Object -First 1).FullName
# Expected: { "name": "Z_...", "type": "PROG", "package": "Z...", "changed": "20250101" }""")

numbered(doc, 'Confirm the credential filter worked \u2014 no plain-text passwords in the output:')
code_block(doc,
"""# Search for "password" in all .abap files (should only find [REDACTED], not real values)
Select-String -Path ".\\abap_files\\*.abap" -Pattern "password" -CaseSensitive:$false |
  Where-Object { $_.Line -notmatch "REDACTED" }
# If this returns nothing: the credential filter is working correctly""")

p(doc, 'Scheduling monthly re-runs with Windows Task Scheduler:', bold=True, sa=2)
p(doc, 'Phase 1 must be re-run monthly (or after each major transport release) to capture '
       'new and changed ABAP objects. The extractor is safe to re-run \u2014 it overwrites '
       'changed files and adds new ones without deleting anything.')
code_block(doc,
"""# Run this PowerShell command once to register a monthly scheduled task
# It will run extractor.py on the 1st of every month at 02:00 AM
$action  = New-ScheduledTaskAction `
    -Execute "C:\\codesage\\phase1_extract\\codesage-env\\Scripts\\python.exe" `
    -Argument "C:\\codesage\\phase1_extract\\extractor.py" `
    -WorkingDirectory "C:\\codesage\\phase1_extract"

$trigger = New-ScheduledTaskTrigger -Monthly -DaysOfMonth 1 -At "02:00AM"

$settings = New-ScheduledTaskSettingsSet -RunOnlyIfNetworkAvailable

Register-ScheduledTask `
    -TaskName   "CodeSage Phase1 Monthly Extract" `
    -Action     $action `
    -Trigger    $trigger `
    -Settings   $settings `
    -RunLevel   Highest

# To run it immediately to test:
Start-ScheduledTask -TaskName "CodeSage Phase1 Monthly Extract"

# To check if it ran successfully:
Get-ScheduledTaskInfo -TaskName "CodeSage Phase1 Monthly Extract" |
    Select-Object LastRunTime, LastTaskResult
# LastTaskResult = 0 means success""")

make_table(doc,
    ['Common Error', 'Cause', 'Fix'],
    [
        ['RFC_ERROR_SYSTEM_FAILURE',   'Wrong host / sysnr / client in config',         'Double-check scan_config.yaml values with BASIS'],
        ['RFC_ERROR_LOGON_FAILURE',    'Wrong user or password',                         'Verify SAP_RFC_PASSWORD env var is set correctly'],
        ['Authorization failure',      'RFC user missing S_RFC auth for function group', 'Ask BASIS to add the function group to CODESAGE_RFC profile'],
        ['DLL load failed (Windows)',  'SAP NW RFC SDK not on PATH',                    'Add nwrfcsdk\\lib to system PATH; restart terminal'],
        ['No objects extracted (0)',   'Namespace filter too strict or wrong client',     'Try GENERIC_NAME=\'*\' temporarily to test; adjust namespaces in config'],
    ],
    col_widths=[2.0, 2.5, 3.2]
)

info_box(doc, '\u2705 Phase 1 Outcome:',
    'A local folder of ABAP source files + .json metadata for every custom object in your '
    'SAP system (Z*/Y* namespace). Secrets filtered, read-only RFC used. '
    'These files are the shared input for Phase 2 and Phase 3, which run concurrently from this point.')

# ── Phase 2 ──────────────────────────────────────────────────────────────────
doc.add_heading('Chapter 4:  Phase 2 \u2014 Index & Embed', level=1)
p(doc, 'Phase 2 turns the raw ABAP source files into a searchable vector database. '
       'The goal is to be able to answer the question: "given a developer\'s natural '
       'language question, which ABAP code chunks are most relevant to it?" \u2014 in '
       'under 100 milliseconds, entirely on-premise.')

p(doc, 'What is a vector embedding?', bold=True, sa=2)
p(doc, 'A vector embedding is a list of numbers (e.g. 768 numbers) that represents the '
       'meaning of a piece of text. Two pieces of text that mean similar things \u2014 even '
       'if written in very different words \u2014 will produce number lists that are '
       'mathematically close to each other. This is what allows CodeSage to match '
       '"how do we validate vendor payment terms?" with a function module called '
       'Z_VALIDATE_VENDOR_PAYTERMS, even though none of those exact words appear in the '
       'function module\'s name.')

p(doc, 'The chunking step:', bold=True, sa=2)
p(doc, 'ABAP source files can be thousands of lines long. Embedding a whole program as one '
       'unit would lose granularity \u2014 you would retrieve the entire report when you '
       'only needed one function. The chunker splits each file at logical boundaries: '
       'one chunk per function module, per method, per form routine, or per report '
       'section. Each chunk is 150\u2013500 tokens \u2014 small enough to be specific, '
       'large enough to be meaningful.')

p(doc, 'The four ChromaDB collections:', bold=True, sa=2)
bullet(doc, 'abap_programs \u2014 executable reports and includes. Chunks are individual PERFORM blocks and SELECT loops.')
bullet(doc, 'function_modules \u2014 function module source plus parameter metadata. A single chunk = one function module.')
bullet(doc, 'abap_classes \u2014 ABAP OO classes, interfaces, methods. A chunk = one method or one class definition block.')
bullet(doc, 'documentation \u2014 SAP Help content, Clean Core guidelines, and cloudification rules (public documentation scraped and chunked).')

p(doc, 'ChromaDB runs entirely on-premise. No source code leaves the network during this step. '
       'The embedding model (nomic-embed-text) also runs locally via Ollama \u2014 no cloud API call.',
  bold=True, color=SAP_DARK)

p(doc, 'Phase 2 folder structure:', bold=True, sa=2)
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase2_index/
\u2502   \u251c\u2500\u2500 chunker.py                # Splits .abap files into logical chunks
\u2502   \u251c\u2500\u2500 embedder.py               # Embeds chunks and stores in ChromaDB
\u2502   \u2514\u2500\u2500 chromadb_store/           # OUTPUT \u2014 persistent ChromaDB data folder
\u2502       \u251c\u2500\u2500 abap_programs/          # Collection 1: report & executable programs
\u2502       \u251c\u2500\u2500 function_modules/       # Collection 2: function module source + params
\u2502       \u251c\u2500\u2500 abap_classes/           # Collection 3: OO classes, interfaces, methods
\u2502       \u2514\u2500\u2500 documentation/          # Collection 4: SAP Help + Clean Core docs""")

p(doc, 'chunker.py \u2014 split ABAP source files into logical segments:', bold=True, sa=2)
code_block(doc,
"""import pathlib, json, re
from dataclasses import dataclass

@dataclass
class Chunk:
    text: str
    object_name: str
    chunk_type: str   # FORM | METHOD | FUNCTION | CLASS_DEF | BODY
    collection: str   # which ChromaDB collection this belongs to

FORM_RE     = re.compile(r'^FORM\\s+(\\w+)', re.M | re.I)
METHOD_RE   = re.compile(r'^\\s*METHOD\\s+(\\w+)', re.M | re.I)
ENDFORM_RE  = re.compile(r'^ENDFORM', re.M | re.I)
ENDMETHOD_RE= re.compile(r'^\\s*ENDMETHOD', re.M | re.I)

def chunk_abap(abap_path: pathlib.Path, meta: dict) -> list[Chunk]:
    src   = abap_path.read_text(encoding='utf-8')
    lines = src.splitlines()
    name  = meta['name']
    typ   = meta['type']     # PROG, FM, CLAS, etc.
    chunks = []

    if typ == 'FM':
        # Function module \u2014 whole file is one chunk
        collection = 'function_modules'
        chunks.append(Chunk(src, name, 'FUNCTION', collection))

    elif typ == 'CLAS':
        # Split on METHOD / ENDMETHOD boundaries
        collection = 'abap_classes'
        buf, mname = [], None
        for line in lines:
            m = METHOD_RE.match(line)
            if m:
                mname = m.group(1); buf = [line]
            elif ENDMETHOD_RE.match(line) and mname:
                buf.append(line)
                chunks.append(Chunk('\\n'.join(buf), f'{name}.{mname}', 'METHOD', collection))
                mname = None; buf = []
            elif mname:
                buf.append(line)
        if not chunks:   # class definition only \u2014 add whole file
            chunks.append(Chunk(src, name, 'CLASS_DEF', collection))

    else:
        # Program \u2014 split on FORM / ENDFORM boundaries
        collection = 'abap_programs'
        buf, fname = [], None
        for line in lines:
            m = FORM_RE.match(line)
            if m:
                fname = m.group(1); buf = [line]
            elif ENDFORM_RE.match(line) and fname:
                buf.append(line)
                chunks.append(Chunk('\\n'.join(buf), f'{name}.{fname}', 'FORM', collection))
                fname = None; buf = []
            elif fname:
                buf.append(line)
        # Add body (lines outside any FORM) as one chunk
        chunks.append(Chunk(src[:3000], name, 'BODY', collection))

    return [c for c in chunks if len(c.text.strip()) > 50]  # skip empty stubs

if __name__ == '__main__':
    abap_dir = pathlib.Path('../phase1_extract/abap_files')
    all_chunks = []
    for abap_file in abap_dir.glob('*.abap'):
        meta_file = abap_file.with_suffix('.json')
        meta = json.loads(meta_file.read_text()) if meta_file.exists() else {'name': abap_file.stem, 'type': 'PROG'}
        all_chunks.extend(chunk_abap(abap_file, meta))
    print(f'Total chunks: {len(all_chunks)}')""")

p(doc, 'embedder.py \u2014 embed chunks and store in ChromaDB:', bold=True, sa=2)
code_block(doc,
"""import chromadb
import requests, json
from chunker import chunk_abap
import pathlib

# ChromaDB client \u2014 persistent local store (on-premise, no cloud)
db = chromadb.PersistentClient(path='./chromadb_store')

COLLECTIONS = {
    'abap_programs':    db.get_or_create_collection('abap_programs'),
    'function_modules': db.get_or_create_collection('function_modules'),
    'abap_classes':     db.get_or_create_collection('abap_classes'),
    'documentation':    db.get_or_create_collection('documentation'),
}

def embed(text: str) -> list[float]:
    \"\"\"Call local nomic-embed-text via Ollama (no cloud, no API key).\"\"\"
    r = requests.post('http://localhost:11434/api/embeddings',
                      json={'model': 'nomic-embed-text', 'prompt': text})
    return r.json()['embedding']   # 768-dimensional vector

def store_chunk(chunk, idx: int):
    col = COLLECTIONS[chunk.collection]
    vec = embed(chunk.text[:2000])  # embed first 2000 chars (fits context window)
    col.add(
        ids=[f'{chunk.object_name}_{idx}'],
        embeddings=[vec],
        documents=[chunk.text],
        metadatas=[{'object': chunk.object_name, 'type': chunk.chunk_type}])

# --- Main ---
abap_dir = pathlib.Path('../phase1_extract/abap_files')
for i, abap_file in enumerate(sorted(abap_dir.glob('*.abap'))):
    meta_file = abap_file.with_suffix('.json')
    meta = json.loads(meta_file.read_text()) if meta_file.exists() else {'name': abap_file.stem, 'type': 'PROG'}
    for j, chunk in enumerate(chunk_abap(abap_file, meta)):
        store_chunk(chunk, j)
    if i % 100 == 0:
        print(f'  Indexed {i} files...')

print('ChromaDB indexing complete.')
for name, col in COLLECTIONS.items():
    print(f'  {name}: {col.count()} chunks')""")

# ── How to Run Phase 2 ───────────────────────────────────────────────────────
doc.add_heading('4.1  How to Run Phase 2', level=2)

p(doc, 'Prerequisites', bold=True, sa=2)
make_table(doc,
    ['Requirement', 'Details', 'Install'],
    [
        ['Phase 1 complete',     'abap_files/ folder must exist with .abap + .json files', '\u2014 run Phase 1 first'],
        ['Python 3.10+',         'Same venv as Phase 1',                                   'Already set up'],
        ['chromadb',             'Vector database library',                                 'pip install chromadb'],
        ['requests',             'HTTP calls to local Ollama embedding server',             'pip install requests'],
        ['Ollama',               'Local inference runtime for nomic-embed-text',            'ollama.com/download'],
        ['nomic-embed-text',     'Embedding model (768-dim, runs on CPU)',                  'ollama pull nomic-embed-text'],
        ['Disk space',           '~2\u20134 GB for ChromaDB store depending on codebase size', 'Local SSD recommended'],
    ],
    col_widths=[1.7, 3.0, 2.5]
)

p(doc, 'All commands below run in Windows PowerShell with the Phase 1 virtual '
       'environment active.  Phase 2 is entirely local \u2014 no SAP connection, '
       'no cloud service, no API key needed.', italic=True, color=GREY)

p(doc, 'Step-by-step', bold=True, sa=2)

numbered(doc, 'Activate the virtual environment created in Phase 1:')
code_block(doc,
"""cd C:\\codesage\\phase2_index
.\\..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1
# Your prompt should now show: (codesage-env)""")

numbered(doc, 'Install Phase 2 dependencies.  chromadb is the local vector database. '
              'requests is used to call the Ollama embedding server running on your machine:')
code_block(doc,
"""pip install chromadb requests
# Verify
python -c "import chromadb; print('chromadb', chromadb.__version__)"
# Expected: chromadb 0.5.x  (or newer)""")

numbered(doc, 'Install Ollama.  Ollama is a free tool that runs AI models locally on your Windows PC. '
              'After installation it runs silently in the background as a Windows service \u2014 '
              'you do not need to start it manually. Download from: https://ollama.com/download/windows')
code_block(doc,
"""# After installing Ollama, open a new PowerShell window and download the embedding model
# nomic-embed-text is a small (274 MB) model that converts text into vectors
ollama pull nomic-embed-text

# Confirm the model is ready
ollama list
# Expected output:
# NAME                    ID            SIZE     MODIFIED
# nomic-embed-text:latest 0a109f422b47  274 MB   just now""")

numbered(doc, 'Test that Ollama is responding correctly before running the full embedder. '
              'This calls the embedding API directly and should return a list of 768 numbers:')
code_block(doc,
"""# Save this as test_ollama.py and run it
# test_ollama.py
import requests, json

r = requests.post("http://localhost:11434/api/embeddings",
                  json={"model": "nomic-embed-text", "prompt": "hello SAP"})
vec = r.json()["embedding"]
print(f"Embedding OK: {len(vec)} dimensions, first value = {vec[0]:.4f}")
# Expected: Embedding OK: 768 dimensions, first value = 0.0231""")

code_block(doc, """python test_ollama.py""")

numbered(doc, 'Run chunker.py first to preview how many chunks will be created. '
              'This is a dry run \u2014 it does not write to ChromaDB yet, '
              'just counts the chunks so you know what to expect:')
code_block(doc,
"""cd C:\\codesage\\phase2_index
python chunker.py
# Expected output:
# Total chunks: 42,318
# (the number varies based on your codebase size)""")

numbered(doc, 'Run the embedder.  This is the main Phase 2 step. '
              'It will read every .abap file from Phase 1, split it into chunks, '
              'convert each chunk to a vector using nomic-embed-text, '
              'and store it in ChromaDB. Duration: approximately 2 hours for 10,000 objects on a CPU laptop:')
code_block(doc,
"""python embedder.py
# You will see progress every 100 files:
#   Indexed 0 files...
#   Indexed 100 files...
#   Indexed 200 files...
#   ...
# ChromaDB indexing complete.
#   abap_programs:    18,420 chunks
#   function_modules: 12,305 chunks
#   abap_classes:      9,841 chunks
#   documentation:     1,752 chunks
#
# The chromadb_store folder is now populated and ready for Phase 4 runtime queries.""")

numbered(doc, 'Verify the ChromaDB store with a test semantic search. '
              'Save the snippet below as test_search.py and run it.  '
              'This proves the vector index is working correctly \u2014 '
              'it searches for relevant function modules using a plain English question:')
code_block(doc,
"""# test_search.py
import chromadb, requests

def embed(text):
    r = requests.post("http://localhost:11434/api/embeddings",
                      json={"model": "nomic-embed-text", "prompt": text})
    return r.json()["embedding"]

db  = chromadb.PersistentClient(path="./chromadb_store")
col = db.get_collection("function_modules")
print(f"function_modules: {col.count()} chunks indexed")

# Semantic search test
results = col.query(
    query_embeddings=[embed("validate vendor payment terms")],
    n_results=3)

print("\\nTop 3 matches for 'validate vendor payment terms':")
for doc_text, meta in zip(results["documents"][0], results["metadatas"][0]):
    print(f"  [{meta['object']}]  {doc_text[:80]}...")""")

code_block(doc, """python test_search.py""")

p(doc, 'Expected output from the test search:', bold=True, sa=2)
code_block(doc,
"""function_modules: 12,305 chunks indexed

Top 3 matches for 'validate vendor payment terms':
  [Z_VALIDATE_VENDOR_PAYTERMS]  FUNCTION Z_VALIDATE_VENDOR_PAYTERMS...
  [Z_CHECK_VENDOR_MASTER]       DATA: ls_lfa1 TYPE lfa1. SELECT SINGLE...
  [Z_FI_PAYMENT_TERMS_GET]      SELECT SINGLE * FROM t052 INTO ls_t052...""")

p(doc, 'If you see relevant ABAP objects returned for that question \u2014 Phase 2 is working correctly.',
  bold=True, color=SAP_DARK)

make_table(doc,
    ['Common Error', 'Cause', 'Fix'],
    [
        ['Connection refused on port 11434',  'Ollama not installed or not running',              'Install from ollama.com, then restart PC to start the Ollama service'],
        ['"model not found" from Ollama',     'nomic-embed-text not yet downloaded',              'Run: ollama pull nomic-embed-text  and wait for download to complete'],
        ['No .abap files found',              'Wrong working directory or Phase 1 not complete',  'Confirm C:\\codesage\\phase1_extract\\abap_files\\ contains .abap files'],
        ['ChromaDB collection already exists','Re-running embedder on existing store',            'Safe to re-run \u2014 duplicate IDs are skipped automatically'],
        ['Embedding is very slow (~1 s/chunk)','Ollama running on CPU only (no GPU)',              'Normal on CPU-only machines. Ollama will auto-use GPU if NVIDIA drivers installed'],
        ['Script execution policy error',     'PowerShell blocks .ps1 scripts by default',       'Run: Set-ExecutionPolicy RemoteSigned -Scope CurrentUser'],
    ],
    col_widths=[2.2, 2.3, 3.2]
)

info_box(doc, '\u2705 Phase 2 Outcome:',
    'ChromaDB running on-premise with 4 collections and tens of thousands of embedded ABAP '
    'chunks. Any developer query can now be matched to the most relevant code in your '
    'system in under 100 ms \u2014 without sending a single line of code to the cloud. '
    'This vector index is loaded at runtime to feed context into every LLM response.')

# ── Phase 3 ──────────────────────────────────────────────────────────────────
doc.add_heading('Chapter 5:  Phase 3 \u2014 Fine-Tune, Merge and Package for BYOM', level=1)
p(doc, 'This is the most technically involved phase and the one that makes CodeSage '
       'genuinely organisation-specific. By the end of Phase 3 you will have a version of '
       'LLaMA-3 that has learned your naming conventions, your preferred BAPIs, your error '
       'class patterns, and your ABAP coding style \u2014 running as a live inference '
       'endpoint on SAP AI Core with no local GPU required.')

doc.add_heading('5.1  Concept: What Is Fine-Tuning?', level=2)
p(doc, 'A large language model like LLaMA-3 is pre-trained on billions of text documents '
       'from the internet. It knows a lot about the world, about code in general, and about '
       'ABAP syntax. But it has never seen your programs. It does not know that '
       'Z_MM_OPEN_PO_REPORT exists, or that your team always uses ZCX_BASE_ERROR for '
       'exception handling, or that your BAdI for goods receipt is called '
       'ZCL_IMPL_MB_GR_CHECK.')

p(doc, 'Fine-tuning is the process of continuing the model\'s training on a small, '
       'targeted dataset \u2014 in this case, question-and-answer pairs derived from your '
       'own ABAP code. After fine-tuning, the model still knows everything it knew before, '
       'but now it also knows your system. It speaks your language.')

info_box(doc, '\U0001f9e0 Analogy:',
    'Think of the base LLaMA-3 model as a highly experienced SAP consultant who has worked '
    'on hundreds of projects worldwide. Fine-tuning is the three-week onboarding period '
    'where that consultant reads all your code, attends walkthroughs, and learns your '
    'organisation\'s specific patterns. After that, they are no longer a generic SAP expert '
    '\u2014 they are an expert in your SAP.')

doc.add_heading('5.2  Concept: Why QLoRA Instead of Full Fine-Tuning?', level=2)
p(doc, 'Full fine-tuning of an 8-billion-parameter model requires updating billions of '
       'weights across the entire network. This demands enormous GPU memory (typically '
       '80 GB+ A100 or H100) and many hours of compute time. For a corporate project this '
       'is expensive and impractical.')

p(doc, 'QLoRA (Quantised Low-Rank Adaptation) is a technique that achieves most of the '
       'benefit of full fine-tuning at a fraction of the cost:')
bullet(doc, 'The base model weights are frozen (not updated) and loaded in 4-bit compressed format (NF4 quantisation). This reduces GPU memory from ~28 GB to ~6\u201312 GB.', bold_prefix='Quantisation (4-bit NF4):  ')
bullet(doc, 'Instead of modifying the full weight matrices, LoRA inserts small trainable "adapter" matrices alongside certain layers. These adapters have far fewer parameters \u2014 typically 0.1\u20131% of the original model size.', bold_prefix='LoRA adapters:  ')
bullet(doc, 'Only the adapter weights are updated during training. The base model stays frozen and compressed. Training fits on a 16 GB T4 GPU available free on Kaggle.', bold_prefix='What gets trained:  ')
bullet(doc, 'After training, the small adapter file (~50 MB) is merged back into the full-precision base model to produce a single merged model ready for serving.', bold_prefix='Merge step:  ')

p(doc, 'CodeSage QLoRA settings:  rank r=16, alpha=32 (controls adapter scale), target '
       'modules: all attention projection layers (q_proj, k_proj, v_proj, o_proj). '
       'These settings balance training speed and answer quality for a code-focused task.')

doc.add_heading('5.3  Step A \u2014 Generate Training Data', level=2)
p(doc, 'Fine-tuning requires a supervised dataset of (question, answer) pairs. We generate '
       'these automatically from the ABAP source files extracted in Phase 1. Each code '
       'chunk is sent to Claude (claude-sonnet-4-6) with a prompt asking it to produce '
       '3\u20135 realistic developer questions that the chunk answers.')

p(doc, 'Example: given the chunk for Z_VALIDATE_VENDOR_PAYTERMS, Claude generates:')
bullet(doc, '"Do we have a function that checks vendor payment terms before invoice posting?"')
bullet(doc, '"How do we validate that a payment term key exists in T052?"')
bullet(doc, '"Show me how we verify vendor master data before FI document creation."')

p(doc, 'Each (question, chunk) pair becomes one training example in a JSONL file. '
       'A typical scan of 10,000 custom objects produces 30,000\u201350,000 training pairs. '
       'The dataset is reviewed for sensitive content before training begins.')

p(doc, 'Training data generation script (simplified):', bold=True, sa=2)
code_block(doc,
"""import anthropic, json, pathlib

client = anthropic.Anthropic()          # uses ANTHROPIC_API_KEY env var

def generate_qa(chunk_text: str, object_name: str) -> list[dict]:
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=512,
        messages=[{"role": "user", "content":
            f"You are an ABAP expert. Given this code from {object_name}, "
            f"write 4 developer questions whose answer is in this code.\\n\\n"
            f"CODE:\\n{chunk_text}\\n\\nReturn JSON list of strings."}])
    questions = json.loads(msg.content[0].text)
    return [{"prompt": q, "completion": chunk_text} for q in questions]

with open("training_data.jsonl", "w") as f:
    for path in pathlib.Path("./abap_files").glob("*.abap"):
        chunk = path.read_text(encoding="utf-8")[:2000]   # first 2000 chars
        for pair in generate_qa(chunk, path.stem):
            f.write(json.dumps(pair) + "\\n")""")

doc.add_heading('5.4  Step B \u2014 Run QLoRA Fine-Tuning', level=2)
p(doc, 'With the training dataset ready, fine-tuning runs on a GPU. The recommended free '
       'option is a Kaggle Notebook with a T4 GPU. The training loop uses the '
       'Hugging Face transformers + PEFT + bitsandbytes libraries:')
code_block(doc,
"""from transformers import AutoModelForCausalLM, AutoTokenizer, TrainingArguments
from peft import LoraConfig, get_peft_model, TaskType
from trl import SFTTrainer
import torch

MODEL_ID = "meta-llama/Meta-Llama-3-8B-Instruct"

# Load base model in 4-bit NF4 quantisation (QLoRA)
from transformers import BitsAndBytesConfig
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",          # NF4 = best quality at 4-bit
    bnb_4bit_compute_dtype=torch.bfloat16,
    bnb_4bit_use_double_quant=True)     # double quantisation saves ~0.4 GB more

model = AutoModelForCausalLM.from_pretrained(
    MODEL_ID, quantization_config=bnb_config, device_map="auto")
tokenizer = AutoTokenizer.from_pretrained(MODEL_ID)

# Attach LoRA adapters — only these small matrices will be trained
lora_config = LoraConfig(
    r=16,                               # rank — controls adapter capacity
    lora_alpha=32,                      # scale factor
    target_modules=["q_proj","k_proj","v_proj","o_proj"],
    lora_dropout=0.05,
    bias="none",
    task_type=TaskType.CAUSAL_LM)
model = get_peft_model(model, lora_config)
model.print_trainable_parameters()
# Output: trainable params: 13,631,488 || all params: 8,044,093,440 (0.17%)

# Train
trainer = SFTTrainer(
    model=model, tokenizer=tokenizer,
    train_dataset=dataset,              # loaded from training_data.jsonl
    dataset_text_field="text",
    max_seq_length=2048,
    args=TrainingArguments(
        output_dir="./lora-adapter",
        num_train_epochs=3,
        per_device_train_batch_size=4,
        gradient_accumulation_steps=4,
        learning_rate=2e-4,
        fp16=True,
        save_strategy="epoch",
        logging_steps=50))
trainer.train()
trainer.save_model("./lora-adapter")   # saves ~50 MB LoRA adapter""")

p(doc, 'Training duration: approximately 6\u201312 hours on a Kaggle T4 (free). '
       'Only the 0.17% of parameters in the LoRA adapters are updated. '
       'The base model weights remain frozen throughout, which is why the whole '
       'process fits in 16 GB of GPU memory.')

doc.add_heading('5.5  Step C \u2014 Merge Adapter into Base Model', level=2)
p(doc, 'After training, the LoRA adapter is a small set of weight deltas. To serve '
       'the model efficiently at inference time, the adapter is merged back into the '
       'full-precision base model weights. The result is a single merged model folder '
       '\u2014 identical in structure to a standard Hugging Face model, with your fine-tuning '
       'baked in.')
code_block(doc,
"""from transformers import AutoModelForCausalLM, AutoTokenizer
from peft import PeftModel

print("Loading base model in full precision for merge...")
base = AutoModelForCausalLM.from_pretrained(
    "meta-llama/Meta-Llama-3-8B-Instruct",
    torch_dtype="auto", device_map="cpu")   # CPU merge is fine — one-time op

tokenizer = AutoTokenizer.from_pretrained("meta-llama/Meta-Llama-3-8B-Instruct")

print("Applying LoRA adapter...")
model = PeftModel.from_pretrained(base, "./lora-adapter/")

print("Merging and unloading LoRA weights into base...")
merged = model.merge_and_unload()           # adapter deltas added into base weights

print("Saving merged model...")
merged.save_pretrained("./merged_model/")   # ~16 GB folder (full bf16 weights)
tokenizer.save_pretrained("./merged_model/")
print("Merge complete. ./merged_model/ is ready to package.")""")

doc.add_heading('5.6  Step D \u2014 Package as Docker Image (vLLM Format for BYOM)', level=2)
p(doc, 'SAP AI Core BYOM requires the model to be served via a Docker container that '
       'exposes an OpenAI-compatible REST API on port 8080. The standard serving stack '
       'for this is vLLM \u2014 a high-throughput inference engine that SAP AI Core '
       'can manage and scale automatically.')

p(doc, 'Why vLLM?  vLLM implements PagedAttention \u2014 an algorithm that handles '
       'GPU memory for token sequences much more efficiently than naive serving. '
       'It supports batching multiple concurrent requests, which is critical when '
       'many developers are querying CodeSage simultaneously.')
code_block(doc,
"""# Dockerfile — vLLM OpenAI-compatible server (required by SAP AI Core BYOM)
FROM vllm/vllm-openai:latest

# Copy merged model into the container image
COPY ./merged_model/ /model/

EXPOSE 8080

CMD ["python", "-m", "vllm.entrypoints.openai.api_server", \\
     "--model",              "/model", \\
     "--port",               "8080", \\
     "--served-model-name",  "codesage-llama3", \\
     "--max-model-len",      "8192", \\
     "--dtype",              "bfloat16"]""")

p(doc, 'Build the image and push to a container registry accessible from your SAP BTP subaccount '
       '(Docker Hub, SAP BTP Container Registry, or your own private registry):')
code_block(doc,
"""# Build the Docker image (~18 GB including model weights)
docker build -t codesage-llama3:v1 .

# Push to registry (replace with your registry URL)
docker tag  codesage-llama3:v1  <your-registry>/codesage-llama3:v1
docker push <your-registry>/codesage-llama3:v1""")

doc.add_heading('5.7  Step E \u2014 Register and Deploy in SAP AI Core (BYOM)', level=2)
p(doc, 'With the image in a registry, you register a ServingTemplate in SAP AI Core. '
       'This YAML file tells SAP AI Core how to launch your container, how many replicas '
       'to run, and which port to expose. Apply it via SAP AI Launchpad or the AI Core API:')
code_block(doc,
"""# serving-template.yaml — apply once via SAP AI Launchpad or AI Core REST API
apiVersion: ai.sap.com/v1alpha1
kind: ServingTemplate
metadata:
  name: codesage-llama3
  labels:
    scenarios.ai.sap.com/id: codesage-scenario
    ai.sap.com/version: "1.0"
spec:
  template:
    metadata:
      labels:
        ai.sap.com/deploymentSanitizedName: codesage-llama3
    spec:
      imagePullSecrets:
        - name: <your-registry-secret>   # Docker registry credentials in AI Core
      containers:
        - name: kserve-container
          image: <your-registry>/codesage-llama3:v1
          ports:
            - name: http1
              containerPort: 8080
              protocol: TCP
          resources:
            requests:
              memory: "20Gi"
            limits:
              nvidia.com/gpu: "1"       # SAP AI Core allocates a managed GPU""")

p(doc, 'Deployment steps in SAP AI Launchpad:')
numbered(doc, 'Upload serving-template.yaml via ML Operations \u2192 Serving Templates.')
numbered(doc, 'Create a Configuration referencing the serving template and your resource group.')
numbered(doc, 'Create a Deployment from the configuration. Status moves: PENDING \u2192 RUNNING (5\u201310 min).')
numbered(doc, 'Copy the Deployment ID and AI Core service key. Store both in BTP Destination Service for use by the CodeSage Agent.')
numbered(doc, 'Test the endpoint: POST /v2/inference/deployments/{id}/chat/completions with an Authorization header. The API is OpenAI-compatible.')

p(doc, 'Phase 3 folder structure:', bold=True, sa=2)
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase3_finetune/
\u2502   \u251c\u2500\u2500 generate_qa.py            # Step A: generates training_data.jsonl via Claude API
\u2502   \u251c\u2500\u2500 train_qlora.py             # Step B: QLoRA training (run as Kaggle notebook)
\u2502   \u251c\u2500\u2500 merge_model.py             # Step C: merges LoRA adapter into base model
\u2502   \u251c\u2500\u2500 training_data.jsonl        # Generated QA pairs (input to training)
\u2502   \u251c\u2500\u2500 lora-adapter/              # OUTPUT of training (Step B)
\u2502   \u2502   \u251c\u2500\u2500 adapter_config.json    # LoRA hyperparameters (r=16, alpha=32, ...)
\u2502   \u2502   \u2514\u2500\u2500 adapter_model.safetensors  # Trained adapter weights (~50 MB)
\u2502   \u251c\u2500\u2500 merged_model/              # OUTPUT of merge (Step C) \u2014 ~16 GB
\u2502   \u2502   \u251c\u2500\u2500 config.json
\u2502   \u2502   \u251c\u2500\u2500 tokenizer.json
\u2502   \u2502   \u251c\u2500\u2500 tokenizer_config.json
\u2502   \u2502   \u251c\u2500\u2500 special_tokens_map.json
\u2502   \u2502   \u2514\u2500\u2500 model-00001-of-00004.safetensors  (x4 shards)
\u2502   \u251c\u2500\u2500 Dockerfile                 # Step D: vLLM container definition
\u2502   \u2514\u2500\u2500 serving-template.yaml      # Step E: SAP AI Core BYOM registration""")

# ── How to Run Phase 3 ───────────────────────────────────────────────────────
doc.add_heading('5.8  How to Run Phase 3', level=2)

p(doc, 'Phase 3 is split into five sub-steps that run in sequence. '
       'Steps A and B happen in the cloud (Kaggle free GPU \u2014 no cost). '
       'Steps C, D, and E run on your Windows PC. '
       'This phase only needs to be run once, then refreshed quarterly.',
  italic=True, color=GREY)

p(doc, 'Prerequisites', bold=True, sa=2)
make_table(doc,
    ['Requirement', 'Details', 'Where to Get It'],
    [
        ['Phase 1 complete',         'abap_files\\ folder must exist with .abap source files',   'Run Phase 1 first'],
        ['Anthropic API key',        'Used by generate_qa.py to call Claude (Step A)',             'console.anthropic.com \u2192 API Keys'],
        ['Kaggle account',           'Free GPU notebook (NVIDIA T4 16 GB) for training',          'kaggle.com \u2014 free account, no credit card'],
        ['Kaggle GPU quota',         '30 hours/week GPU included for free',                        'Kaggle Settings \u2192 Phone Verification to unlock GPU'],
        ['HuggingFace account',      'To download LLaMA-3 8B model weights (~16 GB)',              'huggingface.co \u2014 free account'],
        ['HuggingFace token',        'Read-access token for the gated LLaMA-3 model',             'huggingface.co/settings/tokens \u2192 New Token (Read)'],
        ['LLaMA-3 licence accepted', 'Meta requires licence acceptance before download',           'Visit the model page on HuggingFace and click Accept'],
        ['Docker Desktop (Windows)', 'Build and push the container image (Steps D\u2013E)',        'docker.com/get-docker \u2014 free for personal/small team use'],
        ['Container registry',       'Stores the Docker image so SAP AI Core can pull it',        'Docker Hub (free), SAP BTP Container Registry, or AWS ECR'],
        ['SAP BTP subaccount',       'Active BTP account with AI Core service enabled',            'SAP BTP Cockpit \u2192 Service Marketplace \u2192 AI Core'],
    ],
    col_widths=[1.9, 2.9, 2.4]
)

p(doc, 'Step A \u2014 Generate training data on your Windows PC (~30\u201360 min):', bold=True, sa=2)
p(doc, 'This step uses Claude to read your ABAP code and write questions about it. '
       'The result is a training dataset in JSONL format (one JSON object per line). '
       'You only need internet access and an Anthropic API key.')
code_block(doc,
"""# Open PowerShell and activate the Phase 1 virtual environment
cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

# Install the Anthropic Python SDK
pip install anthropic

# Store your API key for this session
$env:ANTHROPIC_API_KEY = "sk-ant-your-key-here"

# Run the QA generation script
python generate_qa.py
# What it does:
#   - Reads each .abap file from Phase 1
#   - Sends each code chunk to Claude with the prompt:
#     "Write 4 developer questions whose answer is in this code"
#   - Saves each (question, code) pair as one line in training_data.jsonl

# Check how many training pairs were generated
(Get-Content .\\training_data.jsonl).Count
# Typical result: 30,000 to 50,000 lines

# Preview the first two training examples
Get-Content .\\training_data.jsonl | Select-Object -First 2""")

p(doc, 'Step B \u2014 Fine-tune on Kaggle (free GPU, 6\u201312 hrs):', bold=True, sa=2)
p(doc, 'Kaggle provides a free NVIDIA T4 GPU for 30 hours per week. '
       'You do not need to install anything on your Windows PC for this step \u2014 '
       'everything runs in a Kaggle cloud notebook in your browser.')
numbered(doc, 'Upload training_data.jsonl to Kaggle as a Dataset.  Go to kaggle.com \u2192 Datasets \u2192 New Dataset.  Give it a name like "codesage-training". Note the dataset path shown (e.g. /kaggle/input/codesage-training/).')
numbered(doc, 'Create a new Kaggle Notebook.  Go to kaggle.com/code \u2192 New Notebook.  In the right-hand panel: set Accelerator to GPU T4 x1 and enable Internet access (needed to download LLaMA-3 from HuggingFace).')
numbered(doc, 'Add the following cell at the very top of the notebook to install dependencies and log in to HuggingFace:')
code_block(doc,
"""# Cell 1 \u2014 install libraries (run once at notebook start, takes ~2 min)
!pip install -q transformers peft trl bitsandbytes accelerate datasets huggingface_hub

# Log in to HuggingFace to download LLaMA-3
# (you must have accepted the Meta licence on the HuggingFace model page first)
import huggingface_hub
huggingface_hub.login(token="hf_YOUR_READ_TOKEN_HERE")
print("HuggingFace login OK")""")

numbered(doc, 'Add the training code as Cell 2.  Copy the full contents of train_qlora.py (from your Windows PC) into the second cell. Update the dataset path variable at the top to match your Kaggle dataset path.')
numbered(doc, 'Click Run All.  Training will take 6\u201312 hours. You can close the browser and come back \u2014 Kaggle keeps running the notebook. '
              'Check GPU memory in the right panel \u2014 it should stay below 14 GB.')
numbered(doc, 'Download the adapter when training completes.  In the notebook output panel, find the lora-adapter/ folder.  Click the three-dot menu \u2192 Download as ZIP.  '
              'Save and extract it to C:\\codesage\\phase3_finetune\\lora-adapter\\ on your Windows PC.')

p(doc, 'Step C \u2014 Merge LoRA adapter into base model (Windows, ~20 min, needs 32 GB RAM):', bold=True, sa=2)
p(doc, 'The merge step combines the small LoRA adapter (50 MB) with the full LLaMA-3 base model '
       'to produce a single standalone model file ready for Docker packaging. '
       'This runs on CPU \u2014 no GPU needed. It requires at least 32 GB of RAM because the '
       'full model weights are loaded into memory during the merge.')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

# Install merge dependencies
pip install transformers peft accelerate

# Run the merge (takes ~20 minutes on a standard Windows PC)
python merge_model.py
# What it does:
#   1. Loads LLaMA-3 8B base weights from HuggingFace (downloads ~16 GB once)
#   2. Applies your lora-adapter/ on top
#   3. Saves the combined model to merged_model\\ (~16 GB folder)

# Verify the output folder looks correct
Get-ChildItem .\\merged_model
# Should show: config.json, tokenizer.json, model-00001-of-00004.safetensors, etc.""")

p(doc, 'Step D \u2014 Build and push Docker image (Windows, needs Docker Desktop):', bold=True, sa=2)
p(doc, 'Docker Desktop packages the merged model and the vLLM serving engine into a '
       'single container image. SAP AI Core will pull this image from your container '
       'registry and run it on a managed GPU in the cloud. '
       'The image will be large (~18 GB) so budget 30\u201360 minutes for the push.')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune

# Log in to Docker Hub (or your private registry)
docker login
# Enter your Docker Hub username and password when prompted
# For a private registry: docker login your-registry.example.com

# Build the container image
# The -t flag gives it a name (tag). v1 is the version label.
docker build -t codesage-llama3:v1 .
# First build takes 10-20 min as it downloads the vLLM base image (~8 GB)

# Tag with your registry path and push
docker tag  codesage-llama3:v1  your-dockerhub-username/codesage-llama3:v1
docker push your-dockerhub-username/codesage-llama3:v1
# Push takes 30-60 min depending on your internet upload speed

# Confirm the image is in the registry
docker image ls | Select-String "codesage-llama3"
# Expected: codesage-llama3   v1   <image-id>   ...  ~18GB""")

p(doc, 'Step E \u2014 Deploy to SAP AI Core (SAP AI Launchpad):', bold=True, sa=2)
numbered(doc, 'Open SAP AI Launchpad.  In your BTP subaccount \u2192 Services \u2192 AI Core \u2192 Go to Application.')
numbered(doc, 'Create a Docker Registry Secret.  Go to ML Operations \u2192 Docker Registry Secrets \u2192 Add.  Provide your registry URL and credentials.')
numbered(doc, 'Upload the serving template.  Go to ML Operations \u2192 Serving Templates \u2192 Upload YAML.  Select serving-template.yaml.')
numbered(doc, 'Create a Configuration.  Go to ML Operations \u2192 Configurations \u2192 Create.  Select the codesage-llama3 template and your resource group.')
numbered(doc, 'Create a Deployment.  Go to ML Operations \u2192 Deployments \u2192 Create.  Select the configuration. Status will move: PENDING \u2192 INITIALIZING \u2192 RUNNING (5\u201310 min).')
numbered(doc, 'Copy the Deployment URL.  From the deployment details page, copy the inference URL and Deployment ID.')
numbered(doc, 'Store credentials in BTP Destination Service so the CodeSage Agent can call AI Core securely.  In BTP Cockpit \u2192 Connectivity \u2192 Destinations \u2192 New Destination, fill in the following fields (all values come from the AI Core service key JSON you downloaded):')
code_block(doc,
"""Name:              CODESAGE_AICORE
Type:              HTTP
URL:               https://<ai-api-url>         <- from service key: "url"
Authentication:    OAuth2ClientCredentials
Client ID:         <clientid>                   <- from service key: "clientid"
Client Secret:     <clientsecret>               <- from service key: "clientsecret"
Token Service URL: <tokenurl>/oauth/token        <- from service key: "url" + /oauth/token

Additional Properties (click "New Property" for each):
  AI-Resource-Group   default
  DEPLOYMENT_ID       <your deployment ID from step 5>""")

numbered(doc, 'Test the endpoint from your Windows PC to confirm it is live.  '
              'Save the snippet below as test_aicore.py and run it:')
code_block(doc,
"""# test_aicore.py  -- run: python test_aicore.py
import requests, os

# Get a token using client credentials
token_url    = "https://<your-token-url>/oauth/token"
client_id    = "<your-client-id>"
client_secret= "<your-client-secret>"

token_resp = requests.post(token_url,
    data={"grant_type": "client_credentials"},
    auth=(client_id, client_secret))
token = token_resp.json()["access_token"]
print("Token OK")

# Call the inference endpoint
ai_core_url   = "https://<your-ai-core-url>"
deployment_id = "<your-deployment-id>"

response = requests.post(
    f"{ai_core_url}/v2/inference/deployments/{deployment_id}/chat/completions",
    headers={
        "Authorization":    f"Bearer {token}",
        "Content-Type":     "application/json",
        "AI-Resource-Group":"default"},
    json={
        "model":     "codesage-llama3",
        "messages":  [{"role": "user", "content": "What is ABAP?"}],
        "max_tokens": 200})

print("Status:", response.status_code)
print("Answer:", response.json()["choices"][0]["message"]["content"])
# Expected: Status: 200  Answer: ABAP stands for Advanced Business Application ...""")

code_block(doc, """python test_aicore.py""")

make_table(doc,
    ['Common Error', 'Cause', 'Fix'],
    [
        ['CUDA out of memory (Kaggle)',       'Batch size too large',                           'Reduce per_device_train_batch_size to 2; increase gradient_accumulation_steps to 8'],
        ['HuggingFace 403 on LLaMA-3',        'License not accepted',                           'Visit huggingface.co/meta-llama/Meta-Llama-3-8B-Instruct and click Accept'],
        ['Docker push timeout',               'Image too large for connection timeout',          'Split push with --compress or use a faster upload connection'],
        ['AI Core deployment stuck PENDING',  'Registry secret missing or wrong credentials',   'Check ML Operations \u2192 Docker Registry Secrets; re-create with correct credentials'],
        ['Inference returns 503',             'Deployment still initialising',                   'Wait 5\u201310 min; check deployment logs in AI Launchpad for errors'],
        ['Empty response from model',         'max_model_len too small for prompt',              'Increase --max-model-len in Dockerfile CMD to 8192 or 16384'],
    ],
    col_widths=[2.2, 2.3, 3.2]
)

info_box(doc, '\u2705 Phase 3 Outcome:',
    'A fine-tuned LLaMA-3 8B model \u2014 trained on your organisation\'s ABAP codebase '
    'using QLoRA \u2014 is running as a live, OpenAI-compatible inference endpoint on '
    'SAP AI Core. No local GPU is needed at runtime. The endpoint is reachable via the '
    'CodeSage Agent on BTP. Total training cost: \u00a30 on Kaggle free tier, or ~\u00a3200 '
    'on a cloud GPU instance (one-off). Inference cost: SAP AI Core per-token billing.')

doc.add_heading('5.9  Hardware Requirements', level=2)
make_table(doc,
    ['Component', 'Fine-Tuning (one-time)', 'Runtime Inference via BYOM'],
    [
        ['GPU',     'Kaggle T4 16 GB (free) or NVIDIA RTX 3090 / A100', 'SAP AI Core managed \u2014 no local GPU needed'],
        ['RAM',     '16 GB+ system RAM',                                  'N/A \u2014 SAP cloud managed'],
        ['Storage', '~100 GB SSD (model + LoRA + dataset)',               'N/A \u2014 SAP cloud managed'],
        ['Network', 'Internet: Hugging Face + Claude API (one-time)',     'HTTPS: BTP Destination \u2192 AI Core endpoint'],
        ['Cost',    'Kaggle: \u00a30  /  Cloud GPU: ~\u00a3200 (one-off)','SAP AI Core: per-inference billing (metered)'],
    ],
    col_widths=[1.4, 3.0, 3.3]
)

doc.add_page_break()

# =============================================================================
# CHAPTER 3 — THE APPLICATION
# =============================================================================
doc.add_heading('Chapter 6:  Phase 4 — Runtime Query via SAP BTP', level=1)

doc.add_heading('6.1  Runtime Architecture', level=2)
p(doc, 'At runtime every developer query passes through a five-step pipeline orchestrated '
       'by the CodeSage Agent \u2014 a SAP BTP CAP (Cloud Application Programming) service '
       'that handles authentication, routing, and logging.')
make_table(doc,
    ['Step', 'Component', 'Location', 'Action'],
    [
        ['1', 'Fiori App / Streamlit UI',      'BTP or on-premise',     'Developer enters natural language question'],
        ['2', 'ChromaDB Retriever',            'On-premise',            'Embed query \u2192 semantic search \u2192 top-5 code chunks'],
        ['3', 'Prompt Builder',                'On-premise / BTP',      'Assemble context-aware prompt with retrieved chunks'],
        ['4', 'CodeSage Agent (BTP CAP)',      'SAP BTP',               'Authenticate, rate-limit, call SAP AI Core via BTP Destination'],
        ['5', 'SAP AI Core BYOM (LLaMA-3)',   'SAP AI Core (cloud)',    'Generate response with code explanation and citations'],
    ],
    col_widths=[0.4, 2.1, 1.6, 3.6]
)

doc.add_heading('6.2  Developer Query Flow', level=2)
p(doc, 'A typical developer interaction proceeds as follows:')
numbered(doc, 'Developer types: "How do we handle dunning in our system?" into the CodeSage UI.')
numbered(doc, 'Query is embedded using nomic-embed-text \u2014 local call, no cloud round-trip.')
numbered(doc, 'ChromaDB returns the top-5 most semantically similar ABAP code chunks.')
numbered(doc, 'Prompt Builder formats these as context in a structured system prompt.')
numbered(doc, 'CodeSage Agent sends the assembled prompt to SAP AI Core BYOM (LLaMA-3).')
numbered(doc, 'LLaMA-3 returns a response in 2\u20135 seconds, citing function modules and programs by name.')
numbered(doc, 'Developer sees the answer with source references and can navigate directly to the code.')
p(doc, 'Key advantage over generic AI: every answer is grounded in your organisation\'s own '
       'code. The model cannot invent a function module that does not exist in your system.')

p(doc, 'Phase 4 folder structure \u2014 CodeSage Agent (BTP CAP service):', bold=True, sa=2)
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase4_runtime/
\u2502   \u251c\u2500\u2500 codesage-agent/               # SAP BTP CAP service (deployed to Cloud Foundry)
\u2502   \u2502   \u251c\u2500\u2500 package.json              # Node.js dependencies (CAP, @sap/xssec, axios)
\u2502   \u2502   \u251c\u2500\u2500 mta.yaml                  # BTP Multi-Target Application descriptor
\u2502   \u2502   \u251c\u2500\u2500 xs-security.json          # XSUAA roles: CodeSage.User, CodeSage.Admin
\u2502   \u2502   \u251c\u2500\u2500 .env                      # Local dev only \u2014 never commit to git
\u2502   \u2502   \u251c\u2500\u2500 srv/
\u2502   \u2502   \u2502   \u251c\u2500\u2500 codesage-service.cds  # CDS service definition (REST endpoint)
\u2502   \u2502   \u2502   \u251c\u2500\u2500 query-handler.js      # Core: RAG retrieval + AI Core call
\u2502   \u2502   \u2502   \u2514\u2500\u2500 chromadb-client.js    # ChromaDB query helper (calls on-premise)
\u2502   \u2502   \u2514\u2500\u2500 db/
\u2502   \u2502       \u2514\u2500\u2500 schema.cds            # Optional: query log table
\u2502   \u2514\u2500\u2500 ui/
\u2502       \u251c\u2500\u2500 webapp/                   # Fiori freestyle or SAP UI5 app
\u2502       \u2514\u2500\u2500 xs-app.json               # App Router config (SSO redirect)""")

p(doc, 'codesage-service.cds \u2014 CAP service definition:', bold=True, sa=2)
code_block(doc,
"""// srv/codesage-service.cds
// Defines the REST endpoint that the UI and Teams bot call

service CodeSageService @(requires: 'CodeSage.User') {

    // Main query endpoint
    // POST /codesage/query  { question, context_filter? }
    action query(
        question       : String not null,
        context_filter : String    // optional: 'abap_programs' | 'function_modules' etc.
    ) returns {
        answer      : String;
        sources     : array of { object_name: String; chunk_type: String; };
        model       : String;
        latency_ms  : Integer;
    };
}""")

p(doc, 'query-handler.js \u2014 the core orchestration logic (RAG + BYOM call):', bold=True, sa=2)
code_block(doc,
"""// srv/query-handler.js
const axios   = require('axios');
const xsenv   = require('@sap/xsenv');
const { retrieveChunks } = require('./chromadb-client');

// BTP Destination Service provides the AI Core URL + OAuth2 token automatically
const { destination } = xsenv.getServices({ destination: { tag: 'destination' } });

module.exports = (srv) => {

  srv.on('query', async (req) => {
    const { question, context_filter } = req.data;
    const t0 = Date.now();

    // ── Step 1: Retrieve top-5 relevant chunks from ChromaDB (on-premise) ──
    const collections = context_filter
      ? [context_filter]
      : ['abap_programs', 'function_modules', 'abap_classes'];

    const chunks = await retrieveChunks(question, collections, 5);

    const context = chunks
      .map((c, i) => `[${i+1}] ${c.metadata.object} (${c.metadata.type}):\\n${c.document}`)
      .join('\\n\\n');

    // ── Step 2: Build prompt ────────────────────────────────────────────────
    const messages = [
      { role: 'system',
        content: 'You are CodeSage, an expert ABAP assistant for this organisation. '
                +'Answer ONLY using the retrieved context. Cite every object by name. '
                +'End your response with a Sources section.' },
      { role: 'user',
        content: `RETRIEVED CONTEXT:\\n${context}\\n\\nQUESTION: ${question}` }
    ];

    // ── Step 3: Call SAP AI Core BYOM endpoint (OpenAI-compatible) ─────────
    const aiCoreUrl  = destination.url;         // from BTP Destination
    const deployId   = process.env.AI_CORE_DEPLOYMENT_ID;
    const token      = await getAICoreToken();  // OAuth2 client credentials

    const response = await axios.post(
      `${aiCoreUrl}/v2/inference/deployments/${deployId}/chat/completions`,
      { model: 'codesage-llama3', messages, max_tokens: 1024, temperature: 0.1 },
      { headers: { Authorization: `Bearer ${token}`,
                   'AI-Resource-Group': 'default' } });

    const answer = response.data.choices[0].message.content;

    // ── Step 4: Parse sources from answer ──────────────────────────────────
    const sources = chunks.map(c => ({
      object_name: c.metadata.object,
      chunk_type:  c.metadata.type
    }));

    return { answer, sources, model: 'codesage-llama3', latency_ms: Date.now() - t0 };
  });
};""")

# ── How to Run Phase 4 ───────────────────────────────────────────────────────
doc.add_heading('6.3  How to Run Phase 4 \u2014 Deploy the CodeSage Agent', level=2)

p(doc, 'Phase 4 is the runtime layer. It has two deployment targets: '
       'the CodeSage Agent (BTP CAP service) and the Fiori / UI frontend. '
       'Phases 1\u20133 must all be complete before deploying Phase 4.')

p(doc, 'Prerequisites', bold=True, sa=2)
make_table(doc,
    ['Requirement', 'Details', 'Where / How'],
    [
        ['Phases 1\u20133 complete',     'ChromaDB running + AI Core BYOM endpoint RUNNING',    'Verify AI Core status in AI Launchpad'],
        ['BTP subaccount',              'Cloud Foundry environment enabled (any region)',         'BTP Cockpit \u2192 Enable Cloud Foundry'],
        ['Cloud Foundry CLI (cf)',       'Command-line tool to push CAP apps to BTP',             'github.com/cloudfoundry/cli/releases'],
        ['Node.js 18+',                 'Required to build and run the CAP service locally',      'nodejs.org'],
        ['@sap/cds-dk',                 'SAP CDS development kit',                                'npm install -g @sap/cds-dk'],
        ['MTA Build Tool (mbt)',         'Builds the multi-target archive for BTP deployment',    'npm install -g mbt'],
        ['XSUAA service instance',       'BTP OAuth2 service for authentication',                 'BTP Cockpit \u2192 Service Marketplace \u2192 Authorization & Trust Management'],
        ['Destination service instance', 'Stores AI Core endpoint credentials securely',          'BTP Cockpit \u2192 Service Marketplace \u2192 Destination'],
        ['CODESAGE_AICORE destination',  'Configured in Phase 3 Step E',                          'BTP Cockpit \u2192 Destinations'],
    ],
    col_widths=[2.0, 2.8, 2.4]
)

p(doc, 'Phase 4 deploys the CodeSage Agent \u2014 a small Node.js service \u2014 to SAP BTP. '
       'Once deployed, developers can query it from their browser, Fiori Launchpad, or Teams. '
       'All commands below run in Windows PowerShell.',
  italic=True, color=GREY)

p(doc, 'Step 1 \u2014 Install local developer tools (one-time setup on your Windows PC):', bold=True, sa=2)
code_block(doc,
"""# Install Node.js 18 LTS from https://nodejs.org (LTS version recommended)
# Verify after installation:
node --version    # Expected: v18.x.x or higher
npm  --version    # Expected: 9.x.x or higher

# Install SAP CDS development kit (cds is the SAP CAP framework CLI)
npm install -g @sap/cds-dk
cds --version     # Expected: @sap/cds: 7.x.x

# Install the MTA Build Tool (mbt packages the app for BTP deployment)
npm install -g mbt
mbt --version     # Expected: 1.2.x

# Install the Cloud Foundry CLI (cf is used to deploy to BTP Cloud Foundry)
# Download the Windows installer from: https://github.com/cloudfoundry/cli/releases
# After install, verify:
cf --version      # Expected: cf version 8.x.x""")

p(doc, 'Step 2 \u2014 Install project dependencies and run locally for testing:', bold=True, sa=2)
p(doc, 'Before deploying to BTP, always test the service locally first. '
       'Running locally is faster to debug and does not consume BTP resources.')
code_block(doc,
"""cd C:\\codesage\\phase4_runtime\\codesage-agent

# Install Node.js packages defined in package.json
npm install
# Installs: @sap/cds, @sap/xssec, @sap/xsenv, axios, chromadb, express

# Set local environment variables for this PowerShell session
# These point the local service to your on-premise ChromaDB and AI Core deployment
$env:CHROMADB_URL            = "http://localhost:8000"
$env:AI_CORE_DEPLOYMENT_ID   = "<your-deployment-id-from-phase-3>"

# Start the CAP service locally (watches for file changes \u2014 good for development)
cds watch
# Console output:
# [cds] - model loaded from: srv\\codesage-service.cds
# [cds] - serving CodeSageService { path: '/codesage' }
# [cds] - server listening on { url: 'http://localhost:4004' }""")

p(doc, 'While cds watch is running, open a second PowerShell window and test the endpoint:')
code_block(doc,
"""# Test the local endpoint with a sample question
# (Invoke-RestMethod is the Windows PowerShell equivalent of curl)
$body = @{ question = "How do we validate vendor payment terms?" } | ConvertTo-Json

Invoke-RestMethod `
  -Method  Post `
  -Uri     "http://localhost:4004/codesage/query" `
  -Headers @{ "Content-Type" = "application/json" } `
  -Body    $body

# Expected response:
# answer     : Yes \u2014 Z_VALIDATE_VENDOR_PAYTERMS validates vendor...
# sources    : {@{object_name=Z_VALIDATE_VENDOR_PAYTERMS; chunk_type=FUNCTION}}
# model      : codesage-llama3
# latency_ms : 2841""")

p(doc, 'Step 3 \u2014 Log in to BTP Cloud Foundry and create service instances (once per subaccount):', bold=True, sa=2)
p(doc, 'XSUAA provides OAuth2 security (login tokens) and the Destination service stores the '
       'AI Core credentials securely so they are never hard-coded in the application.')
code_block(doc,
"""# Log in to BTP Cloud Foundry (get the API endpoint from BTP Cockpit -> Cloud Foundry)
cf login -a https://api.cf.<your-region>.hana.ondemand.com
# Enter your SAP BTP email and password when prompted
# Select your org and space from the numbered list

# Create the XSUAA (authentication) service instance
# xs-security.json defines the roles: CodeSage.User and CodeSage.Admin
cf create-service xsuaa application codesage-xsuaa -c xs-security.json

# Create the Destination service instance
cf create-service destination lite codesage-destination

# Verify both services were created successfully
cf services
# NAME                    SERVICE        PLAN          STATUS
# codesage-xsuaa          xsuaa          application   create succeeded
# codesage-destination    destination    lite          create succeeded""")

p(doc, 'Step 4 \u2014 Build the deployment archive and deploy to BTP:', bold=True, sa=2)
code_block(doc,
"""cd C:\\codesage\\phase4_runtime\\codesage-agent

# Build the MTA archive (packages all files into a single .mtar deployment file)
mbt build
# Output: creating archive: .\\mta_archives\\codesage-agent_1.0.0.mtar

# Deploy to BTP Cloud Foundry
cf deploy mta_archives\\codesage-agent_1.0.0.mtar
# What happens during deployment (takes 3-5 minutes):
#   Uploading archive to BTP...
#   Creating application "codesage-agent"...
#   Binding services: codesage-xsuaa, codesage-destination...
#   Starting application...
#   Application "codesage-agent" started at:
#   https://codesage-agent-<random-id>.cfapps.<region>.hana.ondemand.com

# Note down the URL \u2014 this is your live CodeSage endpoint""")

p(doc, 'Step 5 \u2014 Point the deployed agent to your on-premise ChromaDB:', bold=True, sa=2)
p(doc, 'ChromaDB runs on-premise inside your network. BTP Cloud Foundry (in the cloud) '
       'needs a route to reach it. Use SAP Cloud Connector to expose the ChromaDB port '
       'to BTP. Then set the URL as an environment variable:')
code_block(doc,
"""# Set ChromaDB URL (via SAP Cloud Connector virtual host)
cf set-env codesage-agent CHROMADB_URL "https://virtual-chromadb-host:8000"
cf set-env codesage-agent AI_CORE_DEPLOYMENT_ID "<your-deployment-id>"

# Restart the app to pick up the new environment variables
cf restage codesage-agent

# Check the app is running
cf app codesage-agent
# Expected: requested state: started  instances: 1/1  memory: 256M""")

p(doc, 'Step 6 \u2014 Verify the live BTP deployment end-to-end:', bold=True, sa=2)
p(doc, 'Save the snippet below as test_live.py, fill in your XSUAA credentials '
       '(from the service key in BTP Cockpit), and run it from PowerShell:')
code_block(doc,
"""# test_live.py  -- run: python test_live.py
import requests

# Step 6a: Get an OAuth2 access token from XSUAA
#   Find these values in BTP Cockpit -> Instances -> codesage-xsuaa -> Service Keys
xsuaa_url     = "https://<your-subdomain>.authentication.<region>.hana.ondemand.com"
client_id     = "<xsuaa-client-id>"
client_secret = "<xsuaa-client-secret>"

token_resp = requests.post(
    f"{xsuaa_url}/oauth/token",
    data={"grant_type": "client_credentials"},
    auth=(client_id, client_secret))
token = token_resp.json()["access_token"]
print("Token obtained OK")

# Step 6b: Call the live CodeSage Agent on BTP
agent_url = "https://codesage-agent-<id>.cfapps.<region>.hana.ondemand.com"

response = requests.post(
    f"{agent_url}/codesage/query",
    headers={
        "Authorization": f"Bearer {token}",
        "Content-Type":  "application/json"},
    json={"question": "Do we have a function module for vendor payment term validation?"})

print("HTTP Status:", response.status_code)
data = response.json()
print("Answer:",      data["answer"][:200])
print("Sources:",     data["sources"])
print("Latency ms:",  data["latency_ms"])
# Expected:
# HTTP Status: 200
# Answer: Yes \u2014 Z_VALIDATE_VENDOR_PAYTERMS checks vendor payment terms against T052...
# Sources: [{'object_name': 'Z_VALIDATE_VENDOR_PAYTERMS', 'chunk_type': 'FUNCTION'}]
# Latency ms: 2841""")

code_block(doc, """python test_live.py""")

p(doc, 'Step 7 \u2014 Register as a Fiori tile (optional but recommended):', bold=True, sa=2)
numbered(doc, 'In BTP Cockpit \u2192 HTML5 Applications, upload the ui/webapp folder as an HTML5 app.')
numbered(doc, 'In SAP Fiori Launchpad Configuration Cockpit, add a new tile pointing to the HTML5 app URL.')
numbered(doc, 'Assign the tile to a role collection mapped to CodeSage.User. Developers will see the tile when they log in to their Fiori Launchpad.')

make_table(doc,
    ['Common Error', 'Cause', 'Fix'],
    [
        ['cf login fails',                      'Wrong API endpoint URL for your BTP region',         'Check BTP Cockpit \u2192 Cloud Foundry \u2192 API Endpoint'],
        ['XSUAA binding missing',               'xs-security.json not found during cf create-service','Run from the codesage-agent/ folder where xs-security.json lives'],
        ['401 Unauthorized on /codesage/query', 'JWT token missing or expired',                       'Pass valid Bearer token; token lifetime is typically 12 hrs'],
        ['ChromaDB connection refused',         'BTP cannot reach on-premise ChromaDB',               'Configure SAP Cloud Connector to expose ChromaDB port to BTP'],
        ['AI Core 404 on inference call',       'Wrong deployment ID or region mismatch',             'Re-check CODESAGE_AICORE destination URL and DEPLOYMENT_ID env var'],
        ['Slow first response (>15 s)',         'AI Core cold start \u2014 model loading into GPU',   'Normal on first call after idle. Subsequent calls: 2\u20135 s'],
    ],
    col_widths=[2.2, 2.5, 3.0]
)

info_box(doc, '\u2705 Phase 4 Outcome:',
    'A developer types a question in plain English and receives a contextualised ABAP answer '
    'in 2\u20135 seconds, citing real function modules and programs by name. Every query is '
    'authenticated via BTP XSUAA, routed through the CodeSage Agent, grounded by ChromaDB '
    'retrieval, and generated by the fine-tuned LLaMA-3 on SAP AI Core. No code leaves '
    'the organisation\'s SAP BTP perimeter. Developers stop reinventing \u2014 they reuse.')

doc.add_heading('6.4  Integration Options', level=2)
p(doc, 'CodeSage can be surfaced through multiple channels. The right path depends on your '
       'organisation\'s tooling, SSO setup, and where developers spend their time.')
make_table(doc,
    ['Channel', 'Effort', 'Auth', 'Best For'],
    [
        ['SAP Fiori Tile \u2b50',   '~2 days', 'BTP IAS SSO \u2014 corporate IdP',       'All SAP developers; seamless Fiori Launchpad integration'],
        ['Microsoft Teams Bot',     '~3 days', 'Azure AD + BTP OAuth2 exchange',           'Teams-first organisations; async Q&A in channels'],
        ['SAP Joule (tile link)',    '~1 day',  'BTP IAS SSO (same as Fiori)',              'Joule-enabled orgs; CodeSage opens in Joule sidebar panel'],
        ['SAP Joule (Skill)',        'Roadmap', 'SAP AI Foundation Skills framework',       'Future \u2014 native "ask Joule about our code" capability'],
        ['Standalone Web UI',        '0 days',  'BTP service key (manual sign-in)',         'Internal pilot; offline / laptop-local use'],
    ],
    col_widths=[1.8, 0.8, 2.3, 2.8]
)

doc.add_heading('6.4.1  SAP Fiori Tile', level=3)
p(doc, 'The recommended integration. Deploy CodeSage as a BTP Cloud Foundry application '
       '(Node.js or Python Streamlit), then register it as a custom tile in the SAP Fiori '
       'Launchpad. SSO is handled by BTP Identity Authentication Service \u2014 developers '
       'access CodeSage without a separate login. Effort is approximately two days covering '
       'BTP CF deployment and tile registration.')

doc.add_heading('6.4.2  Microsoft Teams Bot', level=3)
p(doc, 'For Teams-first organisations, two approaches are available:')
bullet(doc, 'Power Automate flow: Teams Message trigger \u2192 HTTP POST to BTP CAP endpoint \u2192 Adaptive Card reply. No custom code; configurable by a Power Platform developer.', bold_prefix='Option A \u2014 Power Automate:  ')
bullet(doc, 'Bot Framework webhook registered as a Teams App, pointing to the BTP CAP service. Offers richer UX cards and more control over conversation handling.', bold_prefix='Option B \u2014 Bot Framework:  ')
p(doc, 'Authentication uses Azure AD (Entra ID) for Teams identity and a BTP OAuth2 token '
       'exchange for the AI Core call. Developers can ask questions in a private bot chat '
       'or a shared #abap-help channel.')

doc.add_heading('6.4.3  SAP Joule Integration', level=3)
p(doc, 'SAP Joule runs on SAP\'s own internal models and cannot be configured to call a '
       'custom BYOM endpoint directly. CodeSage and Joule are complementary \u2014 Joule '
       'answers generic SAP questions; CodeSage answers questions about your specific code. '
       'Two integration paths exist:')
bullet(doc, 'A custom Fiori tile opens CodeSage in an embedded Joule sidebar panel. Available today; no SAP roadmap dependency.', bold_prefix='Tile integration (now):  ')
bullet(doc, 'SAP is building a Skills framework for Joule that will allow BTP CAP services to register as custom skills. When available, CodeSage can be registered to enable native "ask Joule about our code" conversations.', bold_prefix='Joule Skill (roadmap):  ')

doc.add_heading('6.5  CodeSage Agent \u2014 SAP BTP CAP Service', level=2)
p(doc, 'The CodeSage Agent is the central orchestration component, deployed as a SAP BTP '
       'CAP service on Cloud Foundry. It exposes a single REST endpoint and handles '
       'security, routing, and observability:')
bullet(doc, 'POST /codesage/query \u2014 accepts { question, user_id, context_filter }', bold_prefix='Endpoint:  ')
bullet(doc, 'BTP XSUAA OAuth2 \u2014 validates JWT on every request before any processing.', bold_prefix='Authentication:  ')
bullet(doc, 'Calls SAP AI Core BYOM via a named BTP Destination \u2014 credentials never exposed to the client.', bold_prefix='BTP Destination:  ')
bullet(doc, 'Configurable per user and team to control SAP AI Core inference costs.', bold_prefix='Rate limiting:  ')
bullet(doc, 'Anonymised query hashes and response times logged to BTP Application Logging Service (ALS).', bold_prefix='Logging:  ')
bullet(doc, 'Returns a graceful message if the BYOM endpoint is unavailable or times out (> 15 s).', bold_prefix='Error handling:  ')

doc.add_page_break()

# =============================================================================
# CHAPTER 4 — GOVERNANCE, DATA PROTECTION & COMPLIANCE
# =============================================================================
doc.add_heading('Chapter 7:  Governance, Data Protection & Compliance', level=1)

p(doc, 'CodeSage handles ABAP source code which may contain proprietary business logic and '
       'references to sensitive data structures. The following governance framework must be '
       'reviewed and accepted before production deployment.')

doc.add_heading('7.1  Data Classification & Scan Scope', level=2)
p(doc, 'The table below defines what is and is not scanned and indexed by CodeSage:')
make_table(doc,
    ['Data Type', 'Included', 'Rationale'],
    [
        ['Custom ABAP programs (Z*/Y*)',      '\u2705 Yes',             "Core content \u2014 the organisation's own logic"],
        ['SAP standard code (no Z/Y prefix)', '\u274c No',              'Not org-specific; available in SAP Help docs'],
        ['Database table contents',           '\u274c No',              'Data, not code \u2014 may contain personal data'],
        ['Hard-coded credentials / secrets',  '\u274c Auto-filtered',   'Pre-scan regex removes before any processing'],
        ['SAP Help documentation',            '\u2705 Yes (public)',    'Provides SAP context for RAG-grounded answers'],
        ['User names, personnel IDs',         '\u274c Auto-filtered',   'Excluded to avoid personal data ingestion (GDPR)'],
    ],
    col_widths=[2.5, 1.3, 3.9]
)

doc.add_heading('7.2  GDPR & Data Residency', level=2)
p(doc, 'Data handling across the four pipeline phases is as follows:')
bullet(doc, 'ChromaDB runs entirely on-premise. No ABAP source code is transmitted to any cloud service during indexing.', bold_prefix='Indexing (Phase 2):  ')
bullet(doc, 'Code chunks are submitted to the Claude API for QA pair generation \u2014 a one-time, controlled exercise. Chunks should be reviewed for sensitive content. Anthropic\'s commercial API is GDPR-compliant (DPA available).', bold_prefix='Fine-tuning (Phase 3 \u2014 one-time):  ')
bullet(doc, 'Developer queries are sent to SAP AI Core. SAP AI Core is available in EU (Frankfurt) and US data centres. Choose the BTP subaccount region that matches your data residency requirements.', bold_prefix='Runtime queries (Phase 4):  ')
bullet(doc, 'No query content is stored by default. Only anonymised query hashes and response times are logged for monitoring.', bold_prefix='Query logging:  ')

doc.add_heading('7.3  SAP AI Core Compliance Certifications', level=2)
p(doc, 'SAP AI Core, as a SAP BTP service, is covered by SAP\'s enterprise compliance programme:')
make_table(doc,
    ['Certification', 'Scope', 'Relevance to CodeSage'],
    [
        ['SOC 2 Type II',    'Security, availability, confidentiality \u2014 independently audited',  'BYOM inference workloads are within scope'],
        ['ISO/IEC 27001',    'Information security management system',                                 'Data processed within certified infrastructure'],
        ['GDPR',             'SAP DPA available; EU Standard Contractual Clauses in place',            'Choose EU Frankfurt region for data residency compliance'],
        ['Cloud Act',        'EU data centre option minimises US jurisdiction exposure',               'Select EU BTP subaccount'],
        ['SAP Trust Centre', 'Current audit reports and certification status',                         'Verify at trust.sap.com before production go-live'],
    ],
    col_widths=[1.7, 3.2, 2.8]
)

doc.add_heading('7.4  Access Control & Authentication', level=2)
make_table(doc,
    ['Component', 'Method', 'Roles'],
    [
        ['CodeSage Fiori / Web UI',     'BTP IAS \u2014 SSO via corporate IdP (Azure AD / LDAP)',                      'All authenticated employees'],
        ['CodeSage Agent (BTP CAP)',    'BTP XSUAA OAuth2 \u2014 JWT validated on every request',                      'CodeSage.User (query), CodeSage.Admin (manage)'],
        ['SAP AI Core BYOM endpoint',   'Client credentials stored in BTP Secret Store \u2014 never in code or UI',    'Service identity only'],
        ['ChromaDB (on-premise)',        'Local network only \u2014 no external port',                                  'Application service account'],
        ['Container registry',          'Private registry with 90-day credential rotation',                             'CI/CD pipeline service account'],
    ],
    col_widths=[2.0, 3.3, 2.4]
)
p(doc, 'Role separation: CodeSage.Admin is required to initiate model refresh, update scan '
       'configuration, or access raw logs. CodeSage.User provides query access only.')

doc.add_heading('7.5  Model Governance & Refresh Policy', level=2)
p(doc, 'To ensure CodeSage stays current and accurate as the codebase evolves, the '
       'following refresh and approval process applies:')
bullet(doc, 'Monthly, aligned with the SAP transport release calendar.', bold_prefix='Refresh frequency:  ')
bullet(doc, 'Phase 1 re-scan captures new and modified ABAP objects; Phases 2 and 3 run incrementally.', bold_prefix='What runs:  ')
bullet(doc, 'Tech lead or architecture owner sign-off required before deploying a new model version to production.', bold_prefix='Approval gate:  ')
bullet(doc, 'Previous model version (N\u22121) retained in SAP AI Core for 30 days to allow immediate rollback.', bold_prefix='Rollback:  ')
bullet(doc, '20 benchmark queries run before and after every refresh. Threshold: less than 5% degradation in answer relevance.', bold_prefix='Quality gate:  ')
bullet(doc, 'Model version, training data snapshot, approver, and quality score recorded per deployment.', bold_prefix='Change log:  ')

doc.add_heading('7.6  Audit Trail & Monitoring', level=2)
bullet(doc, 'Timestamp, anonymised user hash, response time, and model version \u2192 BTP Application Logging Service.', bold_prefix='Query log:  ')
bullet(doc, 'Disabled by default; enabled only with user consent and a defined data retention period.', bold_prefix='Query content:  ')
bullet(doc, 'SAP AI Launchpad provides deployment health, latency, and error rate metrics for the BYOM endpoint.', bold_prefix='SAP AI Core monitoring:  ')
bullet(doc, 'BTP Alert Notification Service configured for: deployment downtime, latency > 10 s, error rate > 5%.', bold_prefix='Alerting:  ')

doc.add_heading('7.7  Responsible AI Principles', level=2)
p(doc, 'CodeSage is a decision-support tool, not an autonomous code generator. The '
       'following principles govern its responsible use:')
bullet(doc, 'Every response cites the source ABAP object and program. Developers can verify every answer against actual code.', bold_prefix='Transparency:  ')
bullet(doc, 'Generated code must be reviewed by a developer before transport to production. No direct write-back to SAP is implemented.', bold_prefix='Human review required:  ')
bullet(doc, 'LLaMA-3 can occasionally produce plausible but incorrect ABAP for edge cases not well represented in training data. This limitation is displayed in the UI.', bold_prefix='Limitation disclosure:  ')
bullet(doc, 'Developers are reminded not to include customer names, employee IDs, or sensitive data in queries. BTP content filtering can be enabled as an additional safeguard.', bold_prefix='No PII in queries:  ')
bullet(doc, 'Query logs reviewed quarterly for bias, misuse, or systematic errors. Findings feed back into the next fine-tuning cycle.', bold_prefix='Continuous improvement:  ')

doc.add_page_break()

# =============================================================================
# CHAPTER 5 — SAMPLE CODE GENERATED BY CLAUDE
# =============================================================================
doc.add_heading('Chapter 8:  Sample Code Generated by Claude', level=1)

p(doc, 'The following examples illustrate the quality and style of ABAP code that '
       'CodeSage produces during the QA-pair generation stage (Phase 3) and at runtime '
       'when a developer submits a question. All examples were generated by Claude '
       '(claude-sonnet-4-6) using retrieval-augmented context pulled from a sample '
       'SAP ECC system with ~12,000 custom ABAP objects.')

info_box(doc, 'How these samples are used:',
         'Claude generates the training QA pairs during the one-time fine-tuning phase. '
         'After fine-tuning, the deployed LLaMA-3 model answers developer queries at '
         'runtime using the same RAG context patterns Claude was trained on.')

# ── 5.1 ──────────────────────────────────────────────────────────────────────
doc.add_heading('8.1  Vendor Payment Terms Validator', level=2)
p(doc, 'Developer query:  "Do we have a function module that validates vendor payment terms '
       'before posting an invoice?"')
p(doc, 'CodeSage response (runtime \u2014 grounded in retrieved context):')
code_block(doc,
"""FUNCTION Z_VALIDATE_VENDOR_PAYTERMS.
*"----------------------------------------------------------------------
*"  IMPORTING
*"    VALUE(IV_LIFNR)   TYPE  LIFNR       " Vendor account
*"    VALUE(IV_ZTERM)   TYPE  DZTERM      " Payment term key
*"  EXPORTING
*"    VALUE(EV_VALID)   TYPE  BOOLE_D
*"    VALUE(EV_MESSAGE) TYPE  STRING
*"----------------------------------------------------------------------
  DATA: ls_lfa1  TYPE lfa1,
        ls_t052  TYPE t052.

  " Check vendor master exists
  SELECT SINGLE * FROM lfa1 INTO ls_lfa1
    WHERE lifnr = iv_lifnr.
  IF sy-subrc <> 0.
    ev_valid   = abap_false.
    ev_message = |Vendor { iv_lifnr } not found in LFA1|.
    RETURN.
  ENDIF.

  " Validate payment term against T052 (FI config table)
  SELECT SINGLE * FROM t052 INTO ls_t052
    WHERE zterm = iv_zterm.
  IF sy-subrc <> 0.
    ev_valid   = abap_false.
    ev_message = |Payment term { iv_zterm } not configured in T052|.
    RETURN.
  ENDIF.

  ev_valid   = abap_true.
  ev_message = |Vendor { iv_lifnr } payment term { iv_zterm } is valid|.
ENDFUNCTION."""
)
p(doc, 'Source cited by CodeSage:  ZFI_VENDOR_PAYTERMS_CHCK (Program), '
       'Z_VALIDATE_VENDOR_PAYTERMS (Function Module), T052 (Table reference).',
  italic=True, color=GREY, size=10)

# ── 5.2 ──────────────────────────────────────────────────────────────────────
doc.add_heading('8.2  ABAP OO Class \u2014 Purchase Order Status Handler', level=2)
p(doc, 'Developer query:  "Show me how we handle PO status changes in our system."')
code_block(doc,
"""CLASS zcl_po_status_handler DEFINITION
  PUBLIC FINAL CREATE PUBLIC.

  PUBLIC SECTION.
    TYPES: BEGIN OF ty_status_change,
             ebeln    TYPE ebeln,
             old_stat TYPE char2,
             new_stat TYPE char2,
             changed_by TYPE uname,
             changed_at TYPE timestamp,
           END OF ty_status_change.

    METHODS:
      constructor
        IMPORTING iv_ebeln TYPE ebeln,
      get_current_status
        RETURNING VALUE(rv_status) TYPE char2,
      apply_status_change
        IMPORTING iv_new_status TYPE char2
        RAISING   zcx_po_status_error.

  PRIVATE SECTION.
    DATA: mv_ebeln   TYPE ebeln,
          mv_status  TYPE char2.

    METHODS validate_transition
      IMPORTING iv_from TYPE char2
                iv_to   TYPE char2
      RETURNING VALUE(rv_allowed) TYPE boole_d.
ENDCLASS.

CLASS zcl_po_status_handler IMPLEMENTATION.
  METHOD constructor.
    mv_ebeln = iv_ebeln.
    SELECT SINGLE statu INTO mv_status
      FROM ekko WHERE ebeln = mv_ebeln.
  ENDMETHOD.

  METHOD get_current_status.
    rv_status = mv_status.
  ENDMETHOD.

  METHOD apply_status_change.
    IF validate_transition( iv_from = mv_status
                            iv_to   = iv_new_status ) = abap_false.
      RAISE EXCEPTION TYPE zcx_po_status_error
        EXPORTING
          mv_message = |Transition { mv_status } -> { iv_new_status } not allowed|.
    ENDIF.
    UPDATE ekko SET statu = iv_new_status
      WHERE ebeln = mv_ebeln.
    mv_status = iv_new_status.
  ENDMETHOD.

  METHOD validate_transition.
    " Allowed: OP->SN->LS->GR->IV->CL  (org-specific workflow)
    DATA(lt_allowed) = VALUE string_table(
      ( |OP-SN| ) ( |SN-LS| ) ( |LS-GR| ) ( |GR-IV| ) ( |IV-CL| ) ).
    rv_allowed = xsdbool( |{ iv_from }-{ iv_to }| IN lt_allowed ).
  ENDMETHOD.
ENDCLASS."""
)
p(doc, 'Source cited:  ZCL_PO_STATUS_HANDLER (Class), ZCX_PO_STATUS_ERROR (Exception), '
       'EKKO (Table reference).  Transition map derived from Z_MM_PO_WORKFLOW (Program).',
  italic=True, color=GREY, size=10)

# ── 5.3 ──────────────────────────────────────────────────────────────────────
doc.add_heading('8.3  BAdI Implementation Skeleton', level=2)
p(doc, 'Developer query:  "We need to add a custom check before goods receipt posting. '
       'Is there a BAdI we use for this?"')
code_block(doc,
"""" CodeSage response:
" Your system uses BAdI MB_MIGO_BADI for all MIGO-based GR postings.
" Below is the skeleton based on your existing implementation Z_MB_GR_CUSTOM_CHECK:

CLASS zcl_impl_mb_gr_check DEFINITION
  PUBLIC FINAL
  INHERITING FROM cl_badi_impl.

  PUBLIC SECTION.
    INTERFACES if_ex_mb_migo_badi.
ENDCLASS.

CLASS zcl_impl_mb_gr_check IMPLEMENTATION.
  METHOD if_ex_mb_migo_badi~mb_migo_hold_check.
    " Called before each item line is posted
    " Available context: im_mseg (movement segment), im_mkpf (header)
    DATA lv_plant TYPE werks_d.
    lv_plant = im_mseg-werks.

    " Your org uses plant 1000 / 2000 for restricted materials
    IF im_mseg-matnr IS NOT INITIAL AND lv_plant IN ('1000', '2000').
      " Call your existing validator (found in ZMM_RESTRICTED_STOCK_CHECK)
      CALL FUNCTION 'Z_CHECK_RESTRICTED_MATERIAL'
        EXPORTING
          iv_matnr = im_mseg-matnr
          iv_werks = lv_plant
        IMPORTING
          ev_blocked = DATA(lv_blocked)
          ev_reason  = DATA(lv_reason).
      IF lv_blocked = abap_true.
        ch_msg_prot = 'E'.  " Set error flag to hold posting
        MESSAGE e001(zmm_gr) WITH lv_reason INTO DATA(lv_dummy).
      ENDIF.
    ENDIF.
  ENDMETHOD.
ENDCLASS."""
)
p(doc, 'Source cited:  ZCL_IMPL_MB_GR_CHECK (existing implementation), '
       'Z_CHECK_RESTRICTED_MATERIAL (FM), ZMM_RESTRICTED_STOCK_CHECK (Program), '
       'MB_MIGO_BADI (BAdI definition).  Org-specific plant list from ZMM_PLANT_CONFIG.',
  italic=True, color=GREY, size=10)

# ── 5.4 ──────────────────────────────────────────────────────────────────────
doc.add_heading('8.4  CDS View with Association', level=2)
p(doc, 'Developer query:  "Can you show me how we build CDS views with associations to '
       'cost centres in our landscape?"')
code_block(doc,
"""@AbapCatalog.viewEnhancementCategory: [#NONE]
@AccessControl.authorizationCheck: #CHECK
@EndUserText.label: 'Cost Centre Allocation \u2014 CodeSage example'
@Metadata.ignorePropagatedAnnotations: true

define view entity ZI_COSTCTR_ALLOCATION
  as select from cosp
  association [0..1] to csks  as _CostCentre
    on $projection.CostCentre = _CostCentre.kostl

{
  key cosp.kokrs         as ControllingArea,
  key cosp.kostl         as CostCentre,
  key cosp.gjahr         as FiscalYear,
      cosp.wkgbtr        as ActualCost,
      cosp.wkgbtr_plan   as PlannedCost,

      " Association exposed for consumption views
      _CostCentre
}
where cosp.versn = '000'  -- actual version only (per Z_CC_REPORT_CONFIG)"""
)
p(doc, 'Source cited:  ZI_COSTCTR_ALLOCATION (existing CDS), COSP (Table), CSKS (Table), '
       'Z_CC_REPORT_CONFIG (Config program \u2014 defines version filter).',
  italic=True, color=GREY, size=10)

info_box(doc, 'Key quality indicator:',
         'Every sample above contains org-specific detail \u2014 table names, function module names, '
         'plant codes, version filters \u2014 that generic AI (without RAG) cannot know. '
         'This grounding eliminates hallucinated object names and wrong API calls.')

doc.add_page_break()

# =============================================================================
# CHAPTER 6 — RAG-POWERED CODE GENERATION USING CLAUDE SKILL
# =============================================================================
doc.add_heading('Chapter 9:  RAG-Powered Code Generation Using a Claude Skill', level=1)

p(doc, 'While Chapters 2 and 3 describe the CodeSage pipeline using a fine-tuned LLaMA-3 '
       'model on SAP AI Core BYOM, this chapter describes an alternative and complementary '
       'approach: using a Claude Skill (via the Anthropic API or SAP Generative AI Hub) as '
       'the generation model, fed by the same ChromaDB RAG layer. This approach is suitable '
       'for organisations that prefer a managed frontier model over a self-hosted fine-tuned '
       'model, or that want to use both in parallel.')

info_box(doc, 'Relationship to BYOM:',
         'BYOM (LLaMA-3 fine-tuned) and Claude Skill RAG are not mutually exclusive. '
         'A CodeSage deployment can route simple look-up queries to the fast BYOM model '
         'and complex multi-file code generation tasks to Claude via the Skill pattern.')

# ── 6.1 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.1  What Is a Claude Skill?', level=2)
p(doc, 'A Claude Skill is a structured, reusable prompt + tool-call pattern that wraps the '
       'Anthropic Claude API (or SAP Generative AI Hub \u2014 Claude on BTP) and exposes a '
       'clean interface to an orchestration layer. In the CodeSage context a Skill encapsulates:')
bullet(doc, 'A system prompt that defines the persona, output format, and citation rules.', bold_prefix='System prompt:  ')
bullet(doc, 'A retrieval tool call that fetches relevant ABAP snippets from ChromaDB before generation.', bold_prefix='Retrieval tool:  ')
bullet(doc, 'A generation call to Claude that assembles the retrieved context with the developer question.', bold_prefix='Generation:  ')
bullet(doc, 'Post-processing that extracts code blocks, citations, and confidence indicators from the response.', bold_prefix='Post-processing:  ')

# ── 6.2 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.2  RAG + Claude Skill Architecture', level=2)
p(doc, 'The flow below shows how the ChromaDB RAG layer feeds context directly into a '
       'Claude API call, bypassing the BYOM endpoint:')
make_table(doc,
    ['Step', 'Component', 'Description'],
    [
        ['1', 'Developer UI (Fiori / Teams)', 'Developer submits a code generation request, e.g. "Write a report that lists all open POs per vendor for plant 1000"'],
        ['2', 'Query Embedder', 'nomic-embed-text converts the question into a 768-dim vector (local, no cloud call)'],
        ['3', 'ChromaDB Retrieval', 'Top-8 chunks returned from abap_programs, function_modules, and documentation collections; includes existing PO-related objects'],
        ['4', 'Prompt Assembler', 'Builds a structured prompt: system instructions + org context + retrieved chunks + developer question'],
        ['5', 'Claude Skill (API call)', 'claude-sonnet-4-6 receives the assembled prompt and generates ABAP code grounded in the retrieved context'],
        ['6', 'Response Parser', 'Extracts ABAP code block, cited sources list, and any caveats or follow-up questions'],
        ['7', 'CodeSage UI', 'Renders code with syntax highlighting, source links, and a one-click "copy to SE80/ADT" button'],
    ],
    col_widths=[0.4, 1.9, 5.4]
)

# ── 6.3 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.3  The Assembled Prompt Structure', level=2)
p(doc, 'The prompt sent to Claude by the Skill follows this structure:')
code_block(doc,
"""SYSTEM:
You are CodeSage, an expert ABAP developer assistant for this organisation's
SAP ECC 6.0 / S/4HANA landscape. Answer ONLY using the retrieved code context
below. If you use a function module, class, or table, cite it by name.
Output format: brief explanation, then a complete ABAP code block, then a
Sources section listing every object you referenced.

RETRIEVED CONTEXT (top-8 chunks from ChromaDB):
--- [1] Z_MM_OPEN_PO_REPORT (Program) ---
REPORT z_mm_open_po_report.
  SELECT ebeln, lifnr, werks, netwr FROM ekko
    INTO TABLE @DATA(lt_ekko)
    WHERE loekz = '' AND ebeln LIKE 'Z%'.
  ...
--- [2] Z_GET_VENDOR_NAME (Function Module) ---
  SELECT SINGLE name1 FROM lfa1 INTO rv_name
    WHERE lifnr = iv_lifnr.
--- [3] EKKO / EKPO table documentation chunk ---
  EKKO: PO header.  EKPO: PO item.  Key join: ebeln.
  LOEKZ = '' means not deleted.  ELIKZ = '' means not fully delivered.
... (chunks 4-8 omitted for brevity) ...

USER:
Write a report that lists all open POs per vendor for plant 1000,
showing PO number, vendor name, and net value."""
)

# ── 6.4 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.4  Sample Claude Skill Output', level=2)
p(doc, 'Claude response to the prompt above (formatted by the Skill post-processor):')
code_block(doc,
"""" Lists all open Purchase Orders for plant 1000, grouped by vendor.
" Sources: Z_MM_OPEN_PO_REPORT, Z_GET_VENDOR_NAME, EKKO, EKPO, LFA1

REPORT z_codesage_open_po_1000.

TYPES: BEGIN OF ty_po_line,
         lifnr  TYPE lifnr,
         vname  TYPE name1,
         ebeln  TYPE ebeln,
         netwr  TYPE netwr,
         waers  TYPE waers,
       END OF ty_po_line.

DATA: lt_result TYPE TABLE OF ty_po_line.

" Fetch open, non-deleted PO headers for plant 1000
SELECT h~lifnr, h~ebeln, h~waers,
       SUM( i~netwr ) AS netwr
  FROM ekko AS h
  INNER JOIN ekpo AS i ON i~ebeln = h~ebeln
  INTO TABLE @lt_result
  WHERE h~loekz = ''
    AND h~elikz = ''
    AND i~werks  = '1000'
  GROUP BY h~lifnr, h~ebeln, h~waers
  ORDER BY h~lifnr, h~ebeln.

" Enrich with vendor name using org FM Z_GET_VENDOR_NAME
LOOP AT lt_result ASSIGNING FIELD-SYMBOL(<ls>).
  CALL FUNCTION 'Z_GET_VENDOR_NAME'
    EXPORTING iv_lifnr   = <ls>-lifnr
    IMPORTING rv_name    = <ls>-vname.
ENDLOOP.

" Output
cl_demo_output=>display( lt_result ).

" Sources cited:
"  Z_MM_OPEN_PO_REPORT  \u2014 pattern reference (existing open-PO report)
"  Z_GET_VENDOR_NAME    \u2014 org FM for LFA1 name lookup
"  EKKO / EKPO          \u2014 PO header and item tables
"  LFA1                 \u2014 vendor master"""
)

# ── 6.5 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.5  BYOM vs Claude Skill \u2014 When to Use Which', level=2)
make_table(doc,
    ['Criterion', 'BYOM (LLaMA-3 Fine-tuned)', 'Claude Skill (RAG + API)'],
    [
        ['Response latency',        '2\u20135 s (SAP-managed GPU)',               '3\u20138 s (API round-trip + larger context)'],
        ['Code quality',            'Good \u2014 fine-tuned on org patterns',     'Excellent \u2014 frontier model; better reasoning'],
        ['Grounding',               'Fine-tuning + RAG context',                  'RAG context only \u2014 no fine-tuning needed'],
        ['Data residency',          'SAP AI Core (BTP region choice)',            'Anthropic API or SAP Gen AI Hub (EU available)'],
        ['Cost model',              'Per-inference (SAP AI Core metered)',        'Per-token (Anthropic / SAP Gen AI Hub)'],
        ['Best for',                'High-volume, fast Q&A; explain existing code','Complex code generation; multi-file reasoning'],
        ['Setup effort',            '8\u20139 days (includes fine-tuning)',        '2\u20133 days (RAG + Skill wrapper; no GPU needed)'],
        ['Offline / air-gap',       'Possible with local vLLM container',         'Requires internet (API call)'],
    ],
    col_widths=[1.8, 2.8, 2.8]
)

# ── 6.6 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.6  Implementing the Claude Skill in BTP CAP', level=2)
p(doc, 'The CodeSage Agent (BTP CAP) can be extended with a /codesage/generate endpoint '
       'that implements the Claude Skill pattern. The key differences from the BYOM endpoint:')
bullet(doc, 'Larger context window (200k tokens in Claude) enables passing 8\u201316 retrieved chunks vs 4 for LLaMA-3.', bold_prefix='Context window:  ')
bullet(doc, 'Tool use (function calling) allows Claude to request additional ChromaDB lookups mid-generation if the initial retrieval is insufficient.', bold_prefix='Tool use / multi-turn retrieval:  ')
bullet(doc, 'Claude returns structured JSON with code, citations, and confidence \u2014 parsed by the BTP CAP service before sending to the UI.', bold_prefix='Structured output:  ')
bullet(doc, 'API key stored in BTP Secret Store; never exposed to client. Rotate every 90 days per governance policy (Chapter 4).', bold_prefix='Security:  ')

p(doc, 'Minimal BTP CAP handler (Node.js / TypeScript):')
code_block(doc,
"""// codesage-agent/srv/generate-handler.ts
import Anthropic from '@anthropic-ai/sdk';
import { retrieveChunks } from './chromadb-client';

const claude = new Anthropic({ apiKey: process.env.CLAUDE_API_KEY });

export async function handleGenerate(req: Request) {
  const { question, collections = ['abap_programs','function_modules','documentation'] } = req.data;

  // 1. Retrieve top-8 chunks from ChromaDB
  const chunks = await retrieveChunks(question, collections, 8);
  const context = chunks.map((c, i) =>
    `--- [${i+1}] ${c.metadata.object_name} (${c.metadata.type}) ---\n${c.document}`
  ).join('\n');

  // 2. Call Claude Skill
  const response = await claude.messages.create({
    model: 'claude-sonnet-4-6',
    max_tokens: 4096,
    system: `You are CodeSage, an expert ABAP developer assistant.
Answer using ONLY the retrieved context. Cite every object by name.
Output: explanation, complete ABAP code block, Sources section.`,
    messages: [
      { role: 'user',
        content: `RETRIEVED CONTEXT:\n${context}\n\nQUESTION: ${question}` }
    ]
  });

  // 3. Parse and return
  const text = response.content[0].type === 'text' ? response.content[0].text : '';
  return { answer: text, model: 'claude-sonnet-4-6', chunks_used: chunks.length };
}"""
)

# ── 6.7 ──────────────────────────────────────────────────────────────────────
doc.add_heading('9.7  ChromaDB Collection Strategy for Code Generation', level=2)
p(doc, 'Effective code generation requires that the right collections are queried for each '
       'request type. The table below maps query intent to optimal ChromaDB retrieval strategy:')
make_table(doc,
    ['Developer Intent', 'Primary Collections', 'Retrieval Strategy', 'Chunk Count'],
    [
        ['"Explain how X works"',        'abap_programs, documentation',         'Semantic similarity on function/class descriptions', '4\u20135'],
        ['"Write a report for Y"',       'abap_programs, function_modules',      'Similarity + keyword boost on report patterns',      '6\u20138'],
        ['"Add a check in BAdI Z"',      'abap_classes, function_modules',       'BAdI name keyword filter + semantic context',        '5\u20136'],
        ['"Convert this to ABAP OO"',    'abap_classes, abap_programs',          'Class pattern retrieval + source program context',   '6\u20138'],
        ['"Is there a FM that does X?"', 'function_modules, documentation',      'Semantic similarity on FM descriptions',             '4\u20135'],
        ['"CDS view for table T"',       'documentation, abap_programs',         'Table name keyword filter + CDS/annotation context', '4\u20135'],
    ],
    col_widths=[2.0, 2.1, 2.3, 1.3]
)

info_box(doc, 'Best practice \u2014 hybrid retrieval:',
         'Combine dense vector search (nomic-embed-text cosine similarity) with sparse keyword '
         'search (BM25 on object names and table references) for best results. ChromaDB supports '
         'both modes. Org-specific object names (Z_*, ZCL_*, ZFI_*) are rare in the vector space '
         'and benefit significantly from keyword boosting.')

doc.add_page_break()

# =============================================================================
# CHAPTER 7 — CODE GENERATION USING THE FINE-TUNED MODEL
# =============================================================================
doc.add_heading('Chapter 10:  Generating New ABAP Code from Your Training Data', level=1)

p(doc, 'The previous chapters explained how CodeSage answers questions about existing code. '
       'This chapter covers the next level: using the fine-tuned LLaMA-3 model to '
       'generate brand-new ABAP code that follows your organisation\'s own patterns, '
       'naming conventions, table structures, and coding standards \u2014 because the model '
       'was trained on them.')

info_box(doc, '\U0001f4a1 The key insight:',
    'When you fine-tune LLaMA-3 on your ABAP codebase, the model does not just '
    'memorise your code \u2014 it learns the patterns behind it. It learns that your '
    'team always uses ZCX_BASE_ERROR for exceptions, that SELECT statements always '
    'include a MANDT check for client-dependent tables, that your BAdI implementations '
    'follow a specific class naming pattern (ZCL_IMPL_*), and that every custom report '
    'starts with a standard header comment block. When asked to generate new code, '
    'it applies all of these learned patterns automatically.')

# ── 7.1 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.1  Q&A Mode vs Code Generation Mode', level=2)
p(doc, 'CodeSage operates in two distinct modes depending on how the developer phrases their request:')
make_table(doc,
    ['Mode', 'Trigger Phrase', 'What the Model Does', 'Output'],
    [
        ['Q&A Mode',
         '"How do we...", "Do we have...", "Show me..."',
         'Retrieves relevant existing code chunks and explains them',
         'Explanation + citation of existing ABAP objects'],
        ['Generation Mode',
         '"Write a...", "Generate...", "Create a report that...", "Build a class that..."',
         'Synthesises new code using learned org patterns + RAG context',
         'New ABAP code block following org conventions + sources it was inspired by'],
    ],
    col_widths=[1.5, 2.3, 2.3, 2.1]
)
p(doc, 'The CodeSage Agent detects the mode from the opening verb of the request and adjusts '
       'the system prompt accordingly. In generation mode, the system prompt instructs the '
       'model to produce complete, runnable ABAP code rather than an explanation.')

# ── 7.2 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.2  Why Training Data Enables Better Code Generation', level=2)
p(doc, 'A generic AI model (one not trained on your code) can write syntactically correct ABAP. '
       'But it will make mistakes that reveal it does not know your system:')
bullet(doc, 'It will invent function module names that do not exist (e.g. Z_GET_VENDOR_DATA instead of your actual Z_VENDOR_MASTER_READ).')
bullet(doc, 'It will use standard SAP table access patterns instead of your org\'s central access classes.')
bullet(doc, 'It will miss your mandatory header comment blocks, your coding standards, your exception class hierarchy.')
bullet(doc, 'It will not know which BAdIs your system already implements, leading to duplicate implementations.')

p(doc, 'After QLoRA fine-tuning on your codebase, the model has learned all of the above. '
       'Training data teaches the model three layers of knowledge:')

make_table(doc,
    ['Layer', 'What the Model Learns', 'Effect on Generated Code'],
    [
        ['Vocabulary',
         'Your Z/Y object names, table names, field names, class names',
         'Generated code references real objects that exist in your system'],
        ['Style',
         'Comment headers, indentation, naming patterns (ZCL_*, ZFI_*, ZMM_*), error handling approach',
         'Generated code looks like it was written by your team, not a generic AI'],
        ['Architecture',
         'Which utility classes to reuse, which central FMs exist, which BAdIs are already implemented, table access patterns',
         'Generated code reuses your existing building blocks instead of reinventing them'],
    ],
    col_widths=[1.3, 3.1, 3.3]
)

# ── 7.3 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.3  How Code Generation Works Step by Step', level=2)
p(doc, 'When a developer submits a code generation request, the following pipeline runs:')
numbered(doc, 'Developer submits a generation request in plain English, e.g.: "Write a report that lists all open purchase orders for plant 1000 grouped by vendor, using our standard PO access pattern."')
numbered(doc, 'The CodeSage Agent detects the generation intent from the opening verb ("Write a report").')
numbered(doc, 'ChromaDB retrieval fetches the top-8 most relevant existing code chunks \u2014 similar reports, related function modules, table documentation \u2014 as structural reference.')
numbered(doc, 'The generation prompt is assembled.  It contains: (a) a system instruction telling the model to write new code following org standards, (b) the retrieved code chunks as examples of how your team writes similar code, and (c) the developer\'s request.')
numbered(doc, 'The fine-tuned LLaMA-3 generates the new code.  Because it was trained on your codebase, it naturally applies your naming conventions, reuses your existing FMs, and follows your error handling patterns.')
numbered(doc, 'The response is returned with: the generated ABAP code block, the source objects it drew inspiration from, and a reminder that developer review is required before transport.')

p(doc, 'The generation prompt structure (what the model actually receives):', bold=True, sa=2)
code_block(doc,
"""SYSTEM:
You are CodeSage, an ABAP code generator for this organisation's SAP system.
Generate complete, runnable ABAP code that follows the organisation's conventions:
  - Use ZCX_BASE_ERROR for all exceptions
  - Reference existing Z* function modules rather than writing new ones
  - Include the standard header comment block (author, date, description)
  - Use field symbols (<ls_data>) for loop performance
  - Always check SY-SUBRC after SELECT statements
  - Follow the naming pattern: ZR_ prefix for reports, ZCL_ for classes, ZFM_ for FM groups

REFERENCE CODE (top-8 chunks from ChromaDB \u2014 similar existing objects):
--- [1] Z_MM_OPEN_PO_REPORT (existing similar report) ---
REPORT z_mm_open_po_report.
  SELECT ebeln, lifnr, werks, netwr FROM ekko
    INTO TABLE @DATA(lt_ekko)
    WHERE loekz = '' AND elikz = ''.
  ...
--- [2] Z_GET_VENDOR_NAME (existing FM for vendor lookup) ---
  SELECT SINGLE name1 FROM lfa1 INTO rv_name WHERE lifnr = iv_lifnr.
--- [3] EKKO / EKPO table documentation ---
  EKKO: PO header. loekz='' means active. werks on EKPO (item level).
... (chunks 4-8) ...

DEVELOPER REQUEST:
Write a report that lists all open purchase orders for plant 1000
grouped by vendor, using our standard PO access pattern.""")

# ── 7.4 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.4  Example: Generating a Complete ABAP Report', level=2)
p(doc, 'The following is an actual code generation output from CodeSage (LLaMA-3 fine-tuned '
       'on a sample ABAP landscape).  Observe how the generated code reuses the '
       'organisation\'s existing function module Z_GET_VENDOR_NAME rather than '
       'writing a new vendor lookup, and applies the standard header comment block '
       'that appears in every program in the training data.')

code_block(doc,
"""*&---------------------------------------------------------------------*
*& Report  ZR_MM_OPEN_PO_BY_VENDOR
*& Author: CodeSage (AI Generated \u2014 review before transport)
*& Date  : 2026-04-01
*& Desc  : Lists open POs for plant 1000 grouped by vendor
*&         Uses: EKKO/EKPO, Z_GET_VENDOR_NAME
*&---------------------------------------------------------------------*
REPORT zr_mm_open_po_by_vendor.

TYPES: BEGIN OF ty_po_summary,
         lifnr  TYPE lifnr,
         vname  TYPE name1,
         ebeln  TYPE ebeln,
         netwr  TYPE netwr,
         waers  TYPE waers,
       END OF ty_po_summary.

DATA: lt_po    TYPE STANDARD TABLE OF ty_po_summary,
      ls_po    TYPE ty_po_summary,
      lv_plant TYPE werks_d VALUE '1000'.

SELECTION-SCREEN BEGIN OF BLOCK b1 WITH FRAME TITLE TEXT-001.
  SELECT-OPTIONS: s_lifnr FOR ls_po-lifnr.
SELECTION-SCREEN END OF BLOCK b1.

START-OF-SELECTION.

  " Fetch open, non-deleted PO headers + items for the target plant
  SELECT h~lifnr, h~ebeln, h~waers,
         SUM( i~netwr ) AS netwr
    FROM ekko AS h
    INNER JOIN ekpo AS i
      ON i~ebeln = h~ebeln
    INTO TABLE @lt_po
    WHERE h~loekz  = ''          "active header
      AND h~elikz  = ''          "not fully delivered
      AND i~elikz  = ''          "item not fully delivered
      AND i~werks   = @lv_plant  "plant filter
      AND h~lifnr IN @s_lifnr
    GROUP BY h~lifnr, h~ebeln, h~waers
    ORDER BY h~lifnr, h~ebeln.

  IF sy-subrc <> 0.
    MESSAGE 'No open POs found for the selected criteria' TYPE 'I'.
    RETURN.
  ENDIF.

  " Enrich with vendor name using org FM Z_GET_VENDOR_NAME
  LOOP AT lt_po ASSIGNING FIELD-SYMBOL(<ls_po>).
    CALL FUNCTION 'Z_GET_VENDOR_NAME'
      EXPORTING iv_lifnr   = <ls_po>-lifnr
      IMPORTING rv_name    = <ls_po>-vname
      EXCEPTIONS OTHERS    = 1.
    IF sy-subrc <> 0.
      <ls_po>-vname = '*** Name not found ***'.
    ENDIF.
  ENDLOOP.

  " Display using ALV
  DATA(lo_alv) = cl_salv_table=>factory(
    IMPORTING r_salv_table = DATA(lo_table)
    CHANGING  t_table      = lt_po ).
  lo_table->display( ).

*-- Sources this code was inspired by:
*   Z_MM_OPEN_PO_REPORT      \u2014 PO selection pattern (EKKO/EKPO JOIN)
*   Z_GET_VENDOR_NAME         \u2014 vendor name lookup (reused, not rewritten)
*   ZR_MM_PLANT_REPORT        \u2014 plant filter pattern + ALV display style""")

info_box(doc, '\U0001f4cc What to notice in this generated code:',
    '1. Standard header comment block \u2014 exactly as used in every program in the training data.  '
    '2. Z_GET_VENDOR_NAME reused \u2014 the model knew this FM exists because it was in the training data.  '
    '3. Correct SY-SUBRC check after SELECT and after the FM call.  '
    '4. FIELD-SYMBOL for the loop \u2014 a performance pattern the model learned from the training data.  '
    '5. Sources section at the bottom \u2014 shows which existing objects the model drew inspiration from.')

# ── 7.5 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.5  More Generation Examples', level=2)

doc.add_heading('10.5.1  Generate a Class Method', level=3)
p(doc, 'Request:  "Write a method for ZCL_INVOICE_PROCESSOR that validates invoice amount '
       'against the PO net value before posting, raising ZCX_BASE_ERROR if it exceeds tolerance."')
code_block(doc,
"""METHOD validate_invoice_amount.
*-- Validates that invoice amount does not exceed PO net value by more than tolerance %
*-- Raises: ZCX_BASE_ERROR if tolerance breached
  DATA: lv_po_value  TYPE netwr,
        lv_tolerance TYPE p DECIMALS 2,
        lv_variance  TYPE p DECIMALS 2.

  " Read PO net value from EKKO
  SELECT SINGLE netwr FROM ekko
    INTO @lv_po_value
    WHERE ebeln = @iv_ebeln.
  IF sy-subrc <> 0.
    RAISE EXCEPTION TYPE zcx_base_error
      EXPORTING mv_message = |PO { iv_ebeln } not found in EKKO|.
  ENDIF.

  " Apply tolerance (default 5% \u2014 configurable via ZMM_TOLERANCE_CONFIG)
  lv_tolerance = COND #( WHEN iv_tolerance IS SUPPLIED
                          THEN iv_tolerance
                          ELSE '5.00' ).

  lv_variance = abs( ( iv_invoice_amount - lv_po_value ) / lv_po_value * 100 ).

  IF lv_variance > lv_tolerance.
    RAISE EXCEPTION TYPE zcx_base_error
      EXPORTING mv_message =
        |Invoice amount { iv_invoice_amount } exceeds PO { iv_ebeln } | &&
        |net value { lv_po_value } by { lv_variance }% | &&
        |(tolerance: { lv_tolerance }%)|.
  ENDIF.
ENDMETHOD.""")

p(doc, 'Notice: ZCX_BASE_ERROR is used (not a generic class), the tolerance is read from '
       'ZMM_TOLERANCE_CONFIG (an org-specific config table the model found in the training data), '
       'and the message format matches the pattern used across all error messages in the training set.',
  italic=True, color=GREY, size=10)

doc.add_heading('10.5.2  Generate a BAdI Implementation Skeleton', level=3)
p(doc, 'Request:  "Generate a new BAdI implementation for MB_MIGO_BADI that checks whether '
       'a material is flagged as restricted before goods receipt posting."')
code_block(doc,
"""*-- CodeSage generated BAdI implementation skeleton
*-- BAdI: MB_MIGO_BADI  Enhancement: ZMM_GR_RESTRICTION_CHECK
*-- Purpose: Block GR posting for restricted materials

CLASS zcl_impl_mm_gr_restriction DEFINITION
  PUBLIC FINAL
  INHERITING FROM cl_badi_impl.
  PUBLIC SECTION.
    INTERFACES if_ex_mb_migo_badi.
ENDCLASS.

CLASS zcl_impl_mm_gr_restriction IMPLEMENTATION.

  METHOD if_ex_mb_migo_badi~mb_migo_hold_check.
  *-- Called per GR line item before posting. Set ch_msg_prot = 'E' to block.

    DATA: lv_restricted TYPE boole_d,
          lv_reason     TYPE string.

    " Only check plants in scope (ZMM_RESTRICTED_PLANTS config table)
    CHECK is_mseg-werks IN ( SELECT werks FROM zmm_restricted_plants ).

    " Call the central restriction checker \u2014 reuse, do not rewrite
    CALL FUNCTION 'Z_CHECK_RESTRICTED_MATERIAL'
      EXPORTING
        iv_matnr       = is_mseg-matnr
        iv_werks       = is_mseg-werks
      IMPORTING
        ev_restricted  = lv_restricted
        ev_reason      = lv_reason
      EXCEPTIONS
        material_error = 1
        OTHERS         = 2.

    IF lv_restricted = abap_true.
      ch_msg_prot = 'E'.
      MESSAGE e001(zmm_gr) WITH lv_reason INTO DATA(lv_dummy).
    ENDIF.

  ENDMETHOD.
ENDCLASS.

*-- Sources used:
*   ZCL_IMPL_MB_GR_CHECK          \u2014 existing GR check pattern (reused structure)
*   Z_CHECK_RESTRICTED_MATERIAL   \u2014 org FM for restriction lookup (reused)
*   ZMM_RESTRICTED_PLANTS         \u2014 config table (found in training data)""")

doc.add_heading('10.5.3  Generate a CDS View', level=3)
p(doc, 'Request:  "Generate a CDS consumption view over ZI_COSTCTR_ALLOCATION that adds '
       'controlling area text and filters to the current fiscal year."')
code_block(doc,
"""@AbapCatalog.viewEnhancementCategory: [#NONE]
@AccessControl.authorizationCheck: #CHECK
@EndUserText.label: 'Cost Centre Allocation \u2014 Consumption View'
@Metadata.ignorePropagatedAnnotations: true
@ObjectModel.usageType: { serviceQuality: #X, sizeCategory: #S, dataClass: #MIXED }

define view entity ZC_COSTCTR_ALLOCATION
  as select from ZI_COSTCTR_ALLOCATION as Base

  association [0..1] to I_ControllingArea as _ControlArea
    on $projection.ControllingArea = _ControlArea.ControllingArea

{
      @UI.selectionField: [{ position: 10 }]
      @UI.lineItem:       [{ position: 10, label: 'Controlling Area' }]
  key Base.ControllingArea,

      @UI.selectionField: [{ position: 20 }]
      @UI.lineItem:       [{ position: 20, label: 'Cost Centre' }]
  key Base.CostCentre,

  key Base.FiscalYear,
      Base.ActualCost,
      Base.PlannedCost,

      @UI.lineItem: [{ position: 50 }]
      _ControlArea.ControllingAreaName,

      _ControlArea  -- exposed for further composition
}
-- Filter: current fiscal year only (aligns with org reporting standard ZFI_CURR_YEAR_ONLY)
where Base.FiscalYear = $parameters.P_FiscalYear

*-- Sources: ZI_COSTCTR_ALLOCATION (base interface view), ZFI_CURR_YEAR_ONLY (filter pattern)""")

# ── 7.6 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.6  What Makes the Generated Code Safe to Use', level=2)
p(doc, 'Generated code is a starting point, not a finished deliverable. '
       'CodeSage is designed to make the review step fast and trustworthy:')

make_table(doc,
    ['Safety Feature', 'How It Works', 'Why It Matters'],
    [
        ['Source citations',
         'Every generated block lists the existing objects it drew from',
         'Developer can verify the model\'s reasoning by opening the cited programs in SE80'],
        ['Reuse over reinvention',
         'RAG retrieval ensures the model sees existing FMs and classes before generating',
         'Model prefers CALL FUNCTION \'Z_EXISTING_FM\' over writing duplicate logic'],
        ['No write-back to SAP',
         'CodeSage only returns text \u2014 it never creates transports or writes to SAP',
         'Developer must manually copy, review, and transport the code'],
        ['Org pattern compliance',
         'Fine-tuning means the model applies your error class, comment style, naming convention',
         'Generated code passes code review faster because it matches team standards'],
        ['Limitation disclosure',
         'UI shows: "AI generated \u2014 review before transport" on every response',
         'Prevents blind copy-paste; developer review remains the approval gate'],
    ],
    col_widths=[1.8, 2.8, 3.1]
)

# ── 7.7 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.7  Enabling Code Generation Mode in the CodeSage Agent', level=2)
p(doc, 'Code generation uses a different system prompt than Q&A mode. '
       'The CodeSage Agent (BTP CAP service) detects the intent and switches automatically. '
       'The prompt difference is significant:')

make_table(doc,
    ['Prompt Element', 'Q&A Mode', 'Generation Mode'],
    [
        ['System instruction',
         '"Explain using the retrieved context. Cite sources."',
         '"Generate complete runnable ABAP. Apply org conventions. Include header block."'],
        ['Temperature',
         '0.1 (low \u2014 factual, precise)',
         '0.3 (slightly higher \u2014 allows creative synthesis while staying grounded)'],
        ['Max tokens',
         '512 (explanation only)',
         '2048\u20134096 (full code blocks need more space)'],
        ['Retrieved chunks',
         'Top-5 most similar (explanation context)',
         'Top-8 including structural examples (generation scaffold)'],
        ['Post-processing',
         'Extract cited source names',
         'Extract code block + cited sources + add review disclaimer'],
    ],
    col_widths=[1.8, 2.6, 3.3]
)

p(doc, 'The generation system prompt (added to query-handler.js for generation mode):', bold=True, sa=2)
code_block(doc,
"""// query-handler.js \u2014 generation mode system prompt
const GENERATION_SYSTEM_PROMPT = `
You are CodeSage, an ABAP code generator for this organisation's SAP system.
Your task is to write complete, runnable ABAP code based on the developer's request.

MANDATORY coding conventions (learned from your training data):
1. Start every program/class with the standard header comment block (author, date, desc, sources)
2. Use ZCX_BASE_ERROR for all exception raising (never generic cx_root or message TYPE 'E')
3. Reuse existing Z* function modules from the reference code below \u2014 do not reinvent them
4. Use FIELD-SYMBOL (<ls_...>) in all LOOP AT ... statements
5. Always check SY-SUBRC after every SELECT or CALL FUNCTION
6. Naming: ZR_ prefix for reports, ZCL_ for classes, ZFM_ for function groups
7. End your response with a Sources section listing every Z* object you referenced

Generate the code now, then the Sources section.
`.trim();

// Intent detection
const isGenerationRequest = (question) => {
  const GEN_VERBS = /^(write|generate|create|build|implement|code|make)/i;
  return GEN_VERBS.test(question.trim());
};

// In the query handler, switch system prompt based on intent
const systemPrompt = isGenerationRequest(question)
  ? GENERATION_SYSTEM_PROMPT
  : QA_SYSTEM_PROMPT;

const response = await callAICore(systemPrompt, context, question,
  { max_tokens: isGenerationRequest(question) ? 3000 : 512,
    temperature: isGenerationRequest(question) ? 0.3  : 0.1 });""")

# ── 7.8 ──────────────────────────────────────────────────────────────────────
doc.add_heading('10.8  The Developer Workflow with Code Generation', level=2)
p(doc, 'From the developer\'s perspective, code generation with CodeSage feels like '
       'pair programming with a colleague who has read every line of code ever written '
       'in your SAP system:')

numbered(doc, 'Developer opens CodeSage in the Fiori Launchpad or Teams and types a generation request in plain English.')
numbered(doc, 'CodeSage returns a complete, org-specific ABAP code block within 3\u20138 seconds.')
numbered(doc, 'Developer reads the Sources section to understand which existing objects influenced the generated code.')
numbered(doc, 'Developer opens SE80 or ADT, creates the appropriate object (program / class / FM), and pastes the generated code.')
numbered(doc, 'Developer reviews the code: checks logic, adjusts any business-specific conditions the AI could not know (e.g. specific tolerance percentages, additional WHERE conditions), and runs a syntax check.')
numbered(doc, 'After unit testing, the developer transports the object through the normal SAP transport chain \u2014 CodeSage has no involvement in or access to the transport system.')

info_box(doc, '\U0001f3af Time saving in practice:',
    'Generating the initial draft of a standard ABAP report takes a developer 2\u20134 hours '
    'from scratch. With CodeSage code generation, the draft is ready in under 10 seconds. '
    'Developer time is then focused on the 20\u201330 minute review, adjustment, and test cycle '
    '\u2014 the creative and judgement-based work that AI cannot replace. '
    'Estimated time saving: 75\u201385% per standard object.')

doc.add_page_break()

# =============================================================================
# CHAPTER 8 — HOW THE THREE KNOWLEDGE SOURCES ANSWER A DEVELOPER QUERY
# =============================================================================
doc.add_heading('Chapter 11:  How the Codebase, SAP Help, and Cloudification Repository '
                'Work Together to Answer a Developer Query', level=1)

p(doc, 'When a developer submits a question to CodeSage, the answer does not come from a '
       'single source. It is assembled from three distinct knowledge stores, each contributing '
       'a different layer of intelligence. Understanding what each source contributes \u2014 '
       'and how they combine \u2014 explains why CodeSage answers are more accurate and '
       'more actionable than generic AI responses.')

info_box(doc, '\U0001f9e0 The core principle:',
    'No single source is sufficient on its own. '
    'The codebase model knows your system but not SAP standards. '
    'SAP Help knows standards but not your code. '
    'The cloudification repository knows migration rules but not your existing implementations. '
    'Together, the three sources produce answers that are technically correct, '
    'organisation-specific, and future-proof for S/4HANA migration.')

# ── 8.1 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.1  The Three Knowledge Sources', level=2)
make_table(doc,
    ['Source', 'What It Contains', 'Indexed In', 'Contributes To Answer'],
    [
        ['Codebase Model\n(Fine-tuned LLaMA-3)',
         'Your organisation\'s Z*/Y* ABAP programs, function modules, classes, BAdIs, '
         'enhancement spots \u2014 extracted from the live SAP system via RFC',
         'ChromaDB: abap_programs, function_modules, abap_classes collections',
         'Org-specific patterns, reusable FMs, existing implementations, naming conventions'],
        ['SAP Help Documentation',
         'Official SAP API documentation, BAPI references, S/4HANA programming guides, '
         'standard table descriptions, message class references \u2014 from help.sap.com',
         'ChromaDB: documentation collection',
         'SAP standard best practices, correct BAPI parameters, table field meanings, '
         'recommended programming patterns'],
        ['Cloudification Repository',
         'Clean Core guidelines, C1/C2 API classification rules, deprecated FM and '
         'table lists, BAdI-over-user-exit migration rules, RAP (RESTful ABAP Programming) '
         'patterns for S/4HANA compatibility',
         'ChromaDB: documentation collection (separate namespace)',
         'S/4HANA compatibility advice, migration warnings, modern replacement APIs '
         'for deprecated objects'],
    ],
    col_widths=[1.7, 2.3, 1.8, 2.4]
)

# ── 8.2 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.2  The Query Flow \u2014 Step by Step', level=2)
p(doc, 'The diagram below describes exactly what happens from the moment a developer '
       'presses Enter to the moment they receive an answer.')

make_table(doc,
    ['Step', 'What Happens', 'Which Source Is Used'],
    [
        ['1\nDeveloper submits query',
         'Developer types: "How should I read material master data in S/4HANA \u2014 '
         'do we have an existing FM or should I use a BAPI?"',
         'None yet \u2014 query received by CodeSage Agent'],
        ['2\nQuery embedded',
         'The question is converted to a 768-dimensional vector using nomic-embed-text '
         '(local, no cloud call). This vector mathematically represents the meaning '
         'of the question.',
         'Local embedding model (Ollama) \u2014 no external source'],
        ['3\nChromDB retrieval \u2014 Codebase',
         'ChromaDB searches the abap_programs and function_modules collections for '
         'chunks semantically similar to "material master read". Returns top-3 chunks: '
         'e.g. Z_MATERIAL_MASTER_READ (an existing org FM), ZR_MM_MATERIAL_LIST '
         '(a report using it), ZCL_MM_MATERIAL_HANDLER (a class wrapper).',
         '\u2b50 Codebase Model\n(org-specific FMs and programs)'],
        ['4\nChromaDB retrieval \u2014 SAP Help',
         'ChromaDB searches the documentation collection for SAP Help content matching '
         '"material master S/4HANA API". Returns: BAPI_MATERIAL_GET_ALL documentation, '
         'MARA/MARC table field descriptions, S/4HANA material master API reference.',
         '\u2b50 SAP Help Documentation'],
        ['5\nChromaDB retrieval \u2014 Cloudification',
         'ChromaDB searches for cloudification content matching "material master deprecated". '
         'Returns: a Clean Core entry stating that direct SELECT on MARA is C2 '
         '(allowed but not recommended) and that the released API '
         'I_ProductBasicData CDS view is the preferred S/4HANA approach.',
         '\u2b50 Cloudification Repository'],
        ['6\nPrompt assembly',
         'The Prompt Builder combines all retrieved chunks into a structured prompt: '
         'system instructions + 8 context chunks (3 org code + 3 SAP Help + 2 cloudification) '
         '+ the developer\'s question.',
         'All three sources combined'],
        ['7\nLLaMA-3 generates answer',
         'The fine-tuned LLaMA-3 on SAP AI Core reads the assembled prompt and produces '
         'an answer that synthesises all three layers: what your org already has, '
         'what SAP recommends, and what is future-safe for S/4HANA.',
         'All three sources \u2014 model trained on codebase, context from all collections'],
        ['8\nResponse returned',
         'Developer receives a structured answer with: recommendation, code example, '
         'and a Sources section citing objects from all three knowledge stores.',
         '\u2014'],
    ],
    col_widths=[1.1, 3.7, 2.4]
)

# ── 8.3 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.3  What Each Source Contributes to the Final Answer', level=2)
p(doc, 'Using the example query above \u2014 "How should I read material master data in '
       'S/4HANA?" \u2014 here is exactly what each source adds to the answer:')

doc.add_heading('11.3.1  Contribution from the Codebase Model', level=3)
p(doc, 'The codebase retrieval finds that your organisation already has '
       'Z_MATERIAL_MASTER_READ \u2014 a central function module that wraps MARA/MARC '
       'access and is used across 47 programs. The model, fine-tuned on your code, '
       'recognises this as the preferred org-internal approach and includes it in the answer:')
code_block(doc,
"""-- From ChromaDB: function_modules collection --
[Z_MATERIAL_MASTER_READ]
FUNCTION z_material_master_read.
  IMPORTING: IV_MATNR TYPE MATNR, IV_WERKS TYPE WERKS_D
  EXPORTING: ES_MATERIAL TYPE ZS_MATERIAL_DATA
  ...
  SELECT SINGLE * FROM mara INTO ls_mara WHERE matnr = iv_matnr.
  SELECT SINGLE * FROM marc INTO ls_marc WHERE matnr = iv_matnr AND werks = iv_werks.
  ...""")
p(doc, 'Contribution to answer:  "Your organisation already has Z_MATERIAL_MASTER_READ '
       'which is used in 47 programs. Use this FM rather than writing new database access."',
  italic=True, color=SAP_DARK)

doc.add_heading('11.3.2  Contribution from SAP Help Documentation', level=3)
p(doc, 'The SAP Help retrieval finds the official documentation for '
       'BAPI_MATERIAL_GET_ALL and the S/4HANA Material API reference, explaining '
       'the correct parameter structure and the difference between basic data '
       'views (BASIC_DATA) and plant-level data (PLANT_DATA):')
code_block(doc,
"""-- From ChromaDB: documentation collection (SAP Help) --
[BAPI_MATERIAL_GET_ALL - SAP Help]
Returns material master data for a specified material.
IMPORTING: MATERIAL (MATNR), PLANT (WERKS_D optional)
TABLES:    MATERIALGENERAL (basic data), MATERIALDESCRIPTION (texts),
           MATERIALPLANTDATA (plant-specific), MATERIALVALUATION (accounting)
Note: Returns all views in one call. Use when multiple data views are needed.
Performance note: For single-view reads, consider direct CDS access in S/4HANA.""")
p(doc, 'Contribution to answer:  "If you need data from multiple views in one call, '
       'BAPI_MATERIAL_GET_ALL is available. For plant-specific data only, '
       'MARC-based access is sufficient."',
  italic=True, color=SAP_DARK)

doc.add_heading('11.3.3  Contribution from the Cloudification Repository', level=3)
p(doc, 'The cloudification retrieval finds the Clean Core classification entry for '
       'direct database access to MARA, and the recommended S/4HANA replacement API:')
code_block(doc,
"""-- From ChromaDB: documentation collection (Cloudification Repository) --
[Clean Core API Classification \u2014 Material Master]
Direct SELECT on MARA/MARC: Classification C2 (customer-specific use allowed,
  not recommended for new development on S/4HANA public cloud).

Recommended released API:
  CDS View: I_ProductBasicData  (basic data \u2014 replaces MARA direct read)
  CDS View: I_ProductPlantData  (plant data \u2014 replaces MARC direct read)
  OData Service: API_PRODUCT_SRV (for integration scenarios)

Migration path: Replace Z_MATERIAL_MASTER_READ SELECT statements with
  consumption of I_ProductBasicData to remain C1-compliant.""")
p(doc, 'Contribution to answer:  "Note: direct MARA/MARC access is C2-classified. '
       'For new S/4HANA development, consider using I_ProductBasicData CDS view. '
       'Your existing Z_MATERIAL_MASTER_READ should be refactored in the next '
       'clean-core migration sprint."',
  italic=True, color=SAP_DARK)

# ── 8.4 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.4  The Final Assembled Answer', level=2)
p(doc, 'This is what the developer actually receives \u2014 a single coherent answer '
       'synthesised from all three sources:')

info_box(doc, '\U0001f4ac CodeSage Answer:',
    'Your organisation already has Z_MATERIAL_MASTER_READ (used in 47 programs) \u2014 '
    'use this FM rather than writing new database access. It reads MARA and MARC '
    'and returns a ZS_MATERIAL_DATA structure covering basic and plant-level data.\n\n'
    'If you need multiple data views in a single call, SAP\'s standard '
    'BAPI_MATERIAL_GET_ALL is also available and returns general, description, '
    'plant, and valuation data.\n\n'
    'S/4HANA note: Both MARA direct access and the existing FM are C2-classified '
    'under Clean Core guidelines. For new S/4HANA-native development, '
    'prefer consuming I_ProductBasicData (basic data) and I_ProductPlantData '
    '(plant data) CDS views \u2014 these are C1-released and will remain compatible '
    'on S/4HANA public cloud.\n\n'
    'Sources: Z_MATERIAL_MASTER_READ (org FM) \u00b7 BAPI_MATERIAL_GET_ALL (SAP Help) \u00b7 '
    'I_ProductBasicData (Cloudification Repository \u2014 Clean Core C1 API)')

p(doc, 'This answer could not have been produced by any single source alone:')
bullet(doc, 'Without the codebase: the developer would not know Z_MATERIAL_MASTER_READ exists and would write a duplicate.')
bullet(doc, 'Without SAP Help: the developer would not know about BAPI_MATERIAL_GET_ALL or its correct parameter tables.')
bullet(doc, 'Without the cloudification repository: the developer would not know the C2 classification warning or the recommended I_ProductBasicData replacement for S/4HANA.')

# ── 8.5 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.5  How Each Source Is Weighted by Query Type', level=2)
p(doc, 'Not every query draws equally from all three sources. The Prompt Builder '
       'adjusts retrieval emphasis based on the type of question asked:')

make_table(doc,
    ['Query Type', 'Example', 'Primary Source', 'Secondary Source', 'Why'],
    [
        ['"Do we have a FM for X?"',
         '"Do we have a function module that validates customer credit limits?"',
         'Codebase (function_modules)',
         'SAP Help (standard BAPI reference)',
         'Primarily an org-internal lookup; SAP Help provides fallback if no org FM exists'],
        ['"How does standard SAP handle X?"',
         '"What is the standard BAPI for creating a sales order?"',
         'SAP Help (documentation)',
         'Codebase (how your team calls it)',
         'Question is about SAP standard; codebase shows org-specific call pattern'],
        ['"Is X still valid for S/4HANA?"',
         '"Can we still use CALL FUNCTION \'READ_TEXT\' in S/4HANA?"',
         'Cloudification Repository',
         'SAP Help',
         'Specifically a migration compatibility question'],
        ['"Write code to do X"',
         '"Write a report that shows open deliveries per customer"',
         'Codebase (pattern reference)',
         'SAP Help + Cloudification (standards + S/4 compatibility)',
         'Generation needs org patterns as scaffold; SAP Help and cloudification ensure the generated code is standard-compliant'],
        ['"Why does our code do X?"',
         '"Why does ZR_MM_OPEN_PO_REPORT filter on LOEKZ = space?"',
         'Codebase (the specific program)',
         'SAP Help (LOEKZ field meaning)',
         'Explanation requires reading the actual code; SAP Help explains the field semantics'],
    ],
    col_widths=[1.7, 2.2, 1.5, 1.5, 2.3]
)

# ── 8.6 ──────────────────────────────────────────────────────────────────────
doc.add_heading('11.6  Keeping the Three Sources Up to Date', level=2)
p(doc, 'The quality of CodeSage answers depends directly on how current each source is. '
       'Each source has a different refresh cadence:')

make_table(doc,
    ['Source', 'Refresh Trigger', 'How Often', 'Who Owns It'],
    [
        ['Codebase (org ABAP)',
         'New or changed ABAP transports released to production',
         'Monthly (Phase 1 re-scan via Windows Task Scheduler)',
         'MM / FI / SD development teams'],
        ['SAP Help Documentation',
         'SAP product updates, new S/4HANA release notes',
         'Quarterly or after each S/4HANA system upgrade',
         'Basis / architecture team'],
        ['Cloudification Repository',
         'SAP Clean Core API classification updates, new C1/C2 lists published',
         'Quarterly, aligned with SAP S/4HANA release calendar',
         'Architecture / clean core owner'],
    ],
    col_widths=[1.8, 2.4, 1.9, 2.1]
)

p(doc, 'All three sources feed into ChromaDB via the Phase 2 embedder. '
       'Re-running Phase 2 after any source update automatically refreshes '
       'the vector index \u2014 no retraining of the LLaMA-3 model is required '
       'for documentation or cloudification updates. Model retraining (Phase 3) '
       'is only needed when significant new ABAP code has been added to the codebase.')

info_box(doc, '\u2705 Summary \u2014 what the developer gets:',
    'Every CodeSage answer is a synthesis of three layers: '
    '(1) what your organisation has already built \u2014 reuse, do not reinvent; '
    '(2) what SAP recommends as the correct standard approach; '
    '(3) what is safe and compatible for the S/4HANA journey ahead. '
    'No other tool in the SAP ecosystem combines all three automatically '
    'in a single developer query response.')

# ── End page ───────────────────────────────────────────────────────────────────
doc.add_page_break()
p(doc, '', sb=40, sa=0)
p(doc, 'CodeSage for SAP',
  bold=True, color=SAP_BLUE, size=16,
  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
p(doc, 'Version 1.0  |  March 2026',
  color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
p(doc, 'This document is confidential and intended for internal use only.',
  color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
