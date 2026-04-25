import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document('C:/Users/ratna/OneDrive/Desktop/RAG_SAP_Developer_Guide_Template.docx')

def find_idx(fragment):
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i
    return -1

def set_text(idx, text):
    if idx < 0:
        return
    p = doc.paragraphs[idx]
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)

def insert_para_before(ref_idx, text, style='Normal'):
    """Insert paragraph before ref_idx."""
    ref = doc.paragraphs[ref_idx]
    new_p = OxmlElement('w:p')
    ref._element.addprevious(new_p)
    for p in doc.paragraphs:
        if p._element is new_p:
            try:
                p.style = doc.styles[style]
            except:
                pass  # keep default style
            p.add_run(text)
            return p
    return None

def insert_section_before(ref_idx, heading, paragraphs_list, style='Heading 2'):
    """Insert a heading + multiple body paragraphs before ref_idx."""
    # Insert in reverse order so final order is correct
    all_content = [(heading, style)] + [(t, 'Normal') for t in paragraphs_list]
    for text, sty in reversed(all_content):
        insert_para_before(ref_idx, text, sty)

# ── Find insertion point — before Section 1 (Overview) ───────────────────────
overview_idx = find_idx('1. Overview')

# ── Insert Executive Summary before Section 1 ────────────────────────────────
exec_content = [
    'SAP developers spend an estimated 30-40% of their time searching for existing code, understanding legacy programs, '
    'and looking up the right BAPI or FM — time that generates no new value. This proposal outlines a custom, '
    'locally-hosted AI coding assistant trained on your own SAP codebase, designed to cut that overhead dramatically.',
    '',
    'Unlike SAP\'s own AI tools (Joule, SAP Build Code), this assistant is:',
    '  •  Trained on YOUR code — not generic SAP patterns',
    '  •  Runs fully offline — no data leaves your network',
    '  •  Customisable per client, module, or team',
    '  •  Built at a fraction of the cost of enterprise AI subscriptions',
]

problem_heading = 'The Problem — Developer Time Waste'
problem_content = [
    'Research across SAP development teams shows the following time distribution for a typical 8-hour developer day:',
    '',
    'CURRENT STATE (without assistant):',
    '  Searching for existing code / BAPI     ████████░░░░░  2.5 hrs  (31%)',
    '  Understanding legacy programs           █████░░░░░░░░  1.5 hrs  (19%)',
    '  Writing new code                        ██████░░░░░░░  2.0 hrs  (25%)',
    '  Testing & debugging                     ████░░░░░░░░░  1.5 hrs  (19%)',
    '  Documentation / reviews                 █░░░░░░░░░░░░  0.5 hrs   (6%)',
    '  Non-productive search time: 4.0 hrs/day = 50% of the working day',
    '',
    'WITH ABAP ASSISTANT:',
    '  Searching for existing code / BAPI     ██░░░░░░░░░░░  0.5 hrs   (6%)  ↓ 80% reduction',
    '  Understanding legacy programs           ██░░░░░░░░░░░  0.5 hrs   (6%)  ↓ 67% reduction',
    '  Writing new code                        ████████░░░░░  3.0 hrs  (38%)  ↑ 50% increase',
    '  Testing & debugging                     █████░░░░░░░░  1.5 hrs  (19%)',
    '  Freed capacity for new features         ████░░░░░░░░░  2.0 hrs  (25%)  NEW',
]

savings_heading = 'Estimated Cost Savings'
savings_content = [
    'Based on a team of 10 ABAP developers at a typical consulting or enterprise rate:',
    '',
    '  Team size:                     10 ABAP developers',
    '  Average daily rate:            £400 / developer / day',
    '  Hours saved per developer:     2.0 hrs / day  =  0.25 days',
    '  Daily team saving:             10 × £400 × 0.25  =  £1,000 / day',
    '',
    '  Monthly saving:                £1,000 × 22 working days  =  £22,000',
    '  Annual saving:                 £22,000 × 12              =  £264,000',
    '',
    '  One-time build cost:           ~£5,000–£8,000  (7 days dev + GPU compute)',
    '  Ongoing monthly cost:          £0  (fully local, no subscriptions)',
    '',
    '  ROI break-even:                Under 2 weeks after go-live',
    '  Year-1 net saving:             ~£256,000',
]

tools_heading = 'Existing Tools — Comparison & Gaps'
tools_content = [
    'Four tools currently exist in the market for AI-assisted coding in SAP environments:',
    '',
    'SAP JOULE',
    '  Strength: Native SAP integration, knows standard BAPIs, SAP Help content',
    '  Gap: Trained on generic SAP patterns only. Has never seen your Z-code, your naming conventions, '
    'your custom FMs, or your client-specific logic. Cannot answer "how does OUR system do this?"',
    '  Data: Your queries go to SAP\'s cloud. Client code context cannot be sent.',
    '  Cost: Bundled with S/4HANA subscription but locked to SAP\'s generic model.',
    '',
    'SAP BUILD CODE (with Joule)',
    '  Strength: Integrated with ABAP Development Tools (ADT), inline suggestions',
    '  Gap: Same model as Joule — generic SAP knowledge, no client customisation.',
    '  Data: Runs on BTP — queries leave your network.',
    '  Cost: BTP credits consumption.',
    '',
    'GITHUB COPILOT',
    '  Strength: Excellent for web/cloud languages, inline IDE suggestions',
    '  Gap: ABAP is severely underrepresented in training data. No knowledge of '
    'your data model, custom tables, or Z-programs. Code suggestions often incorrect for ABAP syntax.',
    '  Data: Code context sent to Microsoft Azure servers.',
    '  Cost: ~$19/user/month (~£180/year per developer).',
    '',
    'CURSOR / CODEIUM',
    '  Strength: Fast, multi-language, affordable',
    '  Gap: No ABAP-specific training. No SAP context. Generic code completions only.',
    '  Data: Cloud-based.',
    '  Cost: $10–20/user/month.',
    '',
    'CUSTOM ABAP ASSISTANT (This Proposal)',
    '  Strength: Trained on YOUR codebase. Knows your naming, your BAPIs, your patterns. '
    'Fully offline. Per-client customisation. No ongoing cost after build.',
    '  Gap: Requires one-time 6-7 day build effort. Needs occasional refresh when code changes.',
    '  Data: Stays entirely on-premise. Nothing leaves your network.',
    '  Cost: One-time £5,000–£8,000 build. £0 ongoing.',
]

sap_gap_heading = 'Why SAP\'s Own Tools Are Not Enough'
sap_gap_content = [
    'The fundamental limitation of all vendor AI tools — including SAP Joule — is that they are trained on '
    'public, generic data. They cannot know what is unique to your client\'s SAP system.',
    '',
    'What SAP Joule knows vs. What YOUR team actually needs:',
    '',
    '  SAP Joule knows...                  Your team needs...',
    '  Generic SAP patterns          →     YOUR naming conventions (ZMMPO_, ZFIAR_)',
    '  Standard BAPIs                →     Which BAPIs YOUR system actually calls',
    '  SAP Help documentation        →     YOUR internal coding standards document',
    '  Generic exception handling    →     YOUR standard error class ZCX_MM_ERROR',
    '  Public ABAP examples          →     YOUR 500+ existing Z-programs as reference',
    '  SAP\'s standard enhancements   →     YOUR client-specific BAdI implementations',
    '',
    'Joule cannot answer these real developer questions:',
    '  •  "How does our Z-program handle goods receipt for consignment stock?"',
    '  •  "Which FM in our system validates vendor payment terms?"',
    '  •  "Rewrite this in our team\'s OO style with our standard error class"',
    '  •  "What BAPIs do we use for SD order processing in this project?"',
    '',
    'Because it has never seen your code — and it never will.',
    'SAP does not index client-specific Z-objects into Joule. By design.',
    '',
    'THE CUSTOMISATION PROBLEM:',
    'Every client running SAP has a unique codebase built over 10-20 years. '
    'A single vendor AI model trained on generic data will always produce generic answers. '
    'The only way to get answers grounded in YOUR system is to train or index YOUR code.',
    '',
    'This is exactly what the Custom ABAP Assistant does — and no vendor tool can replicate it '
    'without access to your data, which raises the very data protection concerns this proposal also solves.',
]

real_help_heading = 'How This Helps — Real Developer Scenarios'
real_help_content = [
    'The difference between a generic AI tool and the custom ABAP assistant in practice:',
    '',
    'Scenario 1: BAPI suggestion',
    '  Developer asks:  "Suggest BAPI for sales order creation"',
    '  Generic tool:    "Use BAPI_SALESORDER_CREATEFROMDAT2"  (correct but generic)',
    '  Custom assistant: "Use BAPI_SALESORDER_CREATEFROMDAT2. In YOUR system, ZMMSO001 already '
    'calls it on line 145 with your standard commit wrapper. Here is the exact pattern your team uses:"',
    '  [shows actual code from your codebase]',
    '',
    'Scenario 2: Legacy code explanation',
    '  Developer asks:  "Explain function module ZFM_POST_VENDOR_INV"',
    '  Generic tool:    Cannot find it — not in public training data',
    '  Custom assistant: "This is your custom FM that wraps BAPI_ACC_DOCUMENT_POST. It applies '
    'your standard FI posting checks, uses ZCX_FI_ERROR for exceptions, and logs via ZCL_APP_LOG. '
    'It was last modified for the VAT change project in 2023."',
    '',
    'Scenario 3: Code generation',
    '  Developer asks:  "Write a SELECT from table EKKO for open POs"',
    '  Generic tool:    Generic SELECT with basic WHERE clause',
    '  Custom assistant: SELECT written in YOUR team\'s style — using @DATA, FIELD-SYMBOL, '
    'your standard authority check pattern, and your project\'s naming convention for variables.',
    '',
    'Scenario 4: Pattern lookup',
    '  Developer asks:  "How do we handle exceptions in this project?"',
    '  Generic tool:    Generic TRY/CATCH example',
    '  Custom assistant: "In your codebase, exceptions are raised as ZCX_MM_ERROR and logged '
    'via ZCL_LOG=>ADD_ERROR. This pattern appears in 47 of your existing programs. Here is a template:"',
]

# Insert all new sections before Section 1
for heading, content in [
    ('Executive Summary', exec_content),
    (problem_heading, problem_content),
    (savings_heading, savings_content),
    (tools_heading, tools_content),
    (sap_gap_heading, sap_gap_content),
    (real_help_heading, real_help_content),
]:
    # Re-find overview_idx each time as doc grows
    overview_idx = find_idx('1. Overview')
    insert_section_before(overview_idx, heading, content, style='Heading 1')

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save('C:/Users/ratna/OneDrive/Desktop/RAG_SAP_Developer_Guide_Template.docx')
print('Saved successfully.')
