"""create_blog.py — generates blog.docx: educational blog version of CodeSage."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r'c:\Users\ratna\apps\ABAP 2\blog.docx'

SAP_BLUE  = RGBColor(0x0A, 0x6E, 0xD1)
SAP_DARK  = RGBColor(0x1B, 0x6C, 0xA8)
DARK      = RGBColor(0x1A, 0x1A, 0x1A)
GREY      = RGBColor(0x88, 0x88, 0x88)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x1E, 0x8B, 0x4C)
ORANGE    = RGBColor(0xE6, 0x7E, 0x22)

# ── XML helpers ───────────────────────────────────────────────────────────────

def cell_bg(cell, hex6):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex6)
    tcPr.append(shd)

def cell_border(cell, color='0A6ED1', sz='10'):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for side in ['top', 'bottom', 'left', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        tcB.append(b)
    tcPr.append(tcB)

# ── Content helpers ───────────────────────────────────────────────────────────

def p(doc, text, bold=False, italic=False, color=None, size=None,
      align=None, sb=None, sa=None, style='Normal', indent=None):
    para = doc.add_paragraph(style=style)
    run  = para.add_run(text)
    if bold:   run.bold   = True
    if italic: run.italic = True
    if color:  run.font.color.rgb = color
    if size:   run.font.size = Pt(size)
    if align:  para.alignment = align
    if sb is not None: para.paragraph_format.space_before = Pt(sb)
    if sa is not None: para.paragraph_format.space_after  = Pt(sa)
    if indent is not None: para.paragraph_format.left_indent = Inches(indent)
    return para

def bullet(doc, text, bold_prefix=None, level=0):
    style = 'List Bullet' if level == 0 else 'List Bullet 2'
    para  = doc.add_paragraph(style=style)
    if bold_prefix:
        r = para.add_run(bold_prefix); r.bold = True
    para.add_run(text)
    return para

def numbered(doc, text, bold_prefix=None):
    para = doc.add_paragraph(style='List Number')
    if bold_prefix:
        r = para.add_run(bold_prefix); r.bold = True
    para.add_run(text)
    return para

def code_block(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return para

def info_box(doc, label, body, bg='EAF4FF', border='0A6ED1', label_color=None):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    cell_bg(cell, bg)
    cell_border(cell, border, '8')
    para = cell.paragraphs[0]
    r1 = para.add_run(label + '  '); r1.bold = True
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = label_color or SAP_BLUE
    r2 = para.add_run(body)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    doc.add_paragraph()

def outcome_box(doc, text):
    """Green outcome box."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    cell_bg(cell, 'E9F7EF')
    cell_border(cell, '1E8B4C', '10')
    para = cell.paragraphs[0]
    r1 = para.add_run('\u2705  Phase Outcome:  '); r1.bold = True
    r1.font.name = 'Calibri'; r1.font.size = Pt(11)
    r1.font.color.rgb = GREEN
    r2 = para.add_run(text)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    doc.add_paragraph()

def make_table(doc, headers, rows, col_widths=None, hdr_bg='0A6ED1', font_size=10):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        cell_bg(c, hdr_bg)
        para = c.paragraphs[0]; para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = para.add_run(h)
        r.bold = True; r.font.name = 'Calibri'
        r.font.size = Pt(font_size); r.font.color.rgb = WHITE
    for ri, row_data in enumerate(rows):
        for ci, val in enumerate(row_data):
            c = tbl.rows[ri + 1].cells[ci]
            if ri % 2 == 1: cell_bg(c, 'EAF4FF')
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Calibri'; r.font.size = Pt(font_size)
    if col_widths:
        for row in tbl.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def divider(doc, color=SAP_BLUE):
    p(doc, '\u2500' * 80, color=color, size=8, align=WD_ALIGN_PARAGRAPH.CENTER, sb=6, sa=6)

def phase_header(doc, number, title, tagline):
    """Styled phase chapter opener."""
    tbl  = doc.add_table(rows=1, cols=1)
    cell = tbl.rows[0].cells[0]
    cell_bg(cell, '0A6ED1')
    cell_border(cell, '0A6ED1', '0')
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = para.add_run(f'PHASE {number}  \u2014  {title}\n')
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(18)
    r1.font.color.rgb = WHITE
    r2 = para.add_run(tagline)
    r2.font.name = 'Calibri'; r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0xBD, 0xD7, 0xF5)
    doc.add_paragraph()

ICONS = {
    'powershell': '\U0001f4bb',   # 💻
    'python':     '\U0001f40d',   # 🐍
    'kaggle':     '\u2601\ufe0f',  # ☁️
    'browser':    '\U0001f310',   # 🌐
    'docker':     '\U0001f433',   # 🐳
    'editor':     '\U0001f4dd',   # 📝
    'monitor':    '\U0001f4ca',   # 📊
}
RUN_BG = {
    'powershell': 'EBF5FB',
    'python':     'F4ECF7',
    'kaggle':     'EAF7F0',
    'browser':    'FEF9E7',
    'docker':     'EBF5FB',
    'editor':     'FDF2F8',
    'monitor':    'F0F3F4',
}
RUN_BORDER = {
    'powershell': '2E86C1',
    'python':     '7D3C98',
    'kaggle':     '1E8B4C',
    'browser':    'D4AC0D',
    'docker':     '2874A6',
    'editor':     'A93226',
    'monitor':    '717D7E',
}
RUN_LABEL_TEXT = {
    'powershell': 'Windows PowerShell  (Win + X \u2192 Windows PowerShell)',
    'python':     'Save as .py file  \u2192  run with Python in PowerShell',
    'kaggle':     'Kaggle Notebook cell  (kaggle.com \u2192 your notebook)',
    'browser':    'Web Browser  \u2192  SAP AI Launchpad or BTP Cockpit',
    'docker':     'Windows PowerShell  (Docker Desktop must be running)',
    'editor':     'Text Editor  (Notepad, VS Code, or Notepad++)',
    'monitor':    'Windows PowerShell  \u2192  monitoring / verification',
}

def run_label(doc, kind):
    """Render a small coloured 'Run in:' banner before a code block."""
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    cell_bg(cell, RUN_BG[kind])
    cell_border(cell, RUN_BORDER[kind], '6')
    para = cell.paragraphs[0]
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    r1 = para.add_run(f'{ICONS[kind]}  Run in:  ')
    r1.bold = True; r1.font.name = 'Calibri'; r1.font.size = Pt(9)
    r1.font.color.rgb = RGBColor(*bytes.fromhex(RUN_BORDER[kind]))
    r2 = para.add_run(RUN_LABEL_TEXT[kind])
    r2.font.name = 'Calibri'; r2.font.size = Pt(9)
    r2.font.color.rgb = DARK

def apply_styles(doc):
    for name, sz, col, sb, sa in [
        ('Heading 1', 20, SAP_BLUE, 24, 8),
        ('Heading 2', 14, SAP_DARK, 14, 6),
        ('Heading 3', 12, DARK,     8,  4),
    ]:
        s = doc.styles[name]
        s.font.name = 'Calibri'; s.font.size = Pt(sz)
        s.font.color.rgb = col; s.font.bold = True
        s.paragraph_format.space_before   = Pt(sb)
        s.paragraph_format.space_after    = Pt(sa)
        s.paragraph_format.keep_with_next = True
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.space_after = Pt(5)

# =============================================================================
# BUILD DOCUMENT
# =============================================================================
doc = Document()
for s in doc.sections:
    s.top_margin = s.bottom_margin = Inches(1)
    s.left_margin = s.right_margin = Inches(1.2)
    s.page_width  = Inches(8.27)
    s.page_height = Inches(11.69)
apply_styles(doc)

# =============================================================================
# COVER PAGE
# =============================================================================
p(doc, '', sb=24, sa=0)
p(doc, 'CodeSage for SAP',
  bold=True, color=SAP_BLUE, size=34,
  align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
p(doc, 'Building an AI-Powered ABAP Knowledge Platform from Scratch',
  bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=6)
p(doc, 'An educational step-by-step guide \u2014 fine-tuning LLaMA-3 on your SAP codebase '
       'using QLoRA, ChromaDB, and SAP AI Core BYOM',
  color=GREY, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, sb=0, sa=16)
p(doc, '\u2500' * 72, color=SAP_BLUE, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=16)
p(doc, 'Technologies:  LLaMA-3 8B  \u00b7  QLoRA  \u00b7  ChromaDB  \u00b7  SAP AI Core  \u00b7  BTP CAP  \u00b7  Python',
  color=SAP_DARK, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=4)
p(doc, 'Prerequisites:  Windows PC  \u00b7  Free Kaggle account  \u00b7  SAP BTP subscription',
  color=GREY, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, sb=0, sa=20)

make_table(doc,
    ['Phase', 'What It Builds', 'Runs Where', 'Duration'],
    [
        ['Phase 1', 'Extracts ABAP source from SAP via RFC',          'Windows PC',         '1\u20132 hrs'],
        ['Phase 2', 'Embeds code into ChromaDB vector index',         'Windows PC (CPU)',   '~2 hrs'],
        ['Phase 3', 'Fine-tunes LLaMA-3 with QLoRA, deploys to BYOM', 'Kaggle (free GPU)', '6\u201312 hrs'],
        ['Phase 4', 'Deploys CAP agent on BTP for developer queries', 'SAP BTP',            '1\u20132 hrs'],
    ],
    col_widths=[0.8, 2.9, 1.8, 1.2]
)
doc.add_page_break()

# =============================================================================
# INTRODUCTION — CHAPTER 1 SUMMARY
# =============================================================================
doc.add_heading('Introduction: The Invisible Archive', level=1)

p(doc, 'Every organisation running SAP for more than a few years accumulates thousands of '
       'custom programs, function modules, BAdI implementations, and enhancement spots. '
       'This code represents years of institutional knowledge \u2014 but almost none of it '
       'is discoverable. Ask a developer "do we already have a function module that validates '
       'vendor payment terms?" and the honest answer, in most organisations, is "I don\'t know."')

p(doc, 'The result is predictable:')
bullet(doc, 'Duplicate development \u2014 multiple programs solving the same problem in different ways.')
bullet(doc, 'Slow onboarding \u2014 new partners re-invent solutions the previous team already built.')
bullet(doc, 'Architecture drift \u2014 each team designs in their own style, compounding inconsistency.')
bullet(doc, 'Higher costs \u2014 redundant development effort across every new project.')

p(doc, 'CodeSage for SAP solves this by turning your ABAP codebase into a conversational '
       'knowledge assistant. A developer types a plain English question \u2014 '
       '"How do we handle dunning in our system?" \u2014 and receives a precise answer '
       'in 2\u20135 seconds, citing the actual function modules and programs in your system '
       'by name. Not a generic SAP answer. Your SAP.')

info_box(doc, '\u26a1 How it works:',
    'Phase 1 extracts your ABAP source via RFC. Phase 2 builds a local vector search index '
    '(ChromaDB). Phase 3 fine-tunes LLaMA-3 8B Instruct on your code using QLoRA \u2014 '
    'runs free on Kaggle \u2014 and deploys it to SAP AI Core via BYOM. Phase 4 deploys '
    'a BTP CAP agent that orchestrates retrieval + generation for every developer query. '
    'All code stays inside your SAP BTP perimeter.')

make_table(doc,
    ['The Problem', 'CodeSage Solution'],
    [
        ['"I don\'t know if we already have this."',          'Semantic search over 20,000 ABAP objects in <100 ms'],
        ['New teams re-build existing solutions.',            'Fine-tuned model knows your naming + patterns'],
        ['Documentation is out of date or missing.',         'Answers grounded in live source code via RAG'],
        ['Knowledge leaves with each departing developer.',  'Institutional knowledge encoded in a permanent model'],
    ],
    col_widths=[3.2, 3.5]
)

divider(doc)
doc.add_page_break()

# =============================================================================
# PHASE 1
# =============================================================================
phase_header(doc, 1, 'Scan & Extract',
             'Connect to SAP via RFC \u2014 extract every custom ABAP object as readable files')

doc.add_heading('What This Phase Does', level=2)

p(doc, 'Phase 1 reads your SAP system and saves every custom ABAP object (Z*/Y* namespace) '
       'as plain text files on your local Windows machine. It is entirely read-only \u2014 '
       'nothing is changed in SAP, nothing is sent to the cloud. The output is a folder of '
       '.abap source files and .json metadata sidecars that feed both Phase 2 and Phase 3.')

p(doc, 'Think of it as taking a complete photograph of your codebase at a point in time. '
       'The extractor uses pyrfc \u2014 a Python library that speaks SAP\'s RFC protocol \u2014 '
       'to call standard SAP function modules (RPY_PROGRAM_READ, RPY_FUNCTIONMODULE_READ) '
       'and pull source code line by line. A credential filter strips any hard-coded '
       'passwords before files are written to disk.')

info_box(doc, '\U0001f4a1 Educational note:',
    'RFC (Remote Function Call) is how external programs communicate with SAP. '
    'Think of it as a typed API over a direct TCP socket. pyrfc is the Python wrapper '
    'that handles the protocol details, so your script can call any SAP function module '
    'just like a developer calling it from within ABAP.')

# ── Artifacts ────────────────────────────────────────────────────────────────
doc.add_heading('Artifacts', level=2)

p(doc, 'Folder Structure', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase1_extract/
\u2502   \u251c\u2500\u2500 scan_config.yaml          # Controls which objects/namespaces to scan
\u2502   \u251c\u2500\u2500 extractor.py              # RFC extractor script
\u2502   \u2514\u2500\u2500 abap_files/               # OUTPUT \u2014 created when extractor runs
\u2502       \u251c\u2500\u2500 Z_VALIDATE_VENDOR.abap  # ABAP source (one file per object)
\u2502       \u251c\u2500\u2500 Z_VALIDATE_VENDOR.json  # Metadata sidecar
\u2502       \u251c\u2500\u2500 ZCL_PO_STATUS_HANDLER.abap
\u2502       \u251c\u2500\u2500 ZCL_PO_STATUS_HANDLER.json
\u2502       \u2514\u2500\u2500 ... (5,000 \u2013 20,000 files)""")

p(doc, 'scan_config.yaml \u2014 controls what gets extracted', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""sap:
  host:     "your-sap-host"
  sysnr:    "00"
  client:   "100"
  user:     "CODESAGE_RFC"
  password: "${SAP_RFC_PASSWORD}"    # Never store passwords in plain text

output_dir: "./abap_files"

namespaces:
  - "Z*"
  - "Y*"

object_types:
  - PROG    # Programs / Reports
  - FUGR    # Function Groups (contains Function Modules)
  - CLAS    # ABAP Classes
  - INTF    # Interfaces
  - ENHO    # BAdI Implementations

exclude_packages:
  - "$TMP"  # Skip local dev objects""")

p(doc, 'extractor.py \u2014 RFC extraction script', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""import pyrfc, json, pathlib, yaml, re

cfg = yaml.safe_load(open('scan_config.yaml'))
out = pathlib.Path(cfg['output_dir'])
out.mkdir(exist_ok=True)

# Strip any hard-coded credentials before saving to disk
CRED_RE = re.compile(
    r'(password|passwd|pwd|secret|apikey)\\s*=\\s*[\\'\\"][^\\'\\\"]+[\\'\\"]', re.I)

conn = pyrfc.Connection(
    ashost=cfg['sap']['host'],   sysnr=cfg['sap']['sysnr'],
    client=cfg['sap']['client'], user=cfg['sap']['user'],
    passwd=cfg['sap']['password'])

def save_object(name, obj_type, source, meta):
    source_clean = CRED_RE.sub('[REDACTED]', source)
    (out / f'{name}.abap').write_text(source_clean, encoding='utf-8')
    (out / f'{name}.json').write_text(json.dumps(meta, indent=2), encoding='utf-8')

# Extract Programs
progs = conn.call('RPY_DIRECTORY_FINISH', OBJECT_TYPE='PROG', GENERIC_NAME='Z*')
for obj in progs['TADIR']:
    try:
        r   = conn.call('RPY_PROGRAM_READ', PROG_NAME=obj['OBJ_NAME'])
        src = '\\n'.join(l['LINE'] for l in r['SOURCE'])
        save_object(obj['OBJ_NAME'], 'PROG', src, {
            'name': obj['OBJ_NAME'], 'type': 'PROG',
            'package': obj.get('DEVCLASS',''), 'changed': obj.get('LDATE','')})
    except Exception as e:
        print(f'  SKIP {obj["OBJ_NAME"]}: {e}')

# Extract Function Modules
fugs = conn.call('RPY_DIRECTORY_FINISH', OBJECT_TYPE='FUGR', GENERIC_NAME='Z*')
for obj in fugs['TADIR']:
    try:
        r   = conn.call('RPY_FUNCTIONMODULE_READ', FUNCNAME=obj['OBJ_NAME'])
        src = '\\n'.join(l['LINE'] for l in r.get('SOURCE',[]))
        save_object(obj['OBJ_NAME'], 'FUGR', src, {
            'name': obj['OBJ_NAME'], 'type': 'FM',
            'package': obj.get('DEVCLASS',''), 'changed': obj.get('LDATE','')})
    except Exception as e:
        print(f'  SKIP {obj["OBJ_NAME"]}: {e}')

print(f'Done. {len(list(out.glob("*.abap")))} objects saved to {out}/')""")

# ── Prerequisites + How to Run ────────────────────────────────────────────────
doc.add_heading('Prerequisites & How to Run', level=2)

info_box(doc, '\U0001f4cb All prerequisites in one block:',
    'Python 3.10+  (python.org)  \u00b7  '
    'SAP NW RFC SDK 7.50 for Windows 64-bit  (SAP Support Portal, SAP Note 2573790)  \u00b7  '
    'pip install pyrfc pyyaml  \u00b7  '
    'RFC user CODESAGE_RFC in SAP with S_RFC ACTVT=16 authorisation for function groups '
    'RPY_*, SFES_*, RS_*, SOBJ_*  \u00b7  '
    'TCP port 33XX open between Windows PC and SAP app server  \u00b7  '
    'All commands run in Windows PowerShell.')

p(doc, 'Step 1 \u2014 Install SAP NW RFC SDK', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# 1. Download "SAP NW RFC SDK 7.50 Windows 64-bit" from SAP Support Portal (Note 2573790)
# 2. Extract to C:\\nwrfcsdk\\  (you should see C:\\nwrfcsdk\\lib\\sapnwrfc.dll)
# 3. Add to Windows PATH:
#    Win+S -> "Edit the system environment variables" -> Environment Variables
#    -> System Variables -> Path -> Edit -> New -> C:\\nwrfcsdk\\lib -> OK
# 4. Verify in a NEW PowerShell window:
Get-Item "C:\\nwrfcsdk\\lib\\sapnwrfc.dll"   # Should list the file""")

p(doc, 'Step 2 \u2014 Create virtual environment and install dependencies', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase1_extract

python -m venv codesage-env
.\\codesage-env\\Scripts\\Activate.ps1
# If blocked by execution policy, run this first:
Set-ExecutionPolicy -ExecutionPolicy RemoteSigned -Scope CurrentUser

pip install pyrfc pyyaml
python -c "import pyrfc; print('pyrfc OK')"
python -c "import yaml;  print('pyyaml OK')" """)

p(doc, 'Step 3 \u2014 Store SAP password as environment variable', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Current session only (recommended for dev/test):
$env:SAP_RFC_PASSWORD = "your_sap_password_here"

# Persist across reboots (stored in Windows registry):
[System.Environment]::SetEnvironmentVariable(
    "SAP_RFC_PASSWORD", "your_sap_password_here", "User")""")

p(doc, 'Step 4 \u2014 Test RFC connection', bold=True, sa=2)
run_label(doc, 'python')
code_block(doc,
"""# Save as test_connection.py, then run: python test_connection.py
import pyrfc, os, yaml
cfg  = yaml.safe_load(open('scan_config.yaml'))
conn = pyrfc.Connection(
    ashost=cfg['sap']['host'],   sysnr=cfg['sap']['sysnr'],
    client=cfg['sap']['client'], user=cfg['sap']['user'],
    passwd=os.environ['SAP_RFC_PASSWORD'])
result = conn.call('RFC_PING')
print('RFC connection successful:', result)
# Expected:  RFC connection successful: {}""")

run_label(doc, 'powershell')
code_block(doc, 'python test_connection.py')

p(doc, 'Step 5 \u2014 Run the extractor', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase1_extract
python extractor.py
# Console shows: SKIP messages for deleted/inaccessible objects (normal)
# Finishes with: Done. 8,432 objects saved to .\\abap_files\\""")

p(doc, 'Step 6 \u2014 Verify output', bold=True, sa=2)
run_label(doc, 'monitor')
code_block(doc,
"""# Count extracted files
(Get-ChildItem -Path ".\\abap_files" -Filter "*.abap").Count
# Typical: 5,000 to 20,000

# Confirm no plain-text passwords in output
Select-String -Path ".\\abap_files\\*.abap" -Pattern "password" -CaseSensitive:$false |
  Where-Object { $_.Line -notmatch "REDACTED" }
# If nothing returned: credential filter is working correctly""")

p(doc, 'Step 7 \u2014 Schedule monthly re-run', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""$action  = New-ScheduledTaskAction `
    -Execute "C:\\codesage\\phase1_extract\\codesage-env\\Scripts\\python.exe" `
    -Argument "C:\\codesage\\phase1_extract\\extractor.py" `
    -WorkingDirectory "C:\\codesage\\phase1_extract"
$trigger = New-ScheduledTaskTrigger -Monthly -DaysOfMonth 1 -At "02:00AM"
Register-ScheduledTask `
    -TaskName "CodeSage Phase1 Monthly Extract" `
    -Action $action -Trigger $trigger -RunLevel Highest""")

# ── Outcome ───────────────────────────────────────────────────────────────────
outcome_box(doc,
    'A local folder (abap_files/) containing .abap source and .json metadata for every '
    'custom object in your SAP system (Z*/Y* namespace). Credentials are filtered, '
    'the RFC connection is read-only, and no data leaves your network. '
    'This folder is the shared input for Phase 2 (vector indexing) and Phase 3 (fine-tuning).')

divider(doc)
doc.add_page_break()

# =============================================================================
# PHASE 2
# =============================================================================
phase_header(doc, 2, 'Index & Embed',
             'Chunk ABAP files into logical pieces \u2014 convert them to vectors \u2014 store in ChromaDB')

doc.add_heading('What This Phase Does', level=2)

p(doc, 'Phase 2 transforms the raw ABAP source files from Phase 1 into a searchable semantic '
       'vector database. After this phase, you can ask "which code in our system validates '
       'vendor payment terms?" and get an answer in under 100 milliseconds \u2014 entirely '
       'on-premise, no cloud, no API key.')

p(doc, 'Two steps run in sequence:')
bullet(doc, 'Chunking: splits each ABAP file at logical boundaries (FORM/ENDFORM, METHOD/ENDMETHOD). '
            'One chunk = one function module, one method, or one report section. '
            '150\u2013500 tokens per chunk.', bold_prefix='Step 1 \u2014 chunker.py:  ')
bullet(doc, 'Embedding: converts each chunk to a 768-number vector using the nomic-embed-text model '
            'running locally via Ollama. Stores vectors + original text in ChromaDB (4 collections). '
            'No GPU required \u2014 runs on a standard CPU laptop.',
            bold_prefix='Step 2 \u2014 embedder.py:  ')

info_box(doc, '\U0001f4a1 What is a vector embedding?',
    'A vector is a list of numbers (e.g. 768 floats) that represents the meaning of a text. '
    'Two pieces of code that do similar things will produce vectors that are mathematically '
    'close to each other \u2014 even if they use completely different words. '
    'This is what lets CodeSage match "validate vendor payment terms" to '
    'Z_VALIDATE_VENDOR_PAYTERMS without any keyword matching.')

p(doc, 'The four ChromaDB collections:')
make_table(doc,
    ['Collection', 'Content', 'Chunk Unit'],
    [
        ['abap_programs',    'Executable reports and includes',              'One FORM routine or report body'],
        ['function_modules', 'Function module source + parameters',          'One complete function module'],
        ['abap_classes',     'ABAP OO classes, interfaces, methods',         'One METHOD block'],
        ['documentation',    'SAP Help, Clean Core guidelines (optional)',    'One documentation paragraph'],
    ],
    col_widths=[1.8, 2.6, 2.3]
)

# ── Artifacts ────────────────────────────────────────────────────────────────
doc.add_heading('Artifacts', level=2)

p(doc, 'Folder Structure', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase2_index/
\u2502   \u251c\u2500\u2500 chunker.py                # Splits .abap files into logical chunks
\u2502   \u251c\u2500\u2500 embedder.py               # Embeds chunks and stores in ChromaDB
\u2502   \u2514\u2500\u2500 chromadb_store/           # OUTPUT \u2014 persistent vector database
\u2502       \u251c\u2500\u2500 abap_programs/
\u2502       \u251c\u2500\u2500 function_modules/
\u2502       \u251c\u2500\u2500 abap_classes/
\u2502       \u2514\u2500\u2500 documentation/""")

p(doc, 'chunker.py \u2014 split ABAP source into logical segments', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""import pathlib, json, re
from dataclasses import dataclass

@dataclass
class Chunk:
    text: str; object_name: str; chunk_type: str; collection: str

FORM_RE      = re.compile(r'^FORM\\s+(\\w+)', re.M | re.I)
METHOD_RE    = re.compile(r'^\\s*METHOD\\s+(\\w+)', re.M | re.I)
ENDFORM_RE   = re.compile(r'^ENDFORM', re.M | re.I)
ENDMETHOD_RE = re.compile(r'^\\s*ENDMETHOD', re.M | re.I)

def chunk_abap(abap_path, meta):
    src   = abap_path.read_text(encoding='utf-8')
    lines = src.splitlines()
    name  = meta['name'];  typ = meta['type']
    chunks = []

    if typ == 'FM':                             # Function module = one chunk
        chunks.append(Chunk(src, name, 'FUNCTION', 'function_modules'))

    elif typ == 'CLAS':                         # Class = split on METHOD/ENDMETHOD
        buf, mname = [], None
        for line in lines:
            m = METHOD_RE.match(line)
            if m:
                mname = m.group(1); buf = [line]
            elif ENDMETHOD_RE.match(line) and mname:
                buf.append(line)
                chunks.append(Chunk('\\n'.join(buf), f'{name}.{mname}',
                                    'METHOD', 'abap_classes'))
                mname = None; buf = []
            elif mname:
                buf.append(line)
        if not chunks:
            chunks.append(Chunk(src, name, 'CLASS_DEF', 'abap_classes'))

    else:                                       # Program = split on FORM/ENDFORM
        buf, fname = [], None
        for line in lines:
            m = FORM_RE.match(line)
            if m:
                fname = m.group(1); buf = [line]
            elif ENDFORM_RE.match(line) and fname:
                buf.append(line)
                chunks.append(Chunk('\\n'.join(buf), f'{name}.{fname}',
                                    'FORM', 'abap_programs'))
                fname = None; buf = []
            elif fname:
                buf.append(line)
        chunks.append(Chunk(src[:3000], name, 'BODY', 'abap_programs'))

    return [c for c in chunks if len(c.text.strip()) > 50]

if __name__ == '__main__':
    abap_dir = pathlib.Path('../phase1_extract/abap_files')
    all_chunks = []
    for abap_file in abap_dir.glob('*.abap'):
        meta_file = abap_file.with_suffix('.json')
        meta = json.loads(meta_file.read_text()) if meta_file.exists() \
               else {'name': abap_file.stem, 'type': 'PROG'}
        all_chunks.extend(chunk_abap(abap_file, meta))
    print(f'Total chunks: {len(all_chunks)}')""")

p(doc, 'embedder.py \u2014 embed chunks and store in ChromaDB', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""import chromadb, requests, json, pathlib
from chunker import chunk_abap

db = chromadb.PersistentClient(path='./chromadb_store')

COLLECTIONS = {
    'abap_programs':    db.get_or_create_collection('abap_programs'),
    'function_modules': db.get_or_create_collection('function_modules'),
    'abap_classes':     db.get_or_create_collection('abap_classes'),
    'documentation':    db.get_or_create_collection('documentation'),
}

def embed(text):
    \"\"\"Call local nomic-embed-text via Ollama. No cloud, no API key.\"\"\"
    r = requests.post('http://localhost:11434/api/embeddings',
                      json={'model': 'nomic-embed-text', 'prompt': text})
    return r.json()['embedding']           # 768-dimensional vector

def store_chunk(chunk, idx):
    col = COLLECTIONS[chunk.collection]
    vec = embed(chunk.text[:2000])
    col.add(
        ids=[f'{chunk.object_name}_{idx}'],
        embeddings=[vec],
        documents=[chunk.text],
        metadatas=[{'object': chunk.object_name, 'type': chunk.chunk_type}])

abap_dir = pathlib.Path('../phase1_extract/abap_files')
for i, abap_file in enumerate(sorted(abap_dir.glob('*.abap'))):
    meta_file = abap_file.with_suffix('.json')
    meta = json.loads(meta_file.read_text()) if meta_file.exists() \
           else {'name': abap_file.stem, 'type': 'PROG'}
    for j, chunk in enumerate(chunk_abap(abap_file, meta)):
        store_chunk(chunk, j)
    if i % 100 == 0:
        print(f'  Indexed {i} files...')

print('ChromaDB indexing complete.')
for name, col in COLLECTIONS.items():
    print(f'  {name}: {col.count()} chunks')""")

# ── Prerequisites + How to Run ────────────────────────────────────────────────
doc.add_heading('Prerequisites & How to Run', level=2)

info_box(doc, '\U0001f4cb All prerequisites in one block:',
    'Phase 1 complete (abap_files/ folder exists)  \u00b7  '
    'Python venv from Phase 1 still active  \u00b7  '
    'pip install chromadb requests  \u00b7  '
    'Ollama installed (ollama.com/download \u2014 free, runs as Windows service)  \u00b7  '
    'ollama pull nomic-embed-text  (274 MB download, CPU-only, no GPU needed)  \u00b7  '
    '2\u20134 GB free disk space for the ChromaDB store.')

p(doc, 'Step 1 \u2014 Activate venv and install dependencies', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase2_index
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

pip install chromadb requests
python -c "import chromadb; print('chromadb', chromadb.__version__)" """)

p(doc, 'Step 2 \u2014 Install Ollama and download embedding model', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Install Ollama from ollama.com/download (Windows installer, free)
# After install, open a NEW PowerShell window:
ollama pull nomic-embed-text

ollama list
# NAME                    SIZE
# nomic-embed-text:latest 274 MB""")

p(doc, 'Step 3 \u2014 Test Ollama is responding', bold=True, sa=2)
run_label(doc, 'python')
code_block(doc,
"""# Save as test_ollama.py, run: python test_ollama.py
import requests
r = requests.post("http://localhost:11434/api/embeddings",
                  json={"model": "nomic-embed-text", "prompt": "hello SAP"})
vec = r.json()["embedding"]
print(f"Embedding OK: {len(vec)} dimensions")
# Expected: Embedding OK: 768 dimensions""")

run_label(doc, 'powershell')
code_block(doc, 'python test_ollama.py')

p(doc, 'Step 4 \u2014 Dry-run chunker to preview chunk count', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""python chunker.py
# Expected: Total chunks: 42,318  (varies by codebase size)""")

p(doc, 'Step 5 \u2014 Run the embedder (main Phase 2 step, ~2 hours on CPU)', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""python embedder.py
#   Indexed 0 files...  Indexed 100 files...  ...
# ChromaDB indexing complete.
#   abap_programs:    18,420 chunks
#   function_modules: 12,305 chunks
#   abap_classes:      9,841 chunks""")

p(doc, 'Step 6 \u2014 Test semantic search', bold=True, sa=2)
run_label(doc, 'python')
code_block(doc,
"""# Save as test_search.py, run: python test_search.py
import chromadb, requests

def embed(text):
    r = requests.post("http://localhost:11434/api/embeddings",
                      json={"model": "nomic-embed-text", "prompt": text})
    return r.json()["embedding"]

db  = chromadb.PersistentClient(path="./chromadb_store")
col = db.get_collection("function_modules")
results = col.query(
    query_embeddings=[embed("validate vendor payment terms")],
    n_results=3)
for doc_text, meta in zip(results["documents"][0], results["metadatas"][0]):
    print(f"[{meta['object']}]  {doc_text[:80]}...")
# Expected top result: [Z_VALIDATE_VENDOR_PAYTERMS]  FUNCTION Z_VALIDATE...""")

run_label(doc, 'powershell')
code_block(doc, 'python test_search.py')

# ── Outcome ───────────────────────────────────────────────────────────────────
outcome_box(doc,
    'ChromaDB running on-premise with 4 collections and tens of thousands of embedded ABAP '
    'chunks. Any developer query can be matched to the most relevant code in your system in '
    'under 100 ms \u2014 without sending a single line of code to the cloud. '
    'This vector index is used at runtime (Phase 4) to feed context into every LLM response.')

divider(doc)
doc.add_page_break()

# =============================================================================
# PHASE 3
# =============================================================================
phase_header(doc, 3, 'Fine-Tune, Merge & Package for BYOM',
             'Teach LLaMA-3 your codebase using QLoRA \u2014 deploy to SAP AI Core as a live endpoint')

doc.add_heading('What This Phase Does', level=2)

p(doc, 'Phase 3 is the heart of CodeSage. By the end of it, you have a version of LLaMA-3 '
       'that has learned your naming conventions, your preferred BAPIs, your error class '
       'patterns, and your ABAP coding style \u2014 running as a live inference endpoint '
       'on SAP AI Core with no local GPU required at runtime.')

p(doc, 'The phase has five sequential steps:')
make_table(doc,
    ['Step', 'What It Does', 'Where It Runs', 'Duration'],
    [
        ['A \u2014 Generate training data', 'Claude reads ABAP chunks, writes QA pairs as JSONL',   'Windows PC (Python + Anthropic API)', '30\u201360 min'],
        ['B \u2014 QLoRA fine-tuning',      'Train LoRA adapters on LLaMA-3 8B',                   'Kaggle free GPU (T4 16 GB)',          '6\u201312 hrs'],
        ['C \u2014 Merge adapter',          'Bake LoRA deltas back into base model weights',        'Windows PC (CPU, 32 GB RAM)',         '~20 min'],
        ['D \u2014 Docker image',           'Package merged model + vLLM server as container',      'Windows PC (Docker Desktop)',         '10\u201320 min build + push'],
        ['E \u2014 BYOM deploy',            'Register ServingTemplate + deploy on SAP AI Core',     'SAP AI Launchpad (browser)',          '5\u201310 min'],
    ],
    col_widths=[1.9, 2.6, 2.0, 1.2]
)

info_box(doc, '\U0001f9e0 Why QLoRA and not full fine-tuning?',
    'Full fine-tuning of an 8B model needs an 80 GB A100 GPU \u2014 very expensive. '
    'QLoRA (Quantised Low-Rank Adaptation) freezes the base model in 4-bit compressed '
    'format and only trains tiny "adapter" matrices alongside certain layers. '
    'Only 0.17% of parameters are updated. The result fits on a free Kaggle T4 (16 GB), '
    'costs \u00a30, and reaches 90%+ of full fine-tune quality for code Q&A tasks.')

# ── Artifacts ────────────────────────────────────────────────────────────────
doc.add_heading('Artifacts', level=2)

p(doc, 'Folder Structure', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase3_finetune/
\u2502   \u251c\u2500\u2500 generate_qa.py              # Step A: generates training_data.jsonl
\u2502   \u251c\u2500\u2500 train_qlora.py              # Step B: QLoRA training (Kaggle notebook)
\u2502   \u251c\u2500\u2500 merge_model.py              # Step C: merges adapter into base model
\u2502   \u251c\u2500\u2500 Dockerfile                  # Step D: vLLM container definition
\u2502   \u251c\u2500\u2500 serving-template.yaml       # Step E: SAP AI Core BYOM registration
\u2502   \u251c\u2500\u2500 training_data.jsonl         # Generated QA pairs (OUTPUT of Step A)
\u2502   \u251c\u2500\u2500 lora-adapter/               # OUTPUT of Step B (~50 MB)
\u2502   \u2502   \u251c\u2500\u2500 adapter_config.json     # LoRA config (r=16, alpha=32)
\u2502   \u2502   \u2514\u2500\u2500 adapter_model.safetensors
\u2502   \u2514\u2500\u2500 merged_model/               # OUTPUT of Step C (~16 GB)
\u2502       \u251c\u2500\u2500 config.json
\u2502       \u251c\u2500\u2500 tokenizer.json
\u2502       \u2514\u2500\u2500 model-00001-of-00004.safetensors  (x4 shards)""")

p(doc, 'generate_qa.py \u2014 Step A: create training dataset via Claude', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""import anthropic, json, pathlib

client = anthropic.Anthropic()     # reads ANTHROPIC_API_KEY env var

def generate_qa(chunk_text, object_name):
    msg = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=512,
        messages=[{"role": "user", "content":
            f"You are an ABAP expert. Given this code from {object_name}, "
            f"write 4 developer questions whose answer is in this code.\\n\\n"
            f"CODE:\\n{chunk_text}\\n\\nReturn a JSON list of strings."}])
    questions = json.loads(msg.content[0].text)
    return [{"prompt": q, "completion": chunk_text} for q in questions]

with open("training_data.jsonl", "w") as f:
    for path in pathlib.Path("../phase1_extract/abap_files").glob("*.abap"):
        chunk = path.read_text(encoding="utf-8")[:2000]
        for pair in generate_qa(chunk, path.stem):
            f.write(json.dumps(pair) + "\\n")
# Result: 30,000 - 50,000 lines in training_data.jsonl""")

p(doc, 'train_qlora.py \u2014 Step B: QLoRA fine-tuning (run as Kaggle notebook cell)', bold=True, sa=2)
run_label(doc, 'kaggle')
code_block(doc,
"""from transformers import AutoModelForCausalLM, AutoTokenizer, TrainingArguments, BitsAndBytesConfig
from peft import LoraConfig, get_peft_model, TaskType
from trl import SFTTrainer
import torch

MODEL_ID = "meta-llama/Meta-Llama-3-8B-Instruct"

# Load base model in 4-bit NF4 (reduces GPU memory from ~28 GB to ~6-12 GB)
bnb_config = BitsAndBytesConfig(
    load_in_4bit=True,
    bnb_4bit_quant_type="nf4",
    bnb_4bit_compute_dtype=torch.bfloat16,
    bnb_4bit_use_double_quant=True)

model     = AutoModelForCausalLM.from_pretrained(MODEL_ID,
              quantization_config=bnb_config, device_map="auto")
tokenizer = AutoTokenizer.from_pretrained(MODEL_ID)

# Attach LoRA adapters \u2014 only these tiny matrices will be trained
lora_config = LoraConfig(
    r=16, lora_alpha=32,
    target_modules=["q_proj","k_proj","v_proj","o_proj"],
    lora_dropout=0.05, bias="none",
    task_type=TaskType.CAUSAL_LM)
model = get_peft_model(model, lora_config)
model.print_trainable_parameters()
# Output: trainable params: 13,631,488 (0.17% of 8B)

# Train
trainer = SFTTrainer(
    model=model, tokenizer=tokenizer,
    train_dataset=dataset,
    dataset_text_field="text",
    max_seq_length=2048,
    args=TrainingArguments(
        output_dir="./lora-adapter",
        num_train_epochs=3,
        per_device_train_batch_size=4,
        gradient_accumulation_steps=4,
        learning_rate=2e-4,
        fp16=True, save_strategy="epoch", logging_steps=50))
trainer.train()
trainer.save_model("./lora-adapter")   # ~50 MB adapter file""")

p(doc, 'merge_model.py \u2014 Step C: merge LoRA adapter into base model', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""from transformers import AutoModelForCausalLM, AutoTokenizer
from peft import PeftModel

base      = AutoModelForCausalLM.from_pretrained(
    "meta-llama/Meta-Llama-3-8B-Instruct",
    torch_dtype="auto", device_map="cpu")
tokenizer = AutoTokenizer.from_pretrained("meta-llama/Meta-Llama-3-8B-Instruct")

model  = PeftModel.from_pretrained(base, "./lora-adapter/")
merged = model.merge_and_unload()        # adapter deltas baked into base weights

merged.save_pretrained("./merged_model/")
tokenizer.save_pretrained("./merged_model/")
print("Merge complete. ./merged_model/ (~16 GB) ready for Docker packaging.")""")

p(doc, 'Dockerfile \u2014 Step D: vLLM container for SAP AI Core BYOM', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""FROM vllm/vllm-openai:latest
COPY ./merged_model/ /model/
EXPOSE 8080
CMD ["python", "-m", "vllm.entrypoints.openai.api_server", \\
     "--model",             "/model", \\
     "--port",              "8080", \\
     "--served-model-name", "codesage-llama3", \\
     "--max-model-len",     "8192", \\
     "--dtype",             "bfloat16"]""")

p(doc, 'serving-template.yaml \u2014 Step E: SAP AI Core BYOM registration', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""apiVersion: ai.sap.com/v1alpha1
kind: ServingTemplate
metadata:
  name: codesage-llama3
  labels:
    scenarios.ai.sap.com/id: codesage-scenario
    ai.sap.com/version: "1.0"
spec:
  template:
    spec:
      imagePullSecrets:
        - name: <your-registry-secret>
      containers:
        - name: kserve-container
          image: <your-registry>/codesage-llama3:v1
          ports:
            - containerPort: 8080
          resources:
            requests:
              memory: "20Gi"
            limits:
              nvidia.com/gpu: "1"    # SAP AI Core allocates a managed GPU""")

# ── Prerequisites + How to Run ────────────────────────────────────────────────
doc.add_heading('Prerequisites & How to Run', level=2)

info_box(doc, '\U0001f4cb All prerequisites in one block:',
    'Phase 1 complete (abap_files/ folder)  \u00b7  '
    'Anthropic API key (console.anthropic.com)  \u00b7  '
    'Kaggle free account with phone-verified GPU (kaggle.com)  \u00b7  '
    'HuggingFace account + read token + LLaMA-3 licence accepted  \u00b7  '
    'Docker Desktop for Windows (docker.com, free for personal use)  \u00b7  '
    'Container registry (Docker Hub free tier is sufficient)  \u00b7  '
    'SAP BTP subaccount with AI Core service enabled  \u00b7  '
    '32 GB RAM on Windows PC for the merge step  \u00b7  '
    'pip install anthropic transformers peft bitsandbytes trl accelerate')

p(doc, 'Step A \u2014 Generate training data (Windows PC, 30\u201360 min)', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

pip install anthropic
$env:ANTHROPIC_API_KEY = "sk-ant-your-key-here"

python generate_qa.py
# Reads each .abap file, sends to Claude, saves (question, code) pairs

(Get-Content .\\training_data.jsonl).Count
# Typical: 30,000 to 50,000 lines""")

p(doc, 'Step B \u2014 Fine-tune: choose your GPU option', bold=True, sa=2)

make_table(doc,
    ['Option', 'GPU', 'Cost', 'RAM', 'Duration', 'Best For'],
    [
        ['Option A \u2014 Kaggle (recommended for first run)',
         'NVIDIA T4 16 GB (free)',
         '\u00a30',
         '16 GB GPU + 13 GB RAM',
         '6\u201312 hrs',
         'Anyone, no hardware needed'],
        ['Option B \u2014 Local NVIDIA GPU',
         'RTX 3090 (24 GB) / RTX 4090 (24 GB) / A100 (40\u201380 GB)',
         'Electricity only',
         '24+ GB VRAM recommended',
         '2\u20136 hrs',
         'Faster iteration, larger batch, keep data on-prem'],
    ],
    col_widths=[2.0, 1.8, 0.7, 1.6, 0.9, 1.7]
)

# ── Option A: Kaggle ──────────────────────────────────────────────────────────
p(doc, 'Option A \u2014 Kaggle free GPU (T4 16 GB)', bold=True, color=SAP_BLUE, sa=2)
numbered(doc, 'Upload training_data.jsonl to Kaggle as a Dataset (kaggle.com \u2192 Datasets \u2192 New Dataset).')
numbered(doc, 'Create a new Kaggle Notebook. In the right panel set Accelerator = GPU T4 x1 and enable Internet access.')
numbered(doc, 'In the first notebook cell, install dependencies:')
run_label(doc, 'kaggle')
code_block(doc,
"""# Kaggle Cell 1 \u2014 install + HuggingFace login
!pip install -q transformers peft trl bitsandbytes accelerate datasets huggingface_hub
import huggingface_hub
huggingface_hub.login(token="hf_YOUR_READ_TOKEN")""")

numbered(doc, 'Paste the contents of train_qlora.py into Cell 2 (update the dataset path to your Kaggle input path).')
numbered(doc, 'Click Run All. Training takes 6\u201312 hours. You can close the browser \u2014 Kaggle keeps running.')
numbered(doc, 'When complete, download the lora-adapter/ folder as ZIP. Extract to C:\\codesage\\phase3_finetune\\lora-adapter\\')

info_box(doc, '\u26a0\ufe0f Kaggle limits:',
    '30 GPU hours/week free. Phone verification required to unlock GPU. '
    'Session times out after 12 hours of inactivity \u2014 save checkpoints every epoch '
    '(save_strategy="epoch" in TrainingArguments is already set). '
    'If training is longer than 12 hours, resume from the last checkpoint.')

# ── Option B: Local NVIDIA GPU ────────────────────────────────────────────────
p(doc, 'Option B \u2014 Local NVIDIA GPU (RTX 3090 / 4090 / A100)', bold=True, color=SAP_BLUE, sa=2)
p(doc, 'If you have a Windows machine with a modern NVIDIA GPU you can run the full '
       'training pipeline locally. This keeps all training data on-premise (no upload to '
       'Kaggle) and is typically 2\u20133\u00d7 faster than a T4 due to higher memory bandwidth.')

p(doc, 'Additional prerequisites for local GPU training:', bold=True, sa=2)
make_table(doc,
    ['Requirement', 'Details', 'Where to Get'],
    [
        ['NVIDIA GPU',            'RTX 3090 (24 GB VRAM) or better recommended. GTX 1080 Ti (11 GB) will work but may need smaller batch size.',
                                  'Your Windows machine'],
        ['NVIDIA Driver',         '525+ recommended. Check with: nvidia-smi',
                                  'nvidia.com/drivers or Windows Update'],
        ['CUDA Toolkit 11.8+',    'Required by bitsandbytes (4-bit quantisation). Check with: nvcc --version',
                                  'developer.nvidia.com/cuda-downloads'],
        ['bitsandbytes-windows',  'Windows build of the bitsandbytes quantisation library',
                                  'pip install bitsandbytes (see install step below)'],
        ['~100 GB free disk',     'For base model download (~16 GB) + training data + adapter + merged model',
                                  'Local SSD preferred'],
    ],
    col_widths=[1.8, 3.0, 2.4]
)

p(doc, 'Step B2.1 \u2014 Verify NVIDIA setup in PowerShell', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Check GPU is recognised and driver version
nvidia-smi
# Expected output includes: NVIDIA GeForce RTX 3090 / 4090 / A100
#   Driver Version: 535.xx  |  CUDA Version: 12.x

# Check CUDA toolkit
nvcc --version
# Expected: Cuda compilation tools, release 11.8 (or 12.x)

# If nvcc not found, CUDA Toolkit is not installed.
# Download from: https://developer.nvidia.com/cuda-downloads
# Choose: Windows -> x86_64 -> 11 -> exe (local)""")

p(doc, 'Step B2.2 \u2014 Install Python dependencies for local GPU training', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

# Install PyTorch with CUDA support (match your CUDA version)
# For CUDA 11.8:
pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118
# For CUDA 12.1:
pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu121

# Verify GPU is visible to PyTorch
python -c "import torch; print('CUDA available:', torch.cuda.is_available()); print('GPU:', torch.cuda.get_device_name(0))"
# Expected: CUDA available: True  |  GPU: NVIDIA GeForce RTX 3090

# Install training libraries
pip install transformers peft trl accelerate datasets huggingface_hub

# Install bitsandbytes (4-bit quantisation for QLoRA)
pip install bitsandbytes
# If bitsandbytes gives DLL errors on Windows, use the community Windows build:
pip install bitsandbytes --prefer-binary --extra-index-url=https://jllllll.github.io/bitsandbytes-windows-webui""")

p(doc, 'Step B2.3 \u2014 Download LLaMA-3 base model weights', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Log in to HuggingFace (you must have accepted the Meta LLaMA-3 licence first)
# Visit: https://huggingface.co/meta-llama/Meta-Llama-3-8B-Instruct and click Accept
python -c "import huggingface_hub; huggingface_hub.login(token='hf_YOUR_READ_TOKEN')"

# Download model (~16 GB, one-time)
python -c "
from huggingface_hub import snapshot_download
snapshot_download('meta-llama/Meta-Llama-3-8B-Instruct',
                  local_dir='C:\\\\codesage\\\\phase3_finetune\\\\base_model')
print('Download complete.')
"

# Check folder size (should be ~16 GB)
(Get-ChildItem -Path ".\\base_model" -Recurse | Measure-Object -Property Length -Sum).Sum / 1GB""")

p(doc, 'Step B2.4 \u2014 Update train_qlora.py to use the local model path', bold=True, sa=2)
p(doc, 'The only change needed from the Kaggle version is the MODEL_ID path '
       '\u2014 point it to the local download instead of HuggingFace Hub:')
run_label(doc, 'editor')
code_block(doc,
"""# In train_qlora.py, change this line:
MODEL_ID = "meta-llama/Meta-Llama-3-8B-Instruct"     # Kaggle version (downloads from HF)

# To this for local GPU:
MODEL_ID = "C:\\\\codesage\\\\phase3_finetune\\\\base_model"   # uses local download

# You can also increase batch size on a 24 GB GPU for faster training:
# In TrainingArguments, change:
per_device_train_batch_size=4   # T4 safe value
# To:
per_device_train_batch_size=8   # RTX 3090/4090 with 24 GB VRAM""")

p(doc, 'Step B2.5 \u2014 Run training locally', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

# Monitor GPU memory before starting
nvidia-smi

# Run training
python train_qlora.py

# In a second PowerShell window, monitor GPU usage while training runs:
while ($true) { nvidia-smi --query-gpu=utilization.gpu,memory.used,memory.total --format=csv,noheader; Start-Sleep 10 }
# Expected: gpu-util 95-100%,  memory.used ~14000 MiB / 24576 MiB (RTX 3090)

# Training completes in 2-6 hours (vs 6-12 on Kaggle T4)
# Adapter saved to: .\\lora-adapter\\""")

p(doc, 'Step B2.6 \u2014 Recommended GPU options at a glance', bold=True, sa=2)
make_table(doc,
    ['GPU', 'VRAM', 'Batch Size', 'Est. Training Time', 'Notes'],
    [
        ['NVIDIA RTX 3090',    '24 GB', '8',  '3\u20134 hrs',  'Best value for local training. Widely available.'],
        ['NVIDIA RTX 4090',    '24 GB', '8',  '2\u20133 hrs',  'Fastest consumer GPU. Same VRAM as 3090 but higher bandwidth.'],
        ['NVIDIA RTX 3080',    '10 GB', '2',  '6\u20138 hrs',  'Works but tight. Use gradient_accumulation_steps=8.'],
        ['NVIDIA A100 40 GB',  '40 GB', '16', '1\u20132 hrs',  'Data centre GPU. If available via corporate/cloud.'],
        ['Kaggle T4 (free)',   '16 GB', '4',  '6\u201312 hrs', 'No hardware needed. 30 hrs/week limit.'],
    ],
    col_widths=[1.8, 0.8, 1.0, 1.5, 3.1]
)

info_box(doc, '\U0001f4a1 RTX 3080 (10 GB) tip:',
    'If your GPU has less than 16 GB VRAM, reduce per_device_train_batch_size to 1 or 2 '
    'and increase gradient_accumulation_steps to 16 to maintain the effective batch size. '
    'Also add max_grad_norm=0.3 to TrainingArguments for stability on small batches. '
    'Training will be slower but will fit in 10 GB VRAM.')

p(doc, 'Step C \u2014 Merge adapter on Windows PC (~20 min, needs 32 GB RAM)', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune
..\\phase1_extract\\codesage-env\\Scripts\\Activate.ps1

pip install transformers peft accelerate
python merge_model.py

# Verify output
Get-ChildItem .\\merged_model
# Expected: config.json, tokenizer.json, model-00001-of-00004.safetensors, etc.""")

p(doc, 'Step D \u2014 Build and push Docker image (Docker Desktop required)', bold=True, sa=2)
run_label(doc, 'docker')
code_block(doc,
"""cd C:\\codesage\\phase3_finetune

docker login   # enter Docker Hub credentials
docker build -t codesage-llama3:v1 .    # ~18 GB image, takes 10-20 min first time

docker tag  codesage-llama3:v1  your-dockerhub-username/codesage-llama3:v1
docker push your-dockerhub-username/codesage-llama3:v1  # 30-60 min upload""")

p(doc, 'Step E \u2014 Deploy to SAP AI Core (browser, SAP AI Launchpad)', bold=True, sa=2)
numbered(doc, 'Open SAP AI Launchpad \u2192 ML Operations \u2192 Docker Registry Secrets \u2192 Add (your registry + credentials).')
numbered(doc, 'ML Operations \u2192 Serving Templates \u2192 Upload YAML \u2192 select serving-template.yaml.')
numbered(doc, 'ML Operations \u2192 Configurations \u2192 Create \u2192 select the codesage-llama3 template.')
numbered(doc, 'ML Operations \u2192 Deployments \u2192 Create \u2192 select the configuration. Wait for status: RUNNING (5\u201310 min).')
numbered(doc, 'Copy the Deployment URL and Deployment ID. Store both in BTP \u2192 Connectivity \u2192 Destinations as destination CODESAGE_AICORE.')

p(doc, 'Test the endpoint from PowerShell:', bold=True, sa=2)
run_label(doc, 'python')
code_block(doc,
"""# test_aicore.py  -- python test_aicore.py
import requests
token_resp = requests.post("https://<token-url>/oauth/token",
    data={"grant_type": "client_credentials"},
    auth=("<client-id>", "<client-secret>"))
token = token_resp.json()["access_token"]

response = requests.post(
    "https://<ai-core-url>/v2/inference/deployments/<deployment-id>/chat/completions",
    headers={"Authorization": f"Bearer {token}", "AI-Resource-Group": "default"},
    json={"model": "codesage-llama3",
          "messages": [{"role": "user", "content": "What is ABAP?"}],
          "max_tokens": 200})
print(response.json()["choices"][0]["message"]["content"])
# Expected: ABAP stands for Advanced Business Application Programming...""")

# ── Outcome ───────────────────────────────────────────────────────────────────
outcome_box(doc,
    'A fine-tuned LLaMA-3 8B model \u2014 trained on your organisation\'s ABAP codebase '
    'using QLoRA \u2014 is running as a live, OpenAI-compatible inference endpoint on '
    'SAP AI Core. No local GPU is needed at runtime. '
    'Training cost: \u00a30 on Kaggle free tier. '
    'The endpoint is reachable from BTP via a secure Destination. '
    'Phases 1, 2, and 3 are now all complete \u2014 Phase 4 wires them together.')

divider(doc)
doc.add_page_break()

# =============================================================================
# PHASE 4
# =============================================================================
phase_header(doc, 4, 'Runtime Query via SAP BTP',
             'Deploy the CodeSage Agent \u2014 developers type questions, get ABAP answers in seconds')

doc.add_heading('What This Phase Does', level=2)

p(doc, 'Phase 4 is the user-facing layer. It deploys a small Node.js service (BTP CAP) '
       'called the CodeSage Agent that orchestrates the full query pipeline:')

make_table(doc,
    ['Step', 'What Happens', 'Where'],
    [
        ['1', 'Developer types a plain English question in the UI or Teams bot',       'Browser / Teams'],
        ['2', 'Agent embeds the question using nomic-embed-text (local Ollama)',       'On-premise'],
        ['3', 'ChromaDB returns the top-5 most relevant ABAP code chunks',             'On-premise'],
        ['4', 'Agent builds a structured prompt: context chunks + question',           'BTP CAP service'],
        ['5', 'Prompt is sent to the fine-tuned LLaMA-3 on SAP AI Core BYOM',        'SAP AI Core'],
        ['6', 'Answer returned in 2\u20135 s, citing real program names by name',     'Browser / Teams'],
    ],
    col_widths=[0.4, 4.0, 2.3]
)

info_box(doc, '\U0001f4a1 Why BTP CAP?',
    'CAP (Cloud Application Programming) is SAP\'s opinionated Node.js/Java framework '
    'for BTP services. It handles OAuth2 (XSUAA), multi-tenancy, and Fiori integration '
    'out of the box. The CodeSage Agent is a single CAP service with one action endpoint: '
    'POST /codesage/query. Deploying via mbt + cf deploy takes under 5 minutes.')

# ── Artifacts ────────────────────────────────────────────────────────────────
doc.add_heading('Artifacts', level=2)

p(doc, 'Folder Structure', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""codesage/
\u251c\u2500\u2500 phase4_runtime/
\u2502   \u251c\u2500\u2500 codesage-agent/
\u2502   \u2502   \u251c\u2500\u2500 package.json            # Node.js dependencies
\u2502   \u2502   \u251c\u2500\u2500 mta.yaml                # BTP Multi-Target Application descriptor
\u2502   \u2502   \u251c\u2500\u2500 xs-security.json        # XSUAA roles: CodeSage.User, CodeSage.Admin
\u2502   \u2502   \u251c\u2500\u2500 srv/
\u2502   \u2502   \u2502   \u251c\u2500\u2500 codesage-service.cds  # CDS service endpoint definition
\u2502   \u2502   \u2502   \u251c\u2500\u2500 query-handler.js      # Core: RAG retrieval + AI Core call
\u2502   \u2502   \u2502   \u2514\u2500\u2500 chromadb-client.js    # Calls on-premise ChromaDB
\u2502   \u2502   \u2514\u2500\u2500 db/
\u2502   \u2502       \u2514\u2500\u2500 schema.cds            # Optional: query audit log
\u2502   \u2514\u2500\u2500 ui/
\u2502       \u2514\u2500\u2500 webapp/               # Fiori freestyle UI""")

p(doc, 'codesage-service.cds \u2014 service endpoint definition', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""// srv/codesage-service.cds
service CodeSageService @(requires: 'CodeSage.User') {

    // POST /codesage/query  { question, context_filter? }
    action query(
        question       : String not null,
        context_filter : String    // optional: 'function_modules' | 'abap_classes' etc.
    ) returns {
        answer      : String;
        sources     : array of { object_name: String; chunk_type: String; };
        model       : String;
        latency_ms  : Integer;
    };
}""")

p(doc, 'query-handler.js \u2014 core orchestration: RAG retrieval + AI Core call', bold=True, sa=2)
run_label(doc, 'editor')
code_block(doc,
"""// srv/query-handler.js
const axios  = require('axios');
const xsenv  = require('@sap/xsenv');
const { retrieveChunks } = require('./chromadb-client');

const { destination } = xsenv.getServices({ destination: { tag: 'destination' } });

module.exports = (srv) => {
  srv.on('query', async (req) => {
    const { question, context_filter } = req.data;
    const t0 = Date.now();

    // 1. Retrieve top-5 relevant ABAP chunks from on-premise ChromaDB
    const collections = context_filter
      ? [context_filter]
      : ['abap_programs', 'function_modules', 'abap_classes'];
    const chunks  = await retrieveChunks(question, collections, 5);
    const context = chunks
      .map((c, i) => `[${i+1}] ${c.metadata.object} (${c.metadata.type}):\\n${c.document}`)
      .join('\\n\\n');

    // 2. Build prompt
    const messages = [
      { role: 'system',
        content: 'You are CodeSage, an expert ABAP assistant. '
               + 'Answer ONLY using the retrieved context. '
               + 'Cite every object by name. End with a Sources section.' },
      { role: 'user',
        content: `RETRIEVED CONTEXT:\\n${context}\\n\\nQUESTION: ${question}` }
    ];

    // 3. Call SAP AI Core BYOM (OpenAI-compatible endpoint)
    const token    = await getAICoreToken();
    const deployId = process.env.AI_CORE_DEPLOYMENT_ID;
    const response = await axios.post(
      `${destination.url}/v2/inference/deployments/${deployId}/chat/completions`,
      { model: 'codesage-llama3', messages, max_tokens: 1024, temperature: 0.1 },
      { headers: { Authorization: `Bearer ${token}`, 'AI-Resource-Group': 'default' } });

    const answer  = response.data.choices[0].message.content;
    const sources = chunks.map(c => ({ object_name: c.metadata.object,
                                       chunk_type:  c.metadata.type }));
    return { answer, sources, model: 'codesage-llama3', latency_ms: Date.now() - t0 };
  });
};""")

# ── Prerequisites + How to Run ────────────────────────────────────────────────
doc.add_heading('Prerequisites & How to Run', level=2)

info_box(doc, '\U0001f4cb All prerequisites in one block:',
    'Phases 1\u20133 complete (ChromaDB running + AI Core BYOM endpoint RUNNING)  \u00b7  '
    'Node.js 18+ (nodejs.org)  \u00b7  '
    'npm install -g @sap/cds-dk  \u00b7  '
    'npm install -g mbt  \u00b7  '
    'Cloud Foundry CLI cf (github.com/cloudfoundry/cli/releases)  \u00b7  '
    'SAP BTP subaccount with Cloud Foundry enabled  \u00b7  '
    'XSUAA service instance (BTP \u2192 Service Marketplace \u2192 Authorization & Trust Management)  \u00b7  '
    'Destination service instance + CODESAGE_AICORE destination configured in Phase 3')

p(doc, 'Step 1 \u2014 Install developer tools (one-time)', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Verify Node.js after install
node --version    # Expected: v18.x.x+
npm  --version    # Expected: 9.x.x+

npm install -g @sap/cds-dk
cds --version     # Expected: @sap/cds: 7.x.x

npm install -g mbt
mbt --version     # Expected: 1.2.x

# Install cf CLI from github.com/cloudfoundry/cli/releases (Windows .exe installer)
cf --version      # Expected: cf version 8.x.x""")

p(doc, 'Step 2 \u2014 Install Node.js dependencies and test locally', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cd C:\\codesage\\phase4_runtime\\codesage-agent
npm install

# Set local env vars for testing
$env:CHROMADB_URL           = "http://localhost:8000"
$env:AI_CORE_DEPLOYMENT_ID  = "<your-deployment-id>"

cds watch
# [cds] - serving CodeSageService at http://localhost:4004/codesage""")

p(doc, 'Test in a second PowerShell window while cds watch is running:', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""$body = @{ question = "How do we validate vendor payment terms?" } | ConvertTo-Json
Invoke-RestMethod -Method Post -Uri "http://localhost:4004/codesage/query" `
  -Headers @{ "Content-Type" = "application/json" } -Body $body
# Expected: answer + sources + latency_ms""")

p(doc, 'Step 3 \u2014 Log in to BTP Cloud Foundry and create service instances', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""cf login -a https://api.cf.<your-region>.hana.ondemand.com
# Enter BTP email and password, select your org and space

cf create-service xsuaa        application  codesage-xsuaa        -c xs-security.json
cf create-service destination  lite         codesage-destination

cf services
# Both services should show: create succeeded""")

p(doc, 'Step 4 \u2014 Build MTA archive and deploy to BTP', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""mbt build
# Creates: .\\mta_archives\\codesage-agent_1.0.0.mtar

cf deploy mta_archives\\codesage-agent_1.0.0.mtar
# Takes 3-5 minutes. Ends with:
# Application "codesage-agent" started at:
# https://codesage-agent-<id>.cfapps.<region>.hana.ondemand.com""")

p(doc, 'Step 5 \u2014 Point deployed agent to on-premise ChromaDB', bold=True, sa=2)
run_label(doc, 'powershell')
code_block(doc,
"""# Use SAP Cloud Connector to expose ChromaDB port to BTP, then:
cf set-env codesage-agent CHROMADB_URL           "https://virtual-chromadb-host:8000"
cf set-env codesage-agent AI_CORE_DEPLOYMENT_ID  "<your-deployment-id>"
cf restage codesage-agent

cf app codesage-agent
# Expected: requested state: started  instances: 1/1""")

p(doc, 'Step 6 \u2014 End-to-end live test', bold=True, sa=2)
run_label(doc, 'python')
code_block(doc,
"""# test_live.py  -- python test_live.py
import requests
token_resp = requests.post("<xsuaa-url>/oauth/token",
    data={"grant_type": "client_credentials"},
    auth=("<client-id>", "<client-secret>"))
token = token_resp.json()["access_token"]

response = requests.post(
    "https://codesage-agent-<id>.cfapps.<region>.hana.ondemand.com/codesage/query",
    headers={"Authorization": f"Bearer {token}", "Content-Type": "application/json"},
    json={"question": "Do we have a function module for vendor payment term validation?"})
data = response.json()
print("Answer:",   data["answer"][:200])
print("Sources:",  data["sources"])
print("Latency:", data["latency_ms"], "ms")
# Expected:
# Answer: Yes \u2014 Z_VALIDATE_VENDOR_PAYTERMS validates vendor payment terms against T052...
# Sources: [{'object_name': 'Z_VALIDATE_VENDOR_PAYTERMS', 'chunk_type': 'FUNCTION'}]
# Latency: 2841 ms""")

p(doc, 'Step 7 \u2014 Register as a Fiori Launchpad tile (optional)', bold=True, sa=2)
numbered(doc, 'In BTP Cockpit \u2192 HTML5 Applications, upload the ui/webapp folder.')
numbered(doc, 'In SAP Fiori Launchpad Configuration Cockpit, add a new tile pointing to the HTML5 app URL.')
numbered(doc, 'Assign the tile to a role collection mapped to CodeSage.User. Developers see the tile when they open Fiori Launchpad.')

# ── Outcome ───────────────────────────────────────────────────────────────────
outcome_box(doc,
    'A developer types a question in plain English and receives a contextualised ABAP answer '
    'in 2\u20135 seconds, citing real function modules and programs by name. '
    'Every query is authenticated via BTP XSUAA, grounded by ChromaDB retrieval, and '
    'generated by the fine-tuned LLaMA-3 on SAP AI Core. '
    'No source code leaves the SAP BTP perimeter. Developers stop reinventing \u2014 they reuse.')

divider(doc)
doc.add_page_break()

# =============================================================================
# CLOSING SUMMARY
# =============================================================================
doc.add_heading('You Have Built CodeSage', level=1)

p(doc, 'If you followed all four phases, you now have a production-grade AI assistant '
       'that knows your SAP codebase better than any new team member ever could on day one. '
       'Here is a quick recap of what each phase produced:')

make_table(doc,
    ['Phase', 'What You Built', 'Technology', 'Runs'],
    [
        ['1 \u2014 Scan & Extract',  'ABAP source files from your live SAP system',          'pyrfc + Python',                      'Windows PC (read-only RFC)'],
        ['2 \u2014 Index & Embed',   'Semantic vector index over your entire codebase',       'ChromaDB + nomic-embed-text + Ollama', 'On-premise, CPU only'],
        ['3 \u2014 Fine-Tune',       'LLaMA-3 trained on your code, serving on AI Core',      'QLoRA + HuggingFace + SAP BYOM',      'Kaggle free GPU + SAP cloud'],
        ['4 \u2014 Runtime Query',   'Developer-facing query agent accessible from Fiori',   'BTP CAP + Node.js + XSUAA',           'SAP BTP Cloud Foundry'],
    ],
    col_widths=[1.8, 2.4, 2.2, 2.3]
)

p(doc, 'Keeping it up to date:', bold=True, sa=2)
bullet(doc, 'Re-run Phase 1 monthly (Windows Task Scheduler) to capture new ABAP objects.', bold_prefix='Phase 1:  ')
bullet(doc, 'Re-run Phase 2 after each Phase 1 refresh \u2014 ChromaDB updates automatically.', bold_prefix='Phase 2:  ')
bullet(doc, 'Re-run Phase 3 quarterly or after a major release wave adds significant new code patterns.', bold_prefix='Phase 3:  ')
bullet(doc, 'Phase 4 (BTP agent) requires no re-deployment unless you change the service logic.', bold_prefix='Phase 4:  ')

info_box(doc, '\u2b50 The result:',
    'A developer who joined your organisation today can ask "how do we handle goods '
    'receipt restriction checks?" and receive a precise, cited answer grounded in your '
    'actual ZCL_IMPL_MB_GR_CHECK BAdI implementation \u2014 in under 3 seconds. '
    'Institutional knowledge no longer leaves with the developer who wrote the code.')

doc.save(OUTPUT)
print(f'Saved: {OUTPUT}')
